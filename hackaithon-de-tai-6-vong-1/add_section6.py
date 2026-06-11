#!/usr/bin/env python3
"""Task 6: Add section 5. TÍNH ĐỔI MỚI & KHÁC BIỆT to proposal.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 18, 2: 16, 3: 14}.get(level, 14)
    color = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def make_table(doc, headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            set_font(run, size=10)
            if ri % 2 == 1:
                shading = parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w'))
                cell._tc.get_or_add_tcPr().append(shading)


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '5. TÍNH ĐỔI MỚI & KHÁC BIỆT')

    add_heading(doc, '5.1 So sánh với giải pháp hiện tại', level=2)
    add_para(doc,
        'Để chứng minh tính khác biệt, chúng tôi so sánh VoiceOne với các giải pháp '
        'chatbot DVC/Zalo hiện đang được sử dụng tại các bộ phận một cửa:', indent=True)

    make_table(doc,
        ['Tiêu chí', 'Chatbot DVC/Zalo hiện tại', 'VoiceOne'],
        [
            ['Tương tác chính', 'Text (gõ bàn phím)', 'Voice (giọng nói) — không cần gõ'],
            ['Hỗ trợ người già/KT', 'Hạn chế — cần gõ và đọc chữ', 'Có — giao tiếp hoàn toàn bằng giọng nói'],
            ['Xác thực danh tính', 'Thủ công (OTP/SĐT)', 'AI — eKYC + Face Compare tự động'],
            ['Điền form tự động', 'Không — người dân tự nhập', 'Có — OCR + auto-fill từ CCCD'],
            ['Đo lường hài lòng', 'Khảo sát giấy thụ động', 'Real-time — AI phân tích cảm xúc'],
            ['Tiếng địa phương', 'Không hỗ trợ', 'Có — SmartVoice STT đa vùng miền'],
            ['Kênh tương tác', 'Web + Zalo', 'Kiosk + Web + Mobile'],
        ])
    add_para(doc,
        'VoiceOne có 6/7 tiêu chí vượt trội, tương ứng ~85% khác biệt so với giải pháp '
        'hiện tại — vượt xa ngưỡng 30% yêu cầu của BTC.', indent=True)

    add_heading(doc, '5.2 Bốn điểm đổi mới cốt lõi', level=2)

    add_heading(doc, '5.2.1 Voice-first + Vision', level=3)
    add_para(doc,
        'Không có giải pháp nào tại Việt Nam kết hợp cả nhận dạng giọng nói (Voice) '
        'và thị giác máy tính (Vision) trong cùng một luồng nghiệp vụ cho bộ phận một cửa. '
        'VoiceOne là giải pháp đầu tiên cho phép người dân vừa nói để tra cứu, vừa được '
        'nhận diện khuôn mặt và giấy tờ tự động.')

    add_heading(doc, '5.2.2 Zero UI — Loại bỏ rào cản công nghệ', level=3)
    add_para(doc,
        'Người dân không cần chạm, không cần gõ, không cần hiểu giao diện. Chỉ cần nói — '
        'VoiceOne làm phần còn lại. Điều này đặc biệt quan trọng với người cao tuổi '
        '(chiếm ~15% dân số) và người khuyết tật.')

    add_heading(doc, '5.2.3 Vòng phản hồi tự động', level=3)
    add_para(doc,
        'Camera AI không chỉ xác thực danh tính mà còn phân tích cảm xúc khuôn mặt '
        'sau giao dịch. Dữ liệu được tổng hợp thành báo cáo hài lòng theo thời gian '
        'thực, giúp lãnh đạo ra quyết định cải tiến dựa trên bằng chứng.')

    add_heading(doc, '5.2.4 Orchestration đa API VNPT', level=3)
    add_para(doc,
        'VoiceOne tích hợp 4 API VNPT trong một luồng nghiệp vụ thống nhất: '
        'SmartVoice (STT/TTS) + Smartbot (NLP) + eKYC (OCR/Compare/Liveness) + '
        'SmartVision (Face/Sentiment). Đây là sự kết hợp chưa từng có.')

    doc.save(doc_path)
    print(f'✅ Task 6: Section "Đổi mới & Khác biệt" added')


if __name__ == '__main__':
    main()
