from docx import Document
import os
doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposal.docx')
doc = Document(doc_path)
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')
print(f'Total sections: {len(doc.sections)}')
print()
print('=== Content ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'P{i}: {text[:150]}')
print()
print('=== Table content ===')
for t_idx, table in enumerate(doc.tables):
    print(f'Table {t_idx}:')
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip()[:40] for cell in row.cells]
        print(f'  Row {r_idx}: {cells}')
