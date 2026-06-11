#!/usr/bin/env python3
"""Add section 1. Đặt vấn đề to existing proposal.docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return p


import docx.enum.text


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    
    doc = Document(doc_path)
    
    # ---- Add page break after cover page ----
    p_break = add_page_break(doc)
    
    # ---- Section heading: 1. Đặt vấn đề ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('1. ĐẶT VẤN ĐỀ')
    set_font(run, size=18, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)
    
    # ---- 1.1 Bối cảnh ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('1.1 Bối cảnh')
    set_font(run, size=16, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)  # ~0.5 inch
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        'Chuyển đổi số hành chính công là nhiệm vụ trọng tâm của Chính phủ giai đoạn 2026-2030. '
        'Cổng Dịch vụ công Quốc gia đã đạt hơn 4.000 thủ tục hành chính trực tuyến, nhưng tỷ lệ '
        'người dân sử dụng còn thấp (~30%) do rào cản công nghệ và giao diện phức tạp.'
    )
    set_font(run, size=14)
    
    # ---- 1.2 Ba Pain-point chính ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('1.2 Ba Pain-point chính')
    set_font(run, size=16, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    
    # Create the pain-point table
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table style
    table.style = 'Table Grid'
    
    # Column widths: # narrow, Pain-point medium, Minh chứng wide, Đối tượng narrow
    col_widths = [Cm(0.8), Cm(4.5), Cm(8.0), Cm(3.0)]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width
    
    # Header row
    headers = ['#', 'Pain-point', 'Minh chứng', 'Đối tượng chịu ảnh hưởng']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Set cell background to blue
        shading_elm = parse_xml(
            '<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    data = [
        [
            'PP1',
            'Ngôn ngữ hành chính phức tạp, khó tra cứu',
            '65% người >60 tuổi không tự tra cứu được thủ tục online (Bộ TT&TT 2025)',
            'Người già, người khuyết tật'
        ],
        [
            'PP2',
            'Số hóa hồ sơ chưa triệt để, nhập liệu thủ công',
            'Mỗi giao dịch mất 20-30 phút nhập liệu + kiểm tra giấy tờ (UBND TP.HCM 2025)',
            'Cán bộ một cửa'
        ],
        [
            'PP3',
            'Cán bộ một cửa quá tải, hướng dẫn lặp lại',
            '1 cán bộ tiếp ~50-70 lượt/ngày, 60% là hướng dẫn thủ tục',
            'Cán bộ, người dân chờ lâu'
        ]
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            set_font(run, size=11)
            # Alternate row shading
            if row_idx % 2 == 1:
                shading_elm = parse_xml(
                    '<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Add spacing after table
    p = doc.add_paragraph()
    run = p.add_run('')
    set_font(run, size=8)
    
    # ---- 1.3 Tại sao AI là giải pháp? ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('1.3 Tại sao AI là giải pháp?')
    set_font(run, size=16, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        'Ba pain-point trên đều có thể giải quyết bằng AI:'
    )
    set_font(run, size=14)
    
    # Bullet points
    bullets = [
        ('PP1 → ', 'Xử lý ngôn ngữ tự nhiên (NLP): Voice + Smartbot giúp người dân nói thay vì gõ'),
        ('PP2 → ', 'Thị giác máy tính (Computer Vision): OCR + eKYC tự động nhận dạng giấy tờ'),
        ('PP3 → ', 'Tự động hóa quy trình: AI xử lý câu hỏi lặp lại, giảm tải 40%'),
    ]
    for bold_part, normal_part in bullets:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.line_spacing = 1.5
        # Add bullet symbol
        run = p.add_run('• ')
        set_font(run, size=14, bold=True)
        run = p.add_run(bold_part)
        set_font(run, size=14, bold=True)
        run = p.add_run(normal_part)
        set_font(run, size=14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        'Công nghệ AI của VNPT được chọn vì: huấn luyện sẵn tiếng Việt, API sẵn sàng, đáp ứng bảo mật nhà nước.'
    )
    set_font(run, size=14)
    
    # ---- 1.4 Từ vấn đề đến giải pháp ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('1.4 Từ vấn đề đến giải pháp')
    set_font(run, size=16, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        'Xuất phát từ thực tế đó, chúng tôi đề xuất VoiceOne — trợ lý ảo đa kênh (Kiosk + Web + Mobile) '
        'cho phép người dân tương tác hoàn toàn bằng giọng nói với hệ thống dịch vụ công. VoiceOne kết hợp '
        '4 công nghệ AI cốt lõi của VNPT: Xử lý giọng nói (SmartVoice), Hiểu ngôn ngữ (Smartbot), Nhận dạng '
        'giấy tờ (eKYC/SmartReader), và Phân tích hình ảnh (SmartVision) — tạo nên một trải nghiệm không chạm, '
        'không gõ, không rào cản.'
    )
    set_font(run, size=14)
    
    # Save the document
    doc.save(doc_path)
    print(f'Saved updated proposal to {doc_path}')


if __name__ == '__main__':
    main()
