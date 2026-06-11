#!/usr/bin/env python3
"""Create proposal.docx with cover page for VoiceOne team."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def create_proposal():
    doc = Document()

    # Set page size to A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # ---- HEADER with logo ----
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add logo image
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', 'logo-team.png')
    if os.path.exists(logo_path):
        run = header_para.add_run()
        run.add_picture(logo_path, width=Cm(2.54), height=Cm(2.54))  # 100x100px ~ 2.54cm

    # ---- COVER PAGE content ----
    # Add some vertical space at top
    for _ in range(3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('')
        run.font.size = Pt(14)

    # Title line 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DỰ THI HACKATHON ĐỔI MỚI SÁNG TẠO 2026')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Spacing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('')
    run.font.size = Pt(14)

    # Topic
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Đề tài 6: Ứng dụng trí tuệ nhân tạo (AI) nhằm nâng cao\nnăng suất xử lý hồ sơ, thủ tục hành chính\ncho cơ quan nhà nước')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0066CC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run('')
    run.font.size = Pt(14)

    # Spacing
    for _ in range(2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('')
        run.font.size = Pt(14)

    # Product name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VoiceOne')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Trợ lý giọng nói thông minh cho bộ phận một cửa')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Spacing
    for _ in range(2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('')
        run.font.size = Pt(14)

    # Table info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Bảng thi: Bảng B (Challenger)')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Đội thi: ')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run = p.add_run('______________________')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Spacing
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('')
    run.font.size = Pt(12)

    # Members section
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Thành viên:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # 5 member lines
    for i in range(5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{i+1}. ____________________________ — ___________________')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Spacing
    for _ in range(2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('')
        run.font.size = Pt(14)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Ngày nộp: 16/06/2026')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposal.docx')
    doc.save(output_path)
    print(f"Proposal saved to {output_path}")


if __name__ == '__main__':
    create_proposal()
