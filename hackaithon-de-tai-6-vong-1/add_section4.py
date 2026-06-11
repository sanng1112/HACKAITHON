#!/usr/bin/env python3
"""Task 4: Add section 3. Architecture & Wireframe to proposal.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 18, 2: 16}.get(level, 14)
    color = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s>'
            '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/>'
            '</w:pBdr>' % nsdecls('w')
        )
        pPr.append(pBdr)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def add_image(doc, img_path, width_cm=15, caption=None):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        if caption:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(caption)
            set_font(run2, size=11, bold=True, color=RGBColor(0x66, 0x66, 0x66))
    else:
        print(f'  ⚠️ Image not found: {img_path}')


def add_layer(doc, bold_part, normal_part):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(bold_part)
    set_font(run, size=13, bold=True)
    run = p.add_run(normal_part)
    set_font(run, size=13)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('• ')
    set_font(run, size=13)
    run = p.add_run(text)
    set_font(run, size=13)


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    assets_dir = os.path.join(script_dir, 'assets')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '3. THIẾT KẾ TỔNG QUAN')

    add_heading(doc, '3.1 Kiến trúc hệ thống', level=2)
    add_para(doc,
        'VoiceOne được thiết kế theo mô hình microservices 4 tầng, đảm bảo khả năng '
        'mở rộng và bảo trì độc lập giữa các thành phần. Mỗi tầng đảm nhận một vai trò '
        'riêng biệt và giao tiếp qua RESTful API/WebSocket.')

    for b, n in [
        ('Tầng 1 — User Layer:', ' Đa kênh: Kiosk cảm ứng, Web App, Mobile App. Voice-first.'),
        ('Tầng 2 — AI Core (VNPT APIs):', ' SmartVoice STT/TTS, Smartbot, eKYC OCR/Liveness/Compare, SmartVision.'),
        ('Tầng 3 — Processing Layer:', ' Voice Gateway, Intent Engine, Document Processor, Sentiment Analyzer.'),
        ('Tầng 4 — Data Layer:', ' PostgreSQL (giao dịch), Redis (cache), Knowledge Base (thủ tục HC).'),
    ]:
        add_layer(doc, b, n)

    arch_img = os.path.join(assets_dir, 'architecture-diagram.png')
    add_image(doc, arch_img, width_cm=16, caption='Figure 1: Sơ đồ kiến trúc tổng thể VoiceOne')

    add_heading(doc, '3.2 Giao diện người dùng', level=2)
    add_para(doc,
        'Giao diện VoiceOne được thiết kế theo nguyên tắc tối giản, font lớn, '
        'tương phản cao — phù hợp với người già và người khiếm thị.')

    for t in [
        'Tương phản màu tối thiểu 4.5:1 (đạt chuẩn WCAG AA)',
        'Font sans-serif ≥ 18px, dễ đọc trên màn hình cảm ứng',
        'Nút bấm tối thiểu 48x48px, tránh chạm nhầm',
        'Hỗ trợ phím tắt và điều khiển hoàn toàn bằng giọng nói',
        'Dashboard cán bộ: card KPI + biểu đồ xu hướng + bảng real-time',
    ]:
        add_bullet(doc, t)

    wireframe_img = os.path.join(assets_dir, 'wireframe-voice-interface.png')
    add_image(doc, wireframe_img, width_cm=16, caption='Figure 2: Wireframe giao diện VoiceOne')
    userflow_img = os.path.join(assets_dir, 'user-flow.png')
    add_image(doc, userflow_img, width_cm=16, caption='Figure 3: Sơ đồ luồng người dùng')

    add_heading(doc, '3.3 Luồng xử lý nghiệp vụ', level=2)
    add_para(doc,
        'Luồng xử lý: Người dân đến Kiosk → Camera phát hiện → Chào giọng nói → '
        'Người dân nói yêu cầu → STT → Smartbot nhận diện ý định → '
        'Xử lý nghiệp vụ → TTS phản hồi → Kết thúc. Toàn bộ được log đầy đủ.')

    doc.save(doc_path)
    print(f'✅ Task 4: Section "Architecture & Wireframe" added')


if __name__ == '__main__':
    main()
