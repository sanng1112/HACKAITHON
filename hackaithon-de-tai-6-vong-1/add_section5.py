#!/usr/bin/env python3
"""Task 5: Add section 4. TÍNH KHẢ THI to proposal.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 18, 2: 16, 3: 14}.get(level, 14)
    color = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def add_para(doc, text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=14, bold=True)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('• ')
    set_font(run, size=13)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=13, bold=True)
    run = p.add_run(text)
    set_font(run, size=13)


def make_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            set_font(run, size=10)
            if ri % 2 == 1:
                shading = parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w'))
                cell._tc.get_or_add_tcPr().append(shading)


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '4. TÍNH KHẢ THI')

    # 4.1 Nguồn dữ liệu
    add_heading(doc, '4.1 Nguồn dữ liệu', level=2)
    add_para(doc,
        'VoiceOne không yêu cầu dữ liệu huấn luyện riêng — các API VNPT đã được '
        'huấn luyện sẵn trên dữ liệu tiếng Việt. Dữ liệu thủ tục hành chính được '
        'lấy từ Cổng DVC Quốc gia (dữ liệu mở, cập nhật thường xuyên).', indent=True)

    make_table(doc,
        ['Yếu tố', 'Mô tả'],
        [
            ['Dữ liệu huấn luyện', 'API VNPT có sẵn, không cần train thêm.'],
            ['Dữ liệu vận hành', 'Người dùng cung cấp — bảo mật theo Nghị định 13/2023.'],
            ['Dữ liệu thủ tục HC', 'Crawl từ Cổng DVC Quốc gia (api.dichvucong.gov.vn) — dữ liệu mở.'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 4.2 Nhân lực
    add_heading(doc, '4.2 Nhân lực', level=2)
    make_table(doc,
        ['Vai trò', 'Số lượng', 'Kỹ năng chính'],
        [
            ['Project Manager', '1', 'Agile/Scrum, quản lý rủi ro'],
            ['AI Developer', '1', 'Python, REST API, Xử lý ngôn ngữ tự nhiên'],
            ['Fullstack Developer', '1', 'React/Vue, Node.js, PostgreSQL, Docker'],
            ['UI/UX Designer', '1', 'Figma, thiết kế cho người già/khuyết tật'],
            ['Business Analyst', '1', 'Nghiệp vụ hành chính công'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 4.3 Kiến trúc kỹ thuật
    add_heading(doc, '4.3 Kiến trúc kỹ thuật', level=2)
    add_para(doc,
        'Frontend: React (Next.js) + TypeScript, tối ưu cho Kiosk cảm ứng và Web. '
        'Backend: Node.js (Express) / Python (FastAPI) — RESTful API + WebSocket '
        'cho real-time voice streaming. AI Services: kết nối API VNPT qua HTTP/gRPC, '
        'wrapper service layer riêng. Database: PostgreSQL + Redis cache. '
        'DevOps: Docker → AWS EC2/VNPT Cloud, CI/CD GitHub Actions.', indent=True)

    # 4.4 Kế hoạch MVP 7 ngày
    add_heading(doc, '4.4 Kế hoạch MVP 7 ngày (Vòng 2)', level=2)
    make_table(doc,
        ['Ngày', 'Công việc', 'Kết quả'],
        [
            ['1-2', 'Setup dự án + Tích hợp SmartVoice STT/TTS', 'Ghi âm → text + TTS cơ bản'],
            ['3-4', 'Tích hợp Smartbot + eKYC OCR', 'Nhận diện ý định + scan CCCD'],
            ['5-6', 'Xây UI Kiosk + Dashboard + Luồng hội thoại', 'Giao diện + luồng hoàn chỉnh'],
            ['7', 'End-to-end test + Fix bug + Đóng gói', 'MVP deploy được'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 4.5 Chi phí vận hành
    add_heading(doc, '4.5 Chi phí vận hành', level=2)
    make_table(doc,
        ['Hạng mục', 'Chi phí/tháng (VNĐ)', 'Ghi chú'],
        [
            ['Server (2 VPS 4GB RAM)', '~1.000.000', 'AWS EC2 t3.medium / VNPT Cloud'],
            ['API VNPT', '~500.000 - 2.000.000', 'Tùy số lượng request'],
            ['Domain + SSL', '~200.000', '.gov.vn hoặc .vn'],
            ['DevOps tools', 'Miễn phí', 'GitHub Free / Docker Free'],
            ['Tổng vận hành', '~1.700.000 - 3.200.000', '~$70-130/tháng'],
        ])
    add_para(doc,
        'Chi phí setup Kiosk phần cứng: 15-25 triệu đồng (màn hình cảm ứng 22" + case + camera). '
        'So với giải pháp tư vấn CNTT truyền thống (50-100 triệu/tháng), VoiceOne tiết kiệm ≥90%.', indent=True)

    # 4.6 An toàn bảo mật & Pháp lý
    add_heading(doc, '4.6 An toàn bảo mật & Pháp lý', level=2)
    make_table(doc,
        ['Yêu cầu', 'Giải pháp'],
        [
            ['Bảo vệ dữ liệu cá nhân', 'Nghị định 13/2023/NĐ-CP — AES-256, TLS 1.3'],
            ['Xác thực điện tử', 'Nghị định 59/2022/NĐ-CP — eKYC Liveness detection'],
            ['An toàn thông tin', 'Luật An toàn TT mạng 2015 — Audit log, phân quyền'],
            ['Giao dịch điện tử', 'Luật Giao dịch điện tử 2005'],
            ['Minh bạch', 'Log phiên giao dịch (audio + text + kết quả)'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 4.7 Lộ trình phát triển
    add_heading(doc, '4.7 Lộ trình phát triển', level=2)
    roadmap_items = [
        ('Tháng 1-2:', ' MVP → Pilot tại 1-2 UBND phường'),
        ('Tháng 3-4:', ' Feedback → Cải tiến → Scale lên quận/huyện'),
        ('Tháng 5-6:', ' Tích hợp Cổng DVC Quốc gia → Public beta'),
        ('Tháng 7-12:', ' Mở rộng tỉnh → Hợp tác VNPT'),
    ]
    for b, n in roadmap_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(b)
        set_font(run, size=13, bold=True)
        run = p.add_run(n)
        set_font(run, size=13)

    doc.save(doc_path)
    print(f'✅ Task 5: Section "Tính khả thi" added')


if __name__ == '__main__':
    main()
