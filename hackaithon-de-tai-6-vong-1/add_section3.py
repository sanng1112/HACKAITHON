#!/usr/bin/env python3
"""Task 3: Add section 2. GIẢI PHÁP to proposal.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 18, 2: 16}.get(level, 14)
    color = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s>'
            '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/>'
            '</w:pBdr>' % nsdecls('w')
        )
        pPr.append(pBdr)


def add_para(doc, text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=14, bold=True)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('• ')
    set_font(run, size=14)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_font(run, size=14, bold=True)
    run = p.add_run(text)
    set_font(run, size=14)


def add_step(doc, bold_part, normal_part):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(bold_part)
    set_font(run, size=13, bold=True)
    run = p.add_run(normal_part)
    set_font(run, size=13)


def add_component(doc, bold_part, normal_part):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('• ')
    set_font(run, size=13)
    run = p.add_run(bold_part)
    set_font(run, size=13, bold=True)
    run = p.add_run(normal_part)
    set_font(run, size=13)


def make_table(doc, headers, data, col_aligns=None):
    """Create a formatted table with header shading."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            if col_aligns and ci < len(col_aligns):
                p.alignment = col_aligns[ci]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            set_font(run, size=11)
            if ri % 2 == 1:
                shading = parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w'))
                cell._tc.get_or_add_tcPr().append(shading)
    return table


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '2. GIẢI PHÁP')

    add_heading(doc, '2.1 Tổng quan giải pháp', level=2)
    add_para(doc,
        'VoiceOne là một trợ lý ảo đa kênh (Kiosk + Web + Mobile), '
        'cho phép người dân tương tác hoàn toàn bằng giọng nói với hệ thống dịch vụ công. '
        'Khác với chatbot text hiện tại, VoiceOne hiểu và trả lời bằng giọng nói tiếng Việt, '
        'hướng dẫn thủ tục từng bước, tự động điền thông tin từ CCCD qua camera, '
        'và đánh giá mức độ hài lòng sau giao dịch.')
    add_para(doc,
        'Hệ thống được thiết kế theo nguyên lý "Zero UI" — loại bỏ hoàn toàn rào cản '
        'giao diện phức tạp, giúp bất kỳ người dân nào, dù không biết sử dụng công nghệ, '
        'cũng có thể tự thực hiện giao dịch hành chính công một cách dễ dàng.')

    add_heading(doc, '2.2 Tính năng cốt lõi', level=2)
    add_para(doc, 'VoiceOne tích hợp 5 tính năng chính, tận dụng tối đa API VNPT:', indent=True)
    make_table(doc,
        ['Tính năng', 'Mô tả', 'API VNPT'],
        [
            ['Voice Tra cứu', 'Người dân nói → STT → Smartbot xử lý → TTS trả lời', 'SmartVoice (STT,TTS), Smartbot'],
            ['Voice Khai báo', 'Người dân nói → STT → tự động tạo đơn yêu cầu', 'SmartVoice (STT)'],
            ['Scan & Auto-fill', 'Scan CCCD → OCR bóc tách → tự động điền form', 'eKYC (OCR)'],
            ['Xác thực danh tính', 'So sánh khuôn mặt với ảnh trên CCCD', 'eKYC (Compare, Liveness)'],
            ['Đánh giá hài lòng', 'Camera phân tích cảm xúc → báo cáo real-time', 'SmartVision (face)'],
        ],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    add_heading(doc, '2.3 Kịch bản người dùng: Câu chuyện Ông A', level=2)
    add_para(doc,
        'Để minh họa rõ cách VoiceOne vận hành, chúng tôi xin giới thiệu kịch bản '
        'điển hình: Ông Nguyễn Văn A (65 tuổi) đến UBND phường làm thủ tục xác nhận '
        'tình trạng hôn nhân. Toàn bộ quy trình gồm 8 bước:', indent=True)
    for b, n in [
        ('Bước 1 — Phát hiện:', ' Camera tại Kiosk VoiceOne phát hiện Ông A → Hệ thống phát giọng chào.'),
        ('Bước 2 — Tra cứu:', ' Ông A nói yêu cầu → STT → Smartbot nhận diện ý định.'),
        ('Bước 3 — Hướng dẫn:', ' Smartbot xác định thủ tục → TTS hướng dẫn đưa CCCD vào khay scan.'),
        ('Bước 4 — Scan:', ' eKYC OCR tự động nhận dạng, bóc tách thông tin CCCD.'),
        ('Bước 5 — Xác thực:', ' eKYC Compare + Liveness: so sánh khuôn mặt + kiểm tra người thật.'),
        ('Bước 6 — Xác nhận:', ' TTS đọc lại thông tin, Ông A xác nhận → tự động điền form.'),
        ('Bước 7 — Kết quả:', ' Hệ thống kiểm tra hợp lệ → TTS thông báo kết quả + mã hồ sơ.'),
        ('Bước 8 — Đo lường:', ' SmartVision phân tích cảm xúc → ghi nhận hài lòng → dashboard.'),
    ]:
        add_step(doc, b, n)
    add_para(doc,
        'Toàn bộ quy trình chỉ mất 5-7 phút, giảm 70% thời gian so với 20-30 phút '
        'theo cách truyền thống. Ông A không cần chạm màn hình hay gõ phím.')

    add_heading(doc, '2.4 Vai trò các thành phần AI', level=2)
    add_para(doc, 'VoiceOne sử dụng 5 thành phần AI cốt lõi từ VNPT:', indent=True)
    for b, n in [
        ('STT (SmartVoice Speech-to-Text): ', 'Chuyển giọng nói tiếng Việt thành văn bản. Hỗ trợ giọng địa phương.'),
        ('TTS (SmartVoice Text-to-Speech): ', 'Chuyển văn bản thành giọng nói tự nhiên, thân thiện.'),
        ('Smartbot (NLP/Intent): ', 'Nhận diện ý định, tra cứu thủ tục, xử lý hội thoại đa lượt.'),
        ('eKYC OCR + SmartReader:', ' Nhận dạng và bóc tách thông tin giấy tờ, tự động điền form.'),
        ('eKYC Face Recognition:', ' So sánh khuôn mặt (Compare), phát hiện người thật (Liveness), phân tích cảm xúc.'),
    ]:
        add_component(doc, b, n)

    doc.save(doc_path)
    print(f'✅ Task 3: Section "Giải pháp" added')


if __name__ == '__main__':
    main()
