#!/usr/bin/env python3
"""Task 7: Add section 6. TÁC ĐỘNG DỰ KIẾN to proposal.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    size = {1: 18, 2: 16}.get(level, 14)
    color = RGBColor(0x00, 0x66, 0xCC)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def add_para(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def make_table(doc, headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            set_font(run, size=10)
            if ri % 2 == 1:
                shading = parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w'))
                cell._tc.get_or_add_tcPr().append(shading)


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '6. TÁC ĐỘNG DỰ KIẾN')

    # 6.1 TAM-SAM-SOM
    add_heading(doc, '6.1 Phân tích thị trường (TAM-SAM-SOM)', level=2)
    make_table(doc,
        ['Chỉ số', 'Giá trị', 'Cách tính', 'Nguồn'],
        [
            ['TAM', '~15.000 tỷ VNĐ', 'Chi tiêu CNTT HC công 63 tỉnh, ~240 tỷ/tỉnh/năm', 'Bộ TT&TT 2025'],
            ['SAM', '~500 tỷ VNĐ', 'Mảng AI + tự động hóa cho bộ phận một cửa (3-5% TAM)', 'Phân tích nội bộ'],
            ['SOM', '~25 tỷ VNĐ', '5% SAM trong 2 năm đầu (~50-100 UBND)', 'Dự báo thận trọng'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 6.2 Lợi ích xã hội
    add_heading(doc, '6.2 Lợi ích xã hội', level=2)
    make_table(doc,
        ['Lợi ích', 'Chỉ số', 'Giải thích'],
        [
            ['Giảm thời gian GD', 'Giảm 70%', 'Từ 20-30 phút xuống 5-7 phút'],
            ['Người già/KT tự GD', 'Tăng độ phủ 95%', 'Voice + Zero UI giúp nhóm yếu thế'],
            ['Giảm tải cán bộ', 'Giảm 40%', 'AI xử lý 60% câu hỏi lặp lại'],
            ['Minh bạch hóa', '100% GD log', 'Audio + text + video — chống tiêu cực'],
            ['Tăng hài lòng', '72% → 90%', 'Nhờ giảm thời gian chờ + hỗ trợ tận tình'],
        ])

    p = doc.add_paragraph()
    set_font(p.add_run(''), size=6)

    # 6.3 Mô hình doanh thu
    add_heading(doc, '6.3 Mô hình doanh thu (B2G Subscription)', level=2)
    make_table(doc,
        ['Gói', 'Giá (VNĐ/tháng)', 'Dịch vụ'],
        [
            ['Basic', '5.000.000', '1 cửa, 500 giao dịch/tháng, hỗ trợ 8h/ngày'],
            ['Pro', '15.000.000', 'Đa cửa (tối đa 5), không giới hạn GD, hỗ trợ 24/7'],
            ['Enterprise', 'Theo yêu cầu', 'Tùy chỉnh, tích hợp riêng, SLA cam kết'],
        ])
    add_para(doc,
        'Phí triển khai ban đầu: 30-50 triệu đồng/điểm (Kiosk + camera + setup). '
        'Dự báo hòa vốn trong 12 tháng với 20 khách hàng Gói Basic (doanh thu ~100 triệu/tháng). '
        'ROI 3 năm: ~300% với tăng trưởng 10-15 khách hàng mới mỗi quý sau năm 1.', indent=True)

    # Pricing explanation
    add_heading(doc, '6.4 Phân tích cạnh tranh', level=2)
    add_para(doc,
        'Hiện tại chưa có đối thủ cạnh tranh trực tiếp nào cung cấp giải pháp voice-first '
        'tích hợp đa API VNPT cho bộ phận một cửa. Các đối thủ gián tiếp gồm:')
    competitors = [
        'FPT.AI Chatbot: Chỉ text, không có voice + vision, giá ~10-20 triệu/tháng',
        'Zalo OA chatbot: Text + Zalo, không tích hợp eKYC, giá ~3-5 triệu/tháng',
        'Giải pháp gia công CNTT truyền thống: Chi phí 50-100 triệu/tháng, không có AI',
    ]
    for c in competitors:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run('• ')
        set_font(run, size=13)
        run = p.add_run(c)
        set_font(run, size=13)

    doc.save(doc_path)
    print(f'✅ Task 7: Section "Tác động dự kiến" added')


if __name__ == '__main__':
    main()
