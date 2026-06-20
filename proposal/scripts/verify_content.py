#!/usr/bin/env python3
"""
Verify the generated proposal content for HackAIthon 2026 round 1 requirements.
Usage: python3 verify_content.py [path_to_docx]
"""
import os, sys
from docx import Document

def get_all_text(doc):
    """Extract text from paragraphs AND tables."""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return "\n".join(texts)

def get_all_headers_footers(doc):
    texts = []
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            texts.append(p.text)
        footer = section.footer
        for p in footer.paragraphs:
            texts.append(p.text)
    return "\n".join(texts)

def check_content(doc, full_text):
    results = []

    # ── Required sections based on contest rules ──
    sections = [
        ("Tên sản phẩm 'GovOne'", "Tên sản phẩm", "GovOne"),
        ("Chương 1: Đặt vấn đề", "Đặt vấn đề", "Đặt vấn đề"),
        ("Chương 7: Kết luận", "Kết luận", "Kết luận"),
    ]
    for name, _, keyword in sections:
        found = keyword.lower() in full_text.lower()
        results.append((name, "✅" if found else "❌"))

    # ── Scoring criteria alignment ──
    criteria = [
        ("Tính phù hợp (Pain-points)", "pain-point"),
        ("Tính đổi mới (Voice-First)", "Voice-First"),
        ("Tính khả thi (Kiến trúc 4 tầng)", "4 tầng"),
        ("Tính khả thi (MVP 7 ngày)", "7 ngày"),
        ("Tính khả thi (Chi phí vận hành)", "2.250.000"),
        ("Tác động (TAM-SAM-SOM)", "TAM"),
        ("Tác động (Mô hình doanh thu)", "B2G"),
        ("Tác động (Lợi ích xã hội)", "95%"),
        ("Chất lượng (Sơ đồ kiến trúc)", "kiến trúc tổng quan"),
        ("Chất lượng (Wireframe)", "Wireframe"),
        ("Chất lượng (Web UI Screenshot)", "web-login"),
    ]
    for name, keyword in criteria:
        found = keyword.lower() in full_text.lower()
        results.append((name, "✅" if found else "❌"))

    # ── VNPT APIs ──
    apis = ["VNPT eKYC", "SmartVoice", "Smartbot", "SmartReader", "SmartVision"]
    for api in apis:
        found = api.lower() in full_text.lower()
        results.append((f"API: {api}", "✅" if found else "❌"))

    # ── Key terms ──
    terms = [
        "Đề tài 6", "PP1", "Voice-First", "OCR", "eKYC", "Sentiment AI",
        "Dashboard", "Next.js", "FastAPI", "PostgreSQL",
        "Nghị định 13/2023", "Bảo mật", "MVP", "GTM",
    ]
    for term in terms:
        found = term.lower() in full_text.lower()
        results.append((f"Term: {term}", "✅" if found else "❌"))

    # ── Team info ──
    members = ["Nguyễn Ngọc Bình An", "Hoàng Thị Linh Hương", "Nguyễn Đoàn Nhật Minh", "Trần Hoàng Nguyên", "Phạm Lê Việt Đức"]
    for m in members:
        found = m.lower() in full_text.lower()
        results.append((f"Thành viên: {m}", "✅" if found else "❌"))

    # ── Assets ──
    assets_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets")
    assets = ["architecture-diagram.png", "wireframe-kiosk.png", "wireframe-scan.png",
              "wireframe-dashboard.png", "user-flow-citizen.png", "user-flow-officer.png",
              "web-login.png", "web-citizen-dashboard.png", "web-officer-dashboard.png"]
    for a in assets:
        found = os.path.exists(os.path.join(assets_dir, a))
        results.append((f"Asset: {a}", "✅" if found else "❌"))

    return results

def main(doc_path=None):
    if doc_path is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        project_dir = os.path.dirname(script_dir)
        doc_path = os.path.join(project_dir, "proposal.docx")

    if not os.path.exists(doc_path):
        # Try output path
        output_path = os.path.join(os.path.dirname(os.path.dirname(script_dir)), "output", "GovOne_Proposal.docx")
        if os.path.exists(output_path):
            doc_path = output_path
        else:
            print(f"❌ Cannot find proposal file. Checked: {doc_path}")
            return

    print(f"📖 Verifying: {doc_path}")
    doc = Document(doc_path)
    full_text = get_all_text(doc) + "\n" + get_all_headers_footers(doc)

    results = check_content(doc, full_text)

    passed = sum(1 for _, s in results if "✅" in s)
    failed = sum(1 for _, s in results if "❌" in s)

    print(f"\n{'='*60}")
    print(f"  VERIFICATION REPORT")
    print(f"{'='*60}")
    for name, status in results:
        print(f"  {status}  {name}")
    print(f"{'='*60}")
    print(f"  ✅ Passed: {passed}  |  ❌ Failed: {failed}  |  Total: {len(results)}")
    print(f"{'='*60}")

    return passed, failed

if __name__ == "__main__":
    doc_path = sys.argv[1] if len(sys.argv) > 1 else None
    main(doc_path)
