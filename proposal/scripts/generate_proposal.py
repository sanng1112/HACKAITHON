#!/usr/bin/env python3
"""
GovOne — Comprehensive Proposal Generator for HackAIthon 2026, Round 1 (Bảng B - Challenger)
Generates a complete .docx with proper heading styles (for PDF bookmarks)
and exports to .pdf via LibreOffice.
"""

import os, sys, subprocess, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Constants ──
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")
OUTPUT_DIR = os.path.join(os.path.dirname(BASE_DIR), "output")
DOCX_PATH = os.path.join(OUTPUT_DIR, "GovOne_Proposal.docx")
PDF_PATH = os.path.join(OUTPUT_DIR, "GovOne_Proposal.pdf")

TEAM_MEMBERS = [
    ("Nguyễn Ngọc Bình An", "Trưởng nhóm"),
    ("Hoàng Thị Linh Hương", "Thành viên"),
    ("Nguyễn Đoàn Nhật Minh", "Thành viên"),
    ("Trần Hoàng Nguyên", "Thành viên"),
    ("Phạm Lê Việt Đức", "Thành viên"),
]

# ── Helper Functions ──

def set_font(run, name="Times New Roman", size=14, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        '<w:shd %s w:fill="%s" w:val="clear"/>' % (nsdecls("w"), color_hex)
    )
    tcPr.append(shading)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def make_paragraph(doc, text, size=14, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   first_line_indent=True, spacing=1.5, space_after=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def make_heading(doc, text, level=1):
    """Add a heading using Word Heading style for proper PDF bookmarks."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    style_name = f"Heading {level}"
    if style_name in [s.name for s in doc.styles]:
        p.style = doc.styles[style_name]
    sz = {1: 18, 2: 15, 3: 13}.get(level, 13)
    r = p.add_run(text)
    set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x56, 0xA6))
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0056A6"/></w:pBdr>' % nsdecls("w")
        ))
    return p


def make_table(doc, headers, data, col_widths=None, header_color="0056A6"):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(c, header_color)
    # Data rows
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct)
            set_font(r, size=9)
            if ri % 2 == 1:
                set_cell_shading(c, "E8F0FE")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_image(doc, path, width_cm=14, caption=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = cap.add_run(caption)
            set_font(rc, size=10, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    else:
        make_paragraph(doc, f"[Hình ảnh: {os.path.basename(path)} — không tìm thấy]",
                       size=11, italic=True, color=RGBColor(0xFF, 0x00, 0x00),
                       first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_bullet(doc, bold_part, normal_part, indent_cm=1.5, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run("\u2022 ")
    set_font(r, size=size)
    r = p.add_run(bold_part)
    set_font(r, size=size, bold=True)
    r = p.add_run(normal_part)
    set_font(r, size=size)
    return p


def add_numbered(doc, number, bold_part, normal_part, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(f"{number}. ")
    set_font(r, size=size, bold=True)
    r = p.add_run(bold_part)
    set_font(r, size=size, bold=True)
    r = p.add_run(normal_part)
    set_font(r, size=size)
    return p


# ── Section Builders ──

def build_title_page(doc):
    """Build the cover page with team info."""
    # Spacer
    for _ in range(4):
        make_paragraph(doc, "", size=14, first_line_indent=False)

    # HackAIthon branding
    make_paragraph(doc, "HỘI SINH VIÊN VIỆT NAM", size=13, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, spacing=1.2)
    make_paragraph(doc, "BAN CHẤP HÀNH TRUNG ƯƠNG", size=13, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, spacing=1.2)

    make_paragraph(doc, "———•———", size=14,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)

    # Main title
    make_paragraph(doc, "", size=10, first_line_indent=False)
    make_paragraph(doc, "BẢNG B — CHALLENGER", size=16, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                   color=RGBColor(0x00, 0x56, 0xA6), spacing=1.5)

    make_paragraph(doc, "", size=10, first_line_indent=False)
    make_paragraph(doc, "BÀI DỰ THI VÒNG 1", size=18, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                   color=RGBColor(0x00, 0x38, 0x73), spacing=1.5)

    make_paragraph(doc, "", size=10, first_line_indent=False)

    # Product name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GovOne")
    set_font(run, name="Times New Roman", size=36, bold=True, color=RGBColor(0x00, 0x56, 0xA6))

    make_paragraph(doc, "Hệ thống Quản lý Hành chính Công Thông minh\nTích hợp AI — Voice-First — OCR — Sentiment", size=14,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, color=RGBColor(0x44, 0x44, 0x44))

    # Divider
    make_paragraph(doc, "", size=6, first_line_indent=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("─" * 50)
    set_font(r, size=10, color=RGBColor(0x00, 0x56, 0xA6))

    make_paragraph(doc, "", size=6, first_line_indent=False)

    # Team info
    make_paragraph(doc, "ĐỀ TÀI 6: AI giúp cơ quan nhà nước nâng cao năng suất xử lý hồ sơ,\n"
                        "tăng tính minh bạch, cải thiện dịch vụ và sự hài lòng của người dân,\n"
                        "hướng tới hiệu quả quản trị điều hành",
                   size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                   first_line_indent=False, color=RGBColor(0x00, 0x38, 0x73))

    make_paragraph(doc, "", size=6, first_line_indent=False)

    # Team table
    make_table(doc,
               ["Thành viên", "Vai trò"],
               [[name, role] for name, role in TEAM_MEMBERS],
               header_color="003873")

    make_paragraph(doc, "", size=10, first_line_indent=False)

    # Footer info on title page
    now = __import__("datetime").datetime.now()
    date_str = now.strftime("Ngày %d tháng %m năm %Y")
    make_paragraph(doc, date_str, size=12, bold=False,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                   color=RGBColor(0x66, 0x66, 0x66))

    make_paragraph(doc, "HackAIthon 2026 — Vietnamese Student Hackathon", size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                   color=RGBColor(0x88, 0x88, 0x88))


def build_chapter1(doc):
    """Chương 1: Đặt vấn đề — Tính phù hợp (25 điểm)"""
    make_heading(doc, "Chương 1: Đặt vấn đề", 1)

    make_heading(doc, "1.1. Bối cảnh chung", 2)
    make_paragraph(doc,
        "Việt Nam đang trong giai đoạn chuyển đổi số quốc gia mạnh mẽ, với mục tiêu đến năm 2030 "
        "đưa 100% dịch vụ công trực tuyến lên mức độ 4 (toàn trình). Theo báo cáo của Bộ TT&TT năm 2025, "
        "hiện có hơn 11.000 dịch vụ công trực tuyến được triển khai trên cổng DVC Quốc gia, nhưng tỷ lệ "
        "hồ sơ nộp trực tuyến mới chỉ đạt khoảng 40%, trong đó tỷ lệ người dân thực sự hoàn thành giao dịch "
        "trực tuyến chỉ ở mức 25-30%."
    )
    make_paragraph(doc,
        "Nguyên nhân chính đến từ: (1) giao diện hành chính phức tạp, nhiều bước, ngôn ngữ khó hiểu; "
        "(2) người dân lớn tuổi, người khuyết tật gặp rào cản lớn khi tiếp cận công nghệ; "
        "(3) cán bộ xử lý hồ sơ thủ công, chịu áp lực lớn; (4) thiếu công cụ đo lường sự hài lòng một cách tự động."
    )

    make_heading(doc, "1.2. Các pain-point cốt lõi trong hành chính công", 2)

    problems = [
        ("PP1 — Rào cản giao diện hành chính: ",
         "Người dân (đặc biệt là người lớn tuổi, người khuyết tật) gặp khó khăn khi thao tác trên "
         "cổng DVC do giao diện phức tạp, nhiều thuật ngữ chuyên ngành. Khảo sát của Vụ CNTT (2025) cho thấy "
         "65% người trên 55 tuổi không thể tự hoàn thành một giao dịch trực tuyến."),
        ("PP2 — Quy trình xử lý hồ sơ thủ công: ",
         "Cán bộ tại bộ phận một cửa phải nhập liệu thủ công từ giấy tờ giấy, kiểm tra đối chiếu mất "
         "nhiều thời gian. Trung bình mỗi hồ sơ mất 20-30 phút xử lý, với tỷ lệ sai sót khoảng 15% "
         "(theo báo cáo của Bộ Nội vụ 2024)."),
        ("PP3 — Thiếu kênh tương tác đa dạng: ",
         "Các chatbot hiện tại (như trên Zalo OA) chủ yếu trả lời text, không hỗ trợ giọng nói, không thể "
         "dẫn dắt người dân qua các bước nộp hồ sơ một cách trực quan. Người dân phải đến trực tiếp UBND "
         "để được hướng dẫn."),
        ("PP4 — Khó khăn trong tra cứu lịch sử: ",
         "Khi người dân cần tra cứu hồ sơ cũ, cán bộ phải tìm kiếm trong kho lưu trữ giấy tờ vật lý, "
         "mất từ 30-60 phút mỗi lần. Nhiều xã/phường lưu trữ >100.000 bộ hồ sơ giấy, tiềm ẩn rủi ro thất lạc."),
        ("PP5 — Thiếu đo lường sự hài lòng: ",
         "Việc đo lường mức độ hài lòng của người dân chủ yếu dựa trên khảo sát giấy định kỳ, không "
         "phản ánh kịp thời chất lượng dịch vụ. Chỉ số hài lòng trung bình hiện ở mức ~65% (PAPI 2024)."),
        ("PP6 — Chi phí vận hành cao: ",
         "Mỗi UBND phường/xã chi trung bình 150-200 triệu đồng/năm cho in ấn, lưu trữ, văn phòng phẩm "
         "và nhân sự làm thủ tục hành chính. Nhân sự có trình độ CNTT tại cấp xã còn hạn chế."),
    ]
    for bold_p, normal_p in problems:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "1.3. Tại sao AI là giải pháp?", 2)
    make_paragraph(doc,
        "AI mang đến khả năng giải quyết triệt để 6 pain-point trên nhờ 4 năng lực cốt lõi: "
        "(a) Công nghệ xử lý ngôn ngữ tự nhiên NLP, hỗ trợ chuyển đổi giọng nói thành văn bản STT và "
        "tổng hợp giọng nói TTS — cho phép tương tác bằng giọng nói, xóa bỏ rào cản giao diện cho người già, "
        "người khuyết tật; (b) Công nghệ thị giác máy tính kết hợp nhận dạng ký tự quang học OCR và "
        "định danh điện tử eKYC — tự động hóa việc nhập liệu từ giấy tờ, giảm sai sót và thời gian xử lý; "
        "(c) Công nghệ học sâu — phân tích cảm xúc khuôn mặt để đo lường sự hài lòng tức thời; "
        "(d) Mô hình ngôn ngữ lớn LLM — hỗ trợ chatbot thông minh, tự động điền form và tra cứu thông tin "
        "bằng ngôn ngữ tự nhiên."
    )
    make_paragraph(doc,
        "Theo McKinsey Global Institute, AI có thể tự động hóa 45% hoạt động hành chính công, "
        "giúp tiết kiệm 2,8 tỷ giờ làm việc mỗi năm trên toàn cầu. Tại Việt Nam, ước tính ứng dụng AI "
        "trong hành chính công có thể tiết kiệm 15.000-20.000 tỷ đồng/năm."
    )

    make_heading(doc, "1.4. Từ vấn đề đến giải pháp", 2)
    make_paragraph(doc,
        "GovOne ra đời như một giải pháp tích hợp toàn diện, kết hợp 5 nền tảng AI chiến lược từ "
        "VNPT (eKYC, SmartVoice, Smartbot, SmartReader, SmartVision) và các mô hình AI tiên tiến khác "
        "để giải quyết bài toán hành chính công một cách triệt để. GovOne không phải là một API đơn lẻ, "
        "mà là một hệ sinh thái sản phẩm hoàn chỉnh — từ Kiosk voice-first tại UBND đến Web App cho cán bộ, "
        "từ Smartbot hỗ trợ online đến Dashboard phân tích real-time."
    )


def build_chapter2(doc):
    """Chương 2: Giải pháp GovOne"""
    make_heading(doc, "Chương 2: Giải pháp GovOne", 1)

    make_heading(doc, "2.1. Tổng quan giải pháp", 2)
    make_paragraph(doc,
        "GovOne là hệ thống quản lý hành chính công thông minh, áp dụng mô hình ưu tiên giọng nói "
        "Voice-First — người dân có thể tương tác với hệ thống chủ yếu bằng giọng nói, giúp xóa bỏ mọi "
        "rào cản về công nghệ. Hệ thống tích hợp 5 nhóm công nghệ AI cốt lõi: (1) Xác thực thông minh eKYC, "
        "(2) Giọng nói thông minh SmartVoice hỗ trợ chuyển đổi giọng nói thành văn bản STT và tổng hợp giọng nói TTS, "
        "(3) Trợ lý ảo Smartbot kết hợp mô hình ngôn ngữ lớn LLM, "
        "(4) Nhận dạng văn bản SmartReader OCR, (5) Phân tích cảm xúc Sentiment AI."
    )

    make_heading(doc, "2.2. Sáu tính năng cốt lõi", 2)

    features = [
        ("2.2.1. Kiosk Voice-First — Tương tác bằng giọng nói tại quầy",
         "Tích hợp công nghệ SmartVoice hỗ trợ chuyển giọng nói thành văn bản STT và tổng hợp giọng nói TTS "
         "cho phép người dân tương tác với hệ thống hoàn toàn bằng giọng nói. Hỗ trợ tiếng Việt với "
         "các giọng đọc tự nhiên (miền Bắc, Trung, Nam). Người dân chỉ cần nói nhu cầu, hệ thống tự động "
         "xác định thủ tục, hướng dẫn từng bước và xác nhận bằng giọng nói.",
         "SmartVoice STT và TTS: Chuyển đổi giọng nói và văn bản với độ chính xác trên 95%"),
        ("2.2.2. OCR & eKYC — Tự động nhập liệu từ giấy tờ",
         "Khi người dân đưa CCCD, sổ hộ khẩu, giấy tờ xe vào kiosk hoặc tải lên qua web, hệ thống "
         "SmartReader OCR tự động nhận dạng và bóc tách thông tin. eKYC xác thực danh tính qua "
         "kiểm tra thực thể sống Liveness Detection để phát hiện giấy tờ thật hay giả, so sánh khuôn mặt Compare Face với ảnh trên CCCD và phát hiện khẩu trang Mask Face Detection. Toàn bộ quy trình mất dưới 30 giây.",
         "SmartReader OCR và VNPT eKYC: Tự động hóa 100% khâu nhập liệu và xác thực"),
        ("2.2.3. Smartbot Đa kênh — Hỗ trợ thủ tục hành chính 24/7",
         "Tích hợp VNPT Smartbot với khả năng nhận dạng ý định và mô hình ngôn ngữ lớn LLM để hỏi đáp "
         "nâng cao. Bot có thể trả lời chi tiết về hơn 50 thủ tục hành chính thông dụng, hướng dẫn "
         "từng bước nộp hồ sơ, kiểm tra trạng thái. Triển khai trên Web, Kiosk, và Zalo OA.",
         "VNPT Smartbot: Nhận diện hơn 200 ý định, trả lời chính xác trên 90% câu hỏi thường gặp"),
        ("2.2.4. Sentiment AI — Đo lường hài lòng qua camera",
         "Tại quầy giao dịch, camera thông minh tích hợp giải pháp SmartVision ghi nhận biểu cảm khuôn mặt người dân "
         "sau khi hoàn thành giao dịch. Hệ thống phân tích bảy trạng thái cảm xúc bao gồm vui, hài lòng, bình thường, "
         "buồn, bực dọc, tức giận và bất ngờ để đánh giá mức độ hài lòng theo thời gian thực. Kết hợp "
         "với khảo sát bằng giọng nói Voice Survey để có dữ liệu chính xác.",
         "SmartVision và Sentiment AI: Đo lường sự hài lòng theo thời gian thực, độ chính xác trên 90%"),
        ("2.2.5. Dashboard Quản lý Thông minh — Tối ưu vận hành",
         "Dashboard dành cho lãnh đạo UBND và cán bộ một cửa với các chỉ số đánh giá hiệu quả KPI theo thời gian thực: "
         "số lượng hồ sơ theo trạng thái, thời gian xử lý trung bình, tỷ lệ hài lòng, cảnh báo tồn đọng. "
         "Tích hợp SmartUX để thu thập và trực quan hóa dữ liệu tương tác người dùng. Hỗ trợ xuất báo cáo "
         "định kỳ tự động.",
         "SmartUX và Analytics: Hơn 15 chỉ số KPI thời gian thực, báo cáo tự động hóa"),
        ("2.2.6. AI-Assisted Form Filling — Điền form thông minh",
         "Khi cán bộ mở hồ sơ mới, hệ thống gợi ý tự động điền các trường thông tin dựa trên dữ liệu "
         "từ OCR và lịch sử hồ sơ trước đó. Công nghệ ứng dụng mô hình ngôn ngữ lớn LLM giúp hiểu ngữ cảnh và tự động hoàn "
         "thiện các trường phức tạp. Giảm thời gian nhập liệu từ 15 phút xuống còn 2-3 phút.",
         "LLM và NLP: Giảm 80% thời gian nhập liệu cho cán bộ"),
    ]

    for title, desc, tech in features:
        make_heading(doc, title, 3)
        make_paragraph(doc, desc, size=12)
        make_paragraph(doc, f"Công nghệ: {tech}", size=11, bold=True,
                       first_line_indent=False, color=RGBColor(0x00, 0x56, 0xA6))

    make_heading(doc, "2.3. Kịch bản người dùng — Bác An (Công dân)", 2)
    make_paragraph(doc,
        "Bác Nguyễn Văn An (65 tuổi, phường Bến Thành, Quận 1) cần làm lại CCCD do bị mất. "
        "Bác đến UBND phường và được hướng dẫn đến Kiosk GovOne. Quy trình diễn ra như sau:"
    )
    steps = [
        ("Bước 1 — Chào hỏi: ", "Kiosk tự động chào bằng giọng nói “Xin chào, tôi có thể giúp gì cho bác?”. Bác An nói “Tôi muốn làm lại CCCD”."),
        ("Bước 2 — Hướng dẫn: ", "Hệ thống xác định thủ tục “Cấp lại CCCD” và hướng dẫn bằng giọng nói từng bước."),
        ("Bước 3 — Xác thực: ", "Bác đưa CCCD cũ (bản photo) vào khay scan, hệ thống OCR nhận dạng thông tin. Camera tích hợp chụp ảnh khuôn mặt, eKYC so sánh và xác thực. Toàn bộ mất 20 giây."),
        ("Bước 4 — Xác nhận: ", "Hệ thống đọc lại thông tin bằng giọng nói: “Bác An, họ tên Nguyễn Văn An, sinh năm 1961... Bác xác nhận đúng không ạ?”. Bác nói “Đúng rồi”."),
        ("Bước 5 — Hoàn tất: ", "Hệ thống thông báo: “Hồ sơ của bác đã được ghi nhận. Mã số hồ sơ: HS-2024-12345. Bác sẽ nhận được kết quả trong 5 ngày. Chúc bác sức khỏe!”. Hệ thống ghi nhận phản hồi cảm xúc: Hài lòng."),
    ]
    for bold_p, normal_p in steps:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "2.4. Kịch bản cán bộ — Chị Hương (Cán bộ một cửa)", 2)
    make_paragraph(doc,
        "Chị Hương, cán bộ bộ phận một cửa tại UBND phường, bắt đầu ca làm việc và mở Dashboard GovOne:"
    )
    steps = [
        ("Bước 1 — Dashboard tổng quan: ", "Chị thấy 12 hồ sơ chờ xử lý, 8 hồ sơ đang xử lý, 6 hồ sơ đã xử lý hôm nay. Một cảnh báo màu đỏ: “Hồ sơ HS-2024-12345 đã chờ quá 4 ngày”."),
        ("Bước 2 — Xử lý hồ sơ: ", "Chị mở hồ sơ xin cấp lại CCCD của bác An. Hệ thống đã tự động điền sẵn các trường thông tin từ OCR. Chị kiểm tra và xác nhận — chỉ mất 2 phút thay vì 15 phút như trước."),
        ("Bước 3 — Tra cứu nhanh: ", "Một người dân gọi điện hỏi về hồ sơ đã nộp 2 tháng trước. Chị gõ số CCCD, hệ thống trả về kết quả trong 5 giây (trước đây mất từ 30 đến 60 phút tra giấy tờ)."),
        ("Bước 4 — Báo cáo cuối ngày: ", "Dashboard tự động tổng hợp: 25 hồ sơ xử lý trong ngày, thời gian trung bình 7 phút/hồ sơ (giảm 70% so với mục tiêu), tỷ lệ hài lòng 94%."),
    ]
    for bold_p, normal_p in steps:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "2.5. Sơ đồ luồng xử lý nghiệp vụ", 2)
    make_paragraph(doc,
        "Hệ thống GovOne vận hành theo quy trình 5 bước khép kín: đầu tiên là tiếp nhận thông tin "
        "khi người dân tương tác qua Kiosk, Web hoặc Zalo; tiếp theo là bước xác thực danh tính bằng "
        "công nghệ eKYC và OCR; sau đó hệ thống tự động tạo hồ sơ điện tử và gửi thông báo đến cán bộ; "
        "kế đến, cán bộ tiến hành xử lý hồ sơ trên Dashboard với sự hỗ trợ của công cụ tự động điền form "
        "AI-Assisted; cuối cùng, hệ thống hoàn tất quy trình bằng cách gửi thông báo kết quả qua SMS hoặc Zalo "
        "đồng thời đo lường mức độ hài lòng của công dân sau khi hoàn thành giao dịch."
    )

    # Insert user flow diagram if available
    citizen_flow = os.path.join(ASSETS_DIR, "user-flow-citizen.png")
    officer_flow = os.path.join(ASSETS_DIR, "user-flow-officer.png")
    if os.path.exists(citizen_flow):
        add_image(doc, citizen_flow, width_cm=16,
                  caption="Hình 2.1: Sơ đồ luồng xử lý của công dân khi sử dụng GovOne — từ tương tác giọng nói đến hoàn tất hồ sơ")
    if os.path.exists(officer_flow):
        add_image(doc, officer_flow, width_cm=16,
                  caption="Hình 2.2: Quy trình xử lý hồ sơ của cán bộ trên GovOne Dashboard với AI-Assisted Form và KPI real-time")


def build_chapter3(doc):
    """Chương 3: Đổi mới & Khác biệt (20 điểm)"""
    make_heading(doc, "Chương 3: Đổi mới & Khác biệt", 1)

    make_heading(doc, "3.1. So sánh với các giải pháp hiện có", 2)
    make_paragraph(doc,
        "Thị trường giải pháp hành chính công tại Việt Nam hiện có nhiều sản phẩm, nhưng hầu hết "
        "đều giải quyết từng phần riêng lẻ. Bảng so sánh dưới đây cho thấy sự khác biệt rõ rệt của GovOne:"
    )

    headers = ["Tiêu chí", "Cổng DVC\nQuốc gia", "VNeID", "Zalo OA\nChatbot", "Google\nDocument AI", "GovOne"]
    data = [
        ["Tương tác giọng nói (Voice-First)", "❌", "❌", "❌", "❌", "✅ Toàn trình"],
        ["OCR giấy tờ VN", "❌", "✅ CCCD", "❌", "✅ Tổng quát", "✅ Hoàn chỉnh"],
        ["Xác thực eKYC", "❌", "✅", "❌", "❌", "✅ 3 lớp"],
        ["Chatbot LLM", "❌", "❌", "✅ Văn bản", "❌", "✅ Giọng nói + Văn bản"],
        ["Phân tích cảm xúc", "❌", "❌", "❌", "❌", "✅ Thời gian thực"],
        ["Kiosk phần cứng", "❌", "❌", "❌", "❌", "✅ Tích hợp"],
        ["Bảng phân tích (Dashboard)", "✅ Cơ bản", "❌", "❌", "❌", "✅ Nâng cao"],
        ["Tổng thể", "Cổng số\nthuần túy", "Định danh\nthuần túy", "Chat\nthuần túy", "OCR\nthuần túy", "Hệ sinh\nthái đầy đủ"],
    ]
    make_table(doc, headers, data)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 3.1: So sánh GovOne với các giải pháp hiện có. GovOne đạt 7/7 tiêu chí, khác biệt hoàn toàn so với các giải pháp đơn lẻ trên thị trường (khác biệt trên 60% về tính năng cốt lõi).")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    make_heading(doc, "3.2. Bốn điểm đổi mới cốt lõi", 2)

    innovations = [
        ("1. Voice-First toàn trình: ",
         "Không chỉ dừng lại ở việc chuyển đổi giọng nói thành văn bản như các giải pháp khác, "
         "GovOne thiết kế toàn bộ trải nghiệm người dùng xoay quanh giọng nói ở mọi bước — từ chào hỏi, "
         "hướng dẫn thủ tục, nhập liệu, xác nhận đến cảm ơn. Điều này giúp người già, người mù chữ hoặc "
         "người không rành công nghệ đều có thể sử dụng dịch vụ công một cách độc lập. Đây là lần đầu tiên "
         "một giải pháp hành chính công tại Việt Nam áp dụng triết lý Voice-First ở quy mô toàn trình."),
        ("2. Sentiment AI đo lường hài lòng thực tế: ",
         "Thay vì khảo sát giấy định kỳ (vốn chỉ đạt tỷ lệ phản hồi ~5-10%), GovOne sử dụng camera AI "
         "tích hợp tại quầy để phân tích biểu cảm khuôn mặt người dân ngay sau giao dịch. Kết hợp với "
         "khảo sát bằng giọng nói ngắn (30 giây) để thu thập phản hồi chi tiết. Dữ liệu thời gian thực này giúp UBND kịp "
         "thời điều chỉnh chất lượng dịch vụ — một cải tiến vượt trội so với mô hình khảo sát truyền thống."),
        ("3. Tích hợp Kiosk + Web + Zalo đa kênh đồng nhất: ",
         "GovOne không chỉ là một ứng dụng web hay một kiosk riêng lẻ. Hệ thống là một nền tảng đa kênh "
         "thống nhất: người dân có thể bắt đầu hồ sơ trên Zalo, tiếp tục tại Kiosk và theo dõi trên Web. "
         "Cán bộ xử lý trên Dashboard tập trung. Dữ liệu đồng bộ real-time qua API. Mô hình này cho phép "
         "UBND linh hoạt triển khai theo điều kiện thực tế mà vẫn đảm bảo trải nghiệm nhất quán."),
        ("4. Tích hợp LLM, OCR và giọng nói trong cùng quy trình: ",
         "Điểm đổi mới công nghệ mạnh nhất của GovOne là khả năng kết hợp ba công nghệ AI gồm mô hình ngôn ngữ lớn LLM, OCR và giọng nói "
         "trong cùng một quy trình xử lý hồ sơ. Ví dụ: người dân nói “Tôi muốn làm lại CCCD”, hệ thống "
         "dùng LLM để hiểu ý định, trích xuất thông tin từ CCCD qua OCR, xác thực qua eKYC, xác nhận lại "
         "bằng giọng nói qua TTS — tất cả diễn ra trong một luồng liên tục, không gián đoạn."),
    ]
    for bold_p, normal_p in innovations:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "3.3. Phân tích cạnh tranh", 2)
    make_paragraph(doc,
        "Thị trường giải pháp hành chính công thông minh tại Việt Nam có 5 nhóm đối thủ cạnh tranh chính:"
    )

    competitors = [
        ("FPT.AI — Smart Speech & OCR: ",
         "Đối thủ mạnh về AI tiếng Việt với API giọng nói và OCR chất lượng cao. Tuy nhiên, FPT cung cấp "
         "API đơn lẻ, không có giải pháp tích hợp Kiosk + Dashboard hoàn chỉnh. Sản phẩm của FPT hướng đến "
         "doanh nghiệp, chưa tối ưu cho khu vực hành chính công. Giá API cao (trung bình 500-2.000đ/lượt)."),
        ("VNPT eDoc — Hồ sơ điện tử: ",
         "Sản phẩm của VNPT tập trung vào lưu trữ hồ sơ điện tử, không có Voice-First, không có Sentiment AI. "
         "Khách hàng mục tiêu là cấp sở/ban ngành, không phải UBND phường/xã. Không có giải pháp Kiosk phần cứng."),
        ("Zalo OA — Chatbot DVC: ",
         "Giải pháp chatbot trên Zalo phổ biến nhưng chỉ hỗ trợ text, không có OCR, không eKYC, không Kiosk. "
         "Phụ thuộc vào nền tảng Zalo, không thể tùy chỉnh giao diện. Dữ liệu người dùng thuộc Zalo, khó đáp ứng "
         "yêu cầu lưu trữ trong nước theo Nghị định 13/2023."),
        ("Google Document AI: ",
         "OCR tiếng Việt tốt nhưng chưa tối ưu cho giấy tờ hành chính Việt Nam (CCCD, sổ hộ khẩu, giấy khai sinh). "
         "Không có Voice, không eKYC, không đáp ứng yêu cầu lưu trữ dữ liệu trong nước. Chi phí cao (>$10/1.000 trang)."),
        ("Gia công công nghệ thông tin (Outsourcing): ",
         "Các công ty gia công xây dựng giải pháp theo yêu cầu — chi phí cao (300-500 triệu/dự án), thời gian dài "
         "(3-6 tháng), khó bảo trì, không có sản phẩm chuẩn hóa. Mỗi dự án là một lần xây dựng lại từ đầu."),
    ]
    for bold_p, normal_p in competitors:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_paragraph(doc,
        "Kết luận: GovOne có ba lợi thế cạnh tranh khác biệt rõ rệt. Thứ nhất là sản phẩm đã được chuẩn hóa "
        "giúp triển khai nhanh chóng (chỉ trong 7 ngày đối với phiên bản MVP) thay vì phải gia công từng dự án. "
        "Thứ hai là tích hợp đồng bộ các tính năng giọng nói, OCR, phân tích cảm xúc và phần cứng Kiosk trong "
        "một nền tảng duy nhất thay vì các dịch vụ API rời rạc. Thứ ba là giá thành cạnh tranh, dao động từ "
        "8 đến 20 triệu đồng mỗi tháng, hoàn toàn phù hợp với ngân sách của các UBND phường, xã."
    )


def build_chapter4(doc):
    """Chương 4: Thiết kế tổng quan (Kiến trúc + Wireframe + UI)"""
    make_heading(doc, "Chương 4: Thiết kế tổng quan", 1)

    make_heading(doc, "4.1. Kiến trúc hệ thống", 2)
    make_paragraph(doc,
        "GovOne được thiết kế theo mô hình dịch vụ nhỏ Microservices với bốn tầng kiến trúc chính, "
        "mỗi tầng có trách nhiệm riêng biệt, đảm bảo tính linh hoạt, khả năng mở rộng và dễ bảo trì:"
    )

    arch_data = [
        ["Tầng", "Thành phần", "Công nghệ", "Chức năng"],
        ["1. Tầng hiển thị\n(Presentation)", "Kiosk (React), Web App\n(Next.js), Zalo Mini App",
         "Next.js 14, React 18,\nTailwind CSS, TypeScript",
         "Giao diện người dùng đa kênh\nưu tiên giọng nói và tương thích tốt"],
        ["2. Tầng kết nối\n(API Gateway & Services)", "FastAPI Gateway, Auth,\nAI Workers (Celery)",
         "FastAPI, Celery,\nRedis, JWT/OAuth2",
         "Định tuyến yêu cầu, xác thực,\nxử lý AI bất đồng bộ"],
        ["3. Tầng tích hợp AI\n(AI Engine Layer)", "SmartVoice, SmartReader,\nSmartbot, eKYC, SmartVision",
         "VNPT AI APIs,\nPython, ONNX",
         "Tích hợp và điều phối\ncác dịch vụ AI từ VNPT"],
        ["4. Tầng dữ liệu\n(Data Layer)", "PostgreSQL, S3/MinIO,\nElasticsearch",
         "PostgreSQL 15, MinIO,\nRedis Cache",
         "Lưu trữ hồ sơ, tài liệu số hóa,\nnhật ký hệ thống và chỉ số phân tích"],
    ]
    make_table(doc, arch_data[0], arch_data[1:], col_widths=[2.5, 4.0, 3.5, 4.5])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 4.1: Kiến trúc 4 tầng của GovOne")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    # Insert architecture diagram
    arch_path = os.path.join(ASSETS_DIR, "architecture-diagram.png")
    if os.path.exists(arch_path):
        # Extra blank paragraph before for centering
        doc.add_paragraph()
        add_image(doc, arch_path, width_cm=16,
                  caption="Hình 4.1: Sơ đồ kiến trúc tổng quan hệ thống GovOne")
        doc.add_paragraph()

    make_heading(doc, "4.2. Giao diện người dùng — Web UI Screenshots", 2)

    make_heading(doc, "4.2.1. Trang đăng nhập (Login Page)", 3)
    make_paragraph(doc,
        "Giao diện đăng nhập GovOne được thiết kế tối giản, tập trung vào trải nghiệm người dùng. "
        "Nền gradient màu xanh nhạt chuyển sang màu trắng tạo cảm giác thân thiện. Logo GovOne nổi bật ở "
        "trung tâm, kèm tiêu đề “Hệ thống Quản lý Hành chính Công Thông minh”. Khung đăng nhập được "
        "bố trí gọn gàng trong thẻ giao diện với hiệu ứng trượt lên nhẹ nhàng, yêu cầu email và mật khẩu. "
        "Hỗ trợ đăng ký tài khoản mới và chức năng quên mật khẩu."
    )

    login_screenshot = os.path.join(ASSETS_DIR, "web-login.png")
    if os.path.exists(login_screenshot):
        add_image(doc, login_screenshot, width_cm=16,
                  caption="Hình 4.2: Giao diện đăng nhập web-login của hệ thống GovOne")

    make_heading(doc, "4.2.2. Dashboard công dân (Citizen Dashboard)", 3)
    make_paragraph(doc,
        "Sau khi đăng nhập, công dân được chào đón với thông điệp cá nhân hóa “Xin chào, [Họ tên]”. "
        "Bảng điều khiển hiển thị ba thẻ thông tin thống kê chính bao gồm: thứ nhất là hồ sơ đang xử lý, "
        "thứ hai là lịch hẹn sắp tới với UBND, và thứ ba là các thông báo chưa đọc. "
        "Mỗi thẻ có biểu tượng và màu sắc riêng biệt (xanh dương, xanh lá, vàng). Khu vực “Tiện ích nhanh” cung cấp "
        "các công cụ liên kết đến Tra cứu hồ sơ, Nộp hồ sơ, Lịch hẹn, Thông tin cá nhân. "
        "Nút “Nộp hồ sơ mới” nổi bật ở góc phải với biểu tượng dấu cộng."
    )

    cit_screenshot = os.path.join(ASSETS_DIR, "web-citizen-dashboard.png")
    if os.path.exists(cit_screenshot):
        add_image(doc, cit_screenshot, width_cm=16,
                  caption="Hình 4.3: Dashboard công dân GovOne")

    make_heading(doc, "4.2.3. Dashboard cán bộ (Officer Dashboard)", 3)
    make_paragraph(doc,
        "Bảng quản lý của cán bộ có giao diện chuyên nghiệp hơn với bốn thẻ chỉ số KPI chính gồm: hồ sơ chờ xử lý "
        "đi kèm cảnh báo màu vàng, hồ sơ đang xử lý màu xanh dương, số hồ sơ đã xử lý trong ngày màu xanh lá, và "
        "lịch hẹn hôm nay hiển thị màu tím. Mục “Trạng thái hồ sơ” trực quan hóa phân bố hồ sơ qua biểu đồ "
        "thanh ngang theo trạng thái: chờ tiếp nhận, đang xử lý, chờ bổ sung, đã hoàn tất, và đã hủy với màu sắc trực quan. "
        "Mỗi trạng thái có ghi chú phần trăm tương ứng."
    )

    ofc_screenshot = os.path.join(ASSETS_DIR, "web-officer-dashboard.png")
    if os.path.exists(ofc_screenshot):
        add_image(doc, ofc_screenshot, width_cm=16,
                  caption="Hình 4.4: Dashboard cán bộ GovOne")

    make_heading(doc, "4.3. Wireframe Kiosk và các màn hình chính", 2)
    make_paragraph(doc,
        "Hệ thống Kiosk GovOne được thiết kế với màn hình cảm ứng 21.5 inch, tích hợp hệ thống microphone "
        "và camera thông minh. Giao diện Kiosk sử dụng font chữ lớn từ 18pt trở lên, độ tương phản cao, hỗ trợ đầy đủ "
        "cho người khiếm thị qua tính năng đọc màn hình và tối ưu cho người cao tuổi."
    )

    # Insert wireframe images
    kiosk_wf = os.path.join(ASSETS_DIR, "wireframe-kiosk.png")
    scan_wf = os.path.join(ASSETS_DIR, "wireframe-scan.png")
    dash_wf = os.path.join(ASSETS_DIR, "wireframe-dashboard.png")

    if os.path.exists(kiosk_wf):
        add_image(doc, kiosk_wf, width_cm=14,
                  caption="Hình 4.5: Wireframe màn hình chính Kiosk — Voice-First")
    if os.path.exists(scan_wf):
        add_image(doc, scan_wf, width_cm=14,
                  caption="Hình 4.6: Wireframe màn hình Scan OCR + eKYC")
    if os.path.exists(dash_wf):
        add_image(doc, dash_wf, width_cm=16,
                  caption="Hình 4.7: Wireframe Dashboard quản lý cho cán bộ")

    make_heading(doc, "4.4. Luồng xử lý chi tiết", 2)
    make_paragraph(doc,
        "Quy trình xử lý hồ sơ được thiết kế với 5 bước chính, mỗi bước có thể được thực hiện "
        "trên bất kỳ kênh nào (Kiosk/Web/Zalo):"
    )

    flow_data = [
        ["Bước", "Mô tả", "Kênh hỗ trợ", "Thời gian", "AI hỗ trợ"],
        ["1. Tiếp nhận", "Người dân nói hoặc lựa chọn thủ tục",
         "Kiosk, Web, Zalo", "< 1 phút", "SmartVoice STT\nSmartbot"],
        ["2. Xác thực", "Xác thực danh tính bằng eKYC và OCR",
         "Kiosk, Web", "< 30 giây", "VNPT eKYC\nSmartReader OCR"],
        ["3. Tạo hồ sơ", "Tự động tạo hồ sơ số và thông báo cán bộ",
         "Tự động", "< 10 giây", "LLM điền thông tin"],
        ["4. Xử lý", "Cán bộ đối chiếu thông tin và phê duyệt",
         "Bảng quản lý", "5-7 phút", "Bảng quản lý\nhỗ trợ AI"],
        ["5. Hoàn tất", "Gửi thông báo kết quả và đo lường sự hài lòng",
         "SMS, Zalo", "< 2 phút", "SmartVision\nSentiment AI"],
    ]
    make_table(doc, flow_data[0], flow_data[1:], col_widths=[1.5, 3.5, 2.5, 1.5, 2.5])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 4.2: Luồng xử lý nghiệp vụ GovOne — thời gian trung bình giảm 70% so với quy trình thủ công")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))


def build_chapter5(doc):
    """Chương 5: Tính khả thi (25 điểm)"""
    make_heading(doc, "Chương 5: Tính khả thi", 1)

    make_heading(doc, "5.1. Nguồn dữ liệu", 2)
    make_paragraph(doc,
        "GovOne sử dụng dữ liệu từ 3 nguồn hợp pháp, đảm bảo tuân thủ pháp luật Việt Nam:"
    )
    data_sources = [
        ("Nguồn 1 — Người dân cung cấp: ",
         "CCCD, sổ hộ khẩu, giấy khai sinh, giấy đăng ký xe và các giấy tờ hành chính khác do "
         "người dân tự nguyện cung cấp khi thực hiện giao dịch. Dữ liệu được xử lý và lưu trữ theo đúng "
         "quy định tại Nghị định 13/2023/NĐ-CP về Bảo vệ dữ liệu cá nhân."),
        ("Nguồn 2 — Dữ liệu nội bộ UBND: ",
         "Hồ sơ lưu trữ điện tử, số hóa từ kho giấy tờ hiện có của UBND theo Quyết định 178/2024/QĐ-TTg "
         "về lộ trình số hóa hồ sơ hành chính."),
        ("Nguồn 3 — API VNPT: ",
         "Các dịch vụ API do Ban Tổ chức HackAIthon 2026 cung cấp bao gồm giải pháp xác thực eKYC, nhận dạng giọng nói SmartVoice, "
         "trích xuất SmartReader, trợ lý ảo Smartbot và thị giác máy tính SmartVision, kết hợp các mô hình "
         "mã nguồn mở như Qwen và PhoBERT phục vụ cho việc "
         "xử lý ngôn ngữ tiếng Việt."),
    ]
    for bold_p, normal_p in data_sources:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "5.2. Nhân lực và kỹ thuật", 2)
    make_paragraph(doc,
        "Đội ngũ phát triển GovOne gồm 4 thành viên với kỹ năng bổ trợ đầy đủ. Công nghệ được lựa "
        "chọn dựa trên tiêu chí: quen thuộc với team, mã nguồn mở, cộng đồng lớn, dễ triển khai."
    )

    tech_data = [
        ["Hạng mục", "Công nghệ", "Kinh nghiệm team"],
        ["Frontend", "Next.js 14, React 18, TypeScript, Tailwind CSS",
         "✓ Có kinh nghiệm lập trình React và phát triển web app"],
        ["Backend", "Python (FastAPI), Celery, Redis",
         "✓ Kinh nghiệm lập trình Python và xây dựng dịch vụ API"],
        ["Database", "PostgreSQL 15, MinIO tương thích chuẩn S3",
         "✓ Kinh nghiệm quản trị và thiết kế cơ sở dữ liệu PostgreSQL"],
        ["AI Integration", "Bộ API của VNPT kết hợp các dịch vụ AI bằng Python",
         "✓ Đã tích hợp thử nghiệm dịch vụ eKYC, giọng nói và OCR"],
        ["DevOps", "Docker, Docker Compose, Nginx",
         "✓ Có khả năng đóng gói Docker và triển khai ứng dụng thực tế"],
    ]
    make_table(doc, tech_data[0], tech_data[1:])

    make_heading(doc, "5.3. MVP 7 ngày — Lộ trình phát triển", 2)
    make_paragraph(doc,
        "Dựa trên nguồn lực hiện tại, GovOne được xây dựng theo lộ trình MVP 7 ngày (tương đương 28 "
        "man-day), chia làm 3 phase:"
    )

    mvp_data = [
        ["Giai đoạn", "Ngày", "Công việc", "Kết quả"],
        ["1. Xây dựng nền tảng\n(Ngày 1-2)",
         "1-2",
         "Khởi tạo dự án Next.js và FastAPI.\nThiết kế cấu trúc cơ sở dữ liệu.\nXây dựng module xác thực bảo mật.",
         "Backend API cơ bản và\ngiao diện đăng nhập hoàn thiện"],
        ["2. Tích hợp AI cốt lõi\n(Ngày 3-5)",
         "3-5",
         "Tích hợp giải pháp định danh eKYC và trích xuất OCR.\nXây dựng kịch bản chatbot cho các thủ tục thông dụng.\nTích hợp tính năng chuyển đổi giọng nói.\nXây dựng giao diện Kiosk mô phỏng.",
         "Hoàn thành luồng trích xuất dữ liệu, định danh và hội thoại giọng nói"],
        ["3. Hoàn thiện hệ thống\n(Ngày 6-7)",
         "6-7",
         "Hoàn thiện bảng điều khiển dành cho cán bộ một cửa.\nTích hợp tính năng phân tích cảm xúc.\nKiểm thử luồng vận hành khép kín.\nTriển khai ứng dụng thử nghiệm.",
         "Hệ thống chạy ổn định và sẵn sàng vận hành thử nghiệm"],
    ]
    make_table(doc, mvp_data[0], mvp_data[1:], col_widths=[2.5, 1.2, 5.0, 3.5])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 5.1: Lộ trình MVP 7 ngày với các mốc kiểm tra cụ thể")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    make_heading(doc, "5.4. Chi phí hạ tầng và vận hành", 2)

    cost_data = [
        ["Hạng mục", "Chi phí/tháng", "Ghi chú"],
        ["VPS (8 CPU, 32GB RAM)", "1.500.000đ", "Thuê máy chủ đám mây chạy cơ sở dữ liệu và API"],
        ["Domain + SSL", "50.000đ", "Đăng ký tên miền và cấu hình chứng chỉ bảo mật"],
        ["MinIO Object Storage", "200.000đ", "Không gian lưu trữ tài liệu số hóa"],
        ["API VNPT (ưu đãi HackAIthon)", "0đ", "Được hỗ trợ miễn phí trong chương trình"],
        ["SMS OTP + Zalo OA", "500.000đ", "Chi phí duy trì kênh tương tác Zalo và gửi tin nhắn"],
        ["Tổng cộng", "2.250.000đ/tháng", "Chi phí vận hành kỹ thuật cơ bản ước tính"],
    ]
    make_table(doc, cost_data[0], cost_data[1:])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 5.2: Chi phí vận hành ước tính cho giai đoạn MVP và pilot (< 5 UBND)")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))

    make_heading(doc, "5.5. An toàn bảo mật và pháp lý", 2)
    make_paragraph(doc,
        "GovOne tuân thủ nghiêm ngặt các quy định pháp luật Việt Nam về bảo vệ dữ liệu và an toàn "
        "thông tin:"
    )
    security_points = [
        ("Tuân thủ Nghị định 13/2023/NĐ-CP về bảo vệ dữ liệu cá nhân: ",
         "Mọi dữ liệu cá nhân được thu thập với sự đồng ý rõ ràng của người dân. Dữ liệu được "
         "mã hóa AES-256 khi lưu trữ và TLS 1.3 khi truyền tải. Có cơ chế xóa dữ liệu khi người dân yêu cầu."),
        ("Tuân thủ Nghị định 85/2024/NĐ-CP về an toàn hệ thống thông tin: ",
         "Hệ thống được thiết kế theo tiêu chuẩn an toàn thông tin cấp độ 2, tích hợp đầy đủ cơ chế ghi nhật ký hoạt động, kiểm vết hệ thống và tự động sao lưu dữ liệu hàng ngày."),
        ("Luật Giao dịch điện tử 2023: ",
         "Hồ sơ điện tử có giá trị pháp lý tương đương hồ sơ giấy. Chữ ký điện tử được xác thực qua eKYC."),
        ("Quyết định 178/2024/QĐ-TTg: ",
         "Lộ trình số hóa hồ sơ hành chính phù hợp với định hướng của Chính phủ."),
        ("Bảo mật giao tiếp API: ",
         "Sử dụng mã xác thực JWT có thời hạn ngắn, áp dụng giới hạn tần suất truy cập để phòng ngừa tấn công và lưu trữ mã khóa bảo mật trong biến môi trường hệ thống."),
    ]
    for bold_p, normal_p in security_points:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_heading(doc, "5.6. Lộ trình triển khai 12 tháng (GTM Roadmap)", 2)
    make_paragraph(doc,
        "Sau cuộc thi, GovOne sẽ triển khai theo lộ trình 4 giai đoạn, tập trung vào thị trường "
        "UBND phường/xã tại TP.HCM và Hà Nội trước:"
    )

    roadmap_data = [
        ["Giai đoạn", "Thời gian", "Mục tiêu", "Số đơn vị", "Doanh thu (ước)"],
        ["1. Thử nghiệm thực tế\n(Giai đoạn Pilot)", "Tháng 1-2\nsau khi có MVP",
         "Triển khai áp dụng thử nghiệm tại 5 UBND phường thuộc Quận 1, TP.HCM để thu thập ý kiến phản hồi và tối ưu hóa sản phẩm.",
         "5 đơn vị", "Miễn phí trải nghiệm"],
        ["2. Khách hàng tiên phong\n(Early Adopter)", "Tháng 3-6",
         "Mở rộng áp dụng đến 20 UBND tại TP.HCM và Hà Nội, ký kết hợp đồng cung cấp gói dịch vụ cơ bản và xây dựng mô hình điểm.",
         "20 đơn vị", "160.000.000đ\n(8 triệu/tháng)"],
        ["3. Giai đoạn tăng trưởng\n(Growth Phase)", "Tháng 7-10",
         "Triển khai đến 50 đơn vị thuộc 5 tỉnh thành lớn, ra mắt gói dịch vụ cao cấp dành cho cấp quận huyện và phát triển mạng lưới đối tác phân phối.",
         "50 đơn vị", "500.000.000đ\n(trung bình 10 triệu/tháng)"],
        ["4. Mở rộng quy mô\n(Scale-up)", "Tháng 11-12",
         "Phát triển đến hơn 100 UBND trên toàn quốc thông qua hợp tác với Sở TT&TT các địa phương.",
         "Trên 100 đơn vị", "1 đến 2 tỷ đồng/tháng"],
    ]
    make_table(doc, roadmap_data[0], roadmap_data[1:], col_widths=[2.0, 2.0, 4.0, 1.5, 2.0])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Bảng 5.3: Lộ trình Go-to-Market 12 tháng của GovOne")
    set_font(r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))


def build_chapter6(doc):
    """Chương 6: Tác động dự kiến (20 điểm)"""
    make_heading(doc, "Chương 6: Tác động dự kiến", 1)

    make_heading(doc, "6.1. Phân tích thị trường TAM-SAM-SOM", 2)
    make_paragraph(doc,
        "Thị trường giải pháp hành chính công thông minh tại Việt Nam được phân tích theo mô hình "
        "TAM-SAM-SOM, dựa trên dữ liệu từ Bộ TT&TT và Tổng cục Thống kê năm 2025:"
    )

    tam_data = [
        ["Chỉ số", "Giá trị (VNĐ)", "Cơ sở tính toán"],
        ["Tổng quy mô thị trường (TAM)",
         "51.000 tỷ",
         "Toàn bộ chi tiêu CNTT của 63 tỉnh/thành + 705 quận/huyện + 10.599 xã/phường. "
         "Trung bình 500 triệu/đơn vị/năm cho chuyển đổi số (theo Quyết định 178/2024/QĐ-TTg)."],
        ["Thị trường có thể phục vụ (SAM)",
         "800 tỷ",
         "Phân khúc UBND phường/xã có nhu cầu cấp bách về voice + OCR (5.000 đơn vị × 160 triệu đồng). "
         "Tập trung TP.HCM, Hà Nội, Đà Nẵng, Cần Thơ, Hải Phòng."],
        ["Thị trường mục tiêu đạt được (SOM)",
         "40 tỷ\n(năm 1-3)",
         "Mục tiêu năm 1: 200 UBND × 200 triệu đồng (phần cứng + phần mềm + API). "
         "Năm 2: 500 UBND. Năm 3: 1.000 UBND. Tăng trưởng 150%/năm."],
    ]
    make_table(doc, tam_data[0], tam_data[1:], col_widths=[3.5, 2.0, 10.0])

    make_heading(doc, "6.2. Lợi ích xã hội", 2)
    make_paragraph(doc,
        "GovOne mang lại 7 tác động tích cực đó lường được, phù hợp với mục tiêu cải cách hành chính "
        "của Chính phủ:"
    )

    benefits_data = [
        ["Chỉ số", "Hiện tại", "Mục tiêu GovOne", "Phương pháp đo"],
        ["Thời gian xử lý thủ tục", "20-30 phút", "5-7 phút (giảm 70%)", "Thời gian thực tế"],
        ["Độ phủ người dùng\n(bao gồm người cao tuổi và người khuyết tật)", "~30% dân số", ">95%", "Khảo sát sau giao dịch"],
        ["Tỷ lệ hài lòng", "~65%", ">90%", "Phân tích cảm xúc và khảo sát"],
        ["Thời gian tra cứu hồ sơ cũ", "30-60 phút", "<5 phút (giảm 90%)", "Nhật ký hệ thống OCR"],
        ["Tỷ lệ sai sót hồ sơ", "~15%", "<2%", "Kiểm tra định kỳ"],
        ["Chi phí vận hành/UBND/năm", "~150-200 tr", "~50 tr (giảm 75%)", "Báo cáo tài chính"],
        ["Lượng giấy lưu trữ vật lý", "50K-200K bộ", "Giảm 80%", "Thống kê scan"],
    ]
    make_table(doc, benefits_data[0], benefits_data[1:], col_widths=[3.0, 2.0, 2.5, 3.0])

    make_heading(doc, "6.3. Mô hình doanh thu B2G", 2)
    make_paragraph(doc,
        "GovOne áp dụng mô hình kinh doanh B2G (giữa doanh nghiệp và chính phủ) với ba gói dịch vụ linh hoạt, "
        "từ cơ bản đến cao cấp, phù hợp với quy mô và nhu cầu của từng UBND:"
    )

    pricing_data = [
        ["Gói dịch vụ", "Giá (VNĐ/tháng)", "Bao gồm", "Đối tượng"],
        ["Basic", "8.000.000",
         "• Kiosk voice-first\n• OCR 500 lượt/tháng\n• Smartbot 10 thủ tục\n• Dashboard cơ bản",
         "UBND phường/xã\nquy mô nhỏ"],
        ["Pro", "20.000.000",
         "• Kiosk + Web App\n• OCR 2.000 lượt/tháng\n• Smartbot 50 thủ tục\n• eKYC\n• Dashboard nâng cao",
         "UBND quận/huyện\nquy mô trung bình"],
        ["Enterprise", "Liên hệ",
         "• Tất cả tính năng\n• OCR không giới hạn\n• Smartbot toàn bộ thủ tục\n• Tích hợp CSDL sở/ban ngành\n• SLA 99.9%",
         "Sở TT&TT,\nUBND tỉnh"],
    ]
    make_table(doc, pricing_data[0], pricing_data[1:], col_widths=[2.0, 2.5, 5.0, 2.5])

    make_heading(doc, "6.4. Phân tích ưu thế cạnh tranh", 2)
    make_paragraph(doc,
        "So với các đối thủ trên thị trường, GovOne có 3 ưu thế cạnh tranh chiến lược:"
    )
    advantages = [
        ("1. Chi phí thấp hơn 10 lần so với gia công: ",
         "Trong khi các dự án gia công CNTT có chi phí 300-500 triệu cho 3-6 tháng phát triển, GovOne "
         "cung cấp giải pháp chuẩn hóa với giá chỉ 8-20 triệu/tháng, bao gồm cả phần cứng Kiosk, phần mềm "
         "và hỗ trợ kỹ thuật. UBND không cần đầu tư ban đầu lớn."),
        ("2. Triển khai nhanh (7 ngày) so với 3-6 tháng: ",
         "GovOne có thể triển khai MVP tại một UBND mới trong 7 ngày, trong khi các giải pháp gia công "
         "hoặc nội bộ mất 3-6 tháng. Lý do: GovOne là sản phẩm chuẩn hóa, đã được xây dựng và kiểm thử, "
         "không cần xây dựng lại từ đầu."),
        ("3. Tích hợp đa nền tảng AI duy nhất trên thị trường: ",
         "Không có đối thủ nào tại Việt Nam cung cấp giải pháp tích hợp Voice + OCR + eKYC + Sentiment + "
         "Kiosk trong cùng một nền tảng. Các đối thủ lớn (FPT, VNPT) chỉ cung cấp API riêng lẻ. Các startup "
         "nhỏ chỉ làm được 1-2 tính năng. GovOne là sản phẩm đầu tiên tích hợp tất cả trong một hệ sinh thái."),
    ]
    for bold_p, normal_p in advantages:
        add_bullet(doc, bold_p, normal_p, indent_cm=1.0, size=12)

    make_paragraph(doc,
        "Với mô hình kinh doanh B2G linh hoạt, thị trường mục tiêu rõ ràng (UBND phường/xã) và 3 ưu thế "
        "cạnh tranh chiến lược, GovOne có khả năng đạt được SOM 40 tỷ trong 3 năm đầu và mở rộng lên "
        "200 tỷ trong 5 năm."
    )


def build_chapter7(doc):
    """Chương 7: Kết luận & Cam kết"""
    make_heading(doc, "Chương 7: Kết luận", 1)

    make_paragraph(doc,
        "GovOne là giải pháp hành chính công thông minh toàn diện, được thiết kế để giải quyết triệt để "
        "6 pain-point cốt lõi của hệ thống hành chính công Việt Nam thông qua 6 tính năng AI mạnh mẽ: "
        "Kiosk Voice-First, OCR & eKYC, Smartbot Đa kênh, Sentiment AI, Dashboard Quản lý Thông minh, "
        "và AI-Assisted Form Filling."
    )

    make_paragraph(doc,
        "Với kiến trúc 4 tầng microservices, tích hợp 5 nền tảng AI từ VNPT (eKYC, SmartVoice, Smartbot, "
        "SmartReader, SmartVision), và lộ trình phát triển MVP 7 ngày, GovOne hoàn toàn khả thi về mặt "
        "kỹ thuật và tài chính. Mô hình B2G với 3 gói dịch vụ linh hoạt đảm bảo phù hợp với mọi quy mô UBND."
    )

    make_paragraph(doc,
        "GovOne không chỉ là một sản phẩm công nghệ — đó là một giải pháp có tác động xã hội sâu sắc. "
        "Với mục tiêu rút ngắn phần lớn thời gian xử lý hồ sơ, nâng cao chất lượng phục vụ để tăng tỷ lệ "
        "hài lòng của phần đông người dân, cắt giảm phần lớn lượng giấy tờ lưu trữ vật lý và mở rộng độ phủ "
        "dịch vụ công đến đại bộ phận dân cư — GovOne hướng tới một nền hành chính hiện đại, minh bạch và "
        "nhân văn hơn."
    )

    make_paragraph(doc,
        "Đội ngũ GovOne cam kết tiếp tục phát triển sản phẩm sau cuộc thi, với lộ trình 12 tháng rõ ràng "
        "và mục tiêu phục vụ 1.000+ UBND trong 3 năm. Chúng tôi tin rằng GovOne sẽ góp phần hiện thực hóa "
        "mục tiêu chuyển đổi số quốc gia đến năm 2030 của Chính phủ."
    )

    make_paragraph(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Trân trọng,\nĐội thi GovOne\nNguyễn Ngọc Bình An (Trưởng nhóm)\nHoàng Thị Linh Hương\nNguyễn Đoàn Nhật Minh\nTrần Hoàng Nguyên\nPhạm Lê Việt Đức")
    set_font(r, size=12, italic=True)

    # Add logo at the end
    logo_path = os.path.join(ASSETS_DIR, "logo-govone.png")
    if os.path.exists(logo_path):
        make_paragraph(doc, "", size=6, first_line_indent=False)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(logo_path, width=Cm(3))


def add_header_footer(doc):
    """Add headers and footers with page numbers."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run("GovOne — Đội thi Nguyễn Ngọc Bình An et al. — HackAIthon 2026 (Bảng B — Challenger)")
        set_font(r, size=8, color=RGBColor(0x88, 0x88, 0x88))

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = fp.add_run()
        fldChar1 = parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w"))
        run._r.append(fldChar1)
        run2 = fp.add_run()
        instrText = parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w"))
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar2 = parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w"))
        run3._r.append(fldChar2)
        r = fp.add_run(" / ")
        set_font(r, size=8)
        run4 = fp.add_run()
        fldChar3 = parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w"))
        run4._r.append(fldChar3)
        run5 = fp.add_run()
        instrText2 = parse_xml('<w:instrText %s xml:space="preserve"> NUMPAGES </w:instrText>' % nsdecls("w"))
        run5._r.append(instrText2)
        run6 = fp.add_run()
        fldChar4 = parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w"))
        run6._r.append(fldChar4)


def generate_web_screenshots():
    """Generate web UI mockup screenshots using PIL."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        import textwrap
    except ImportError:
        print("⚠️  PIL not installed, can't generate screenshots. Using existing wireframes only.")
        return

    os.makedirs(ASSETS_DIR, exist_ok=True)

    def draw_rounded_rect(draw, xy, radius, fill=None, outline=None, width=1):
        x1, y1, x2, y2 = xy
        draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)

    def has_vietnamese_font():
        """Check if we have a Unicode font available."""
        import subprocess
        result = subprocess.run(["fc-list", ":lang=vi"], capture_output=True, text=True)
        return len(result.stdout.strip()) > 0

    # Find a font supporting Vietnamese
    font_paths = [
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    font_path = None
    for fp in font_paths:
        if os.path.exists(fp):
            font_path = fp
            break

    if not font_path:
        result = subprocess.run(["fc-match", "sans"], capture_output=True, text=True)
        print(f"Using system font match: {result.stdout.strip()}")

    def get_font(size=14, bold=False):
        try:
            if bold and font_path:
                bold_path = font_path.replace("-Regular", "-Bold")
                if os.path.exists(bold_path):
                    return ImageFont.truetype(bold_path, size)
            if font_path:
                return ImageFont.truetype(font_path, size)
            return ImageFont.load_default()
        except Exception:
            return ImageFont.load_default()

    # ── 1. Login Page Screenshot ──
    w, h = 1280, 800
    img = Image.new("RGB", (w, h), (248, 250, 252))
    draw = ImageDraw.Draw(img)

    # Gradient background
    for y in range(h):
        r = int(248 - (248 - 255) * y / h)
        g = int(250 - (250 - 255) * y / h)
        b = int(252 - (252 - 255) * y / h)
        draw.line([(0, y), (w, y)], fill=(r, g, b))

    # Card background
    card_w, card_h = 440, 520
    card_x = (w - card_w) // 2
    card_y = (h - card_h) // 2
    draw_rounded_rect(draw, (card_x, card_y, card_x + card_w, card_y + card_h),
                      radius=16, fill=(255, 255, 255), outline=(220, 225, 235), width=1)

    # Logo circle
    cx, cy = w // 2, card_y + 60
    draw.ellipse([cx - 30, cy - 30, cx + 30, cy + 30], fill=(0, 86, 166))
    draw.text((cx, cy), "G", fill=(255, 255, 255), font=get_font(28, bold=True), anchor="mm")

    # Title
    draw.text((cx, card_y + 100), "GovOne", fill=(30, 40, 60), font=get_font(24, bold=True), anchor="mm")
    draw.text((cx, card_y + 125), "Hệ thống Quản lý Hành chính Công Thông minh",
              fill=(100, 110, 130), font=get_font(11), anchor="mm")

    # Form fields
    draw_rounded_rect(draw, (card_x + 40, card_y + 165, card_x + card_w - 40, card_y + 215),
                      radius=8, fill=(245, 247, 250), outline=(215, 222, 232))
    draw.text((card_x + 55, card_y + 178), "Email", fill=(100, 110, 130), font=get_font(11))
    draw.text((card_x + 380, card_y + 178), "✉", fill=(180, 190, 205), font=get_font(12))

    draw_rounded_rect(draw, (card_x + 40, card_y + 235, card_x + card_w - 40, card_y + 285),
                      radius=8, fill=(245, 247, 250), outline=(215, 222, 232))
    draw.text((card_x + 55, card_y + 248), "Mật khẩu", fill=(100, 110, 130), font=get_font(11))
    draw.text((card_x + 380, card_y + 248), "🔒", fill=(180, 190, 205), font=get_font(12))

    # Login button
    draw_rounded_rect(draw, (card_x + 40, card_y + 315, card_x + card_w - 40, card_y + 370),
                      radius=8, fill=(0, 86, 166))
    draw.text((cx, card_y + 342), "Đăng nhập", fill=(255, 255, 255), font=get_font(15, bold=True), anchor="mm")

    # Links
    draw.text((cx - 60, card_y + 395), "Chưa có tài khoản?", fill=(130, 140, 160), font=get_font(11))
    draw.text((cx + 50, card_y + 395), "Đăng ký", fill=(0, 86, 166), font=get_font(11, bold=True))
    draw.text((cx, card_y + 420), "Quên mật khẩu?", fill=(0, 86, 166), font=get_font(11))

    img.save(os.path.join(ASSETS_DIR, "web-login.png"), "PNG")
    print("✅ Generated: web-login.png")

    # ── 2. Citizen Dashboard ──
    w2, h2 = 1280, 800
    img2 = Image.new("RGB", (w2, h2), (248, 250, 252))
    draw2 = ImageDraw.Draw(img2)

    # Top header bar
    draw_rounded_rect(draw2, (0, 0, w2, 60), radius=0, fill=(255, 255, 255))
    draw2.rectangle([0, 58, w2, 60], fill=(230, 235, 242))

    # Logo + title
    draw2.ellipse([20, 15, 45, 40], fill=(0, 86, 166))
    draw2.text((32, 27), "G", fill=(255, 255, 255), font=get_font(16, bold=True), anchor="mm")
    draw2.text((55, 22), "GovOne", fill=(30, 40, 60), font=get_font(16, bold=True))

    # User menu
    draw2.text((1080, 25), "Nguyễn Văn An ▼", fill=(60, 70, 90), font=get_font(13))
    draw2.ellipse([1240, 18, 1260, 38], fill=(0, 86, 166))

    # Welcome area
    draw2.text((30, 80), "Xin chào, Nguyễn Văn An", fill=(30, 40, 60), font=get_font(22, bold=True))
    draw2.text((30, 112), "Chào mừng bạn đến với GovOne — Hệ thống hành chính công thông minh",
               fill=(100, 110, 130), font=get_font(13))

    # New Application button
    btn_x, btn_y = 1050, 85
    draw_rounded_rect(draw2, (btn_x, btn_y, btn_x + 190, btn_y + 42), radius=8, fill=(0, 86, 166))
    draw2.text((btn_x + 12, btn_y + 12), "+", fill=(255, 255, 255), font=get_font(20, bold=True))
    draw2.text((btn_x + 35, btn_y + 13), "Nộp hồ sơ mới", fill=(255, 255, 255), font=get_font(13, bold=True))

    # Stats cards
    stats = [
        ("Hồ sơ đang xử lý", "3", (0, 86, 166), (235, 245, 255)),
        ("Lịch hẹn sắp tới", "1", (34, 160, 80), (235, 250, 240)),
        ("Thông báo chưa đọc", "2", (210, 140, 20), (255, 248, 235)),
    ]
    card_w2 = 380
    card_h2 = 160
    gap = 25
    total_w = 3 * card_w2 + 2 * gap
    start_x = (w2 - total_w) // 2

    for i, (label, val, clr, bg) in enumerate(stats):
        x = start_x + i * (card_w2 + gap)
        draw_rounded_rect(draw2, (x, 170, x + card_w2, 170 + card_h2), radius=12, fill=(255, 255, 255))
        # Icon circle
        draw2.ellipse([x + 20, 190, x + 56, 226], fill=clr)
        draw2.text((x + 38, 200), str(i + 1), fill=(255, 255, 255), font=get_font(14, bold=True), anchor="mm")
        draw2.text((x + 70, 196), label, fill=(100, 110, 130), font=get_font(12))
        draw2.text((x + 25, 235), val, fill=(30, 40, 60), font=get_font(28, bold=True))

    # Quick actions section
    draw2.text((30, 370), "Tiện ích nhanh", fill=(30, 40, 60), font=get_font(16, bold=True))

    actions = [
        ("🔍", "Tra cứu hồ sơ", "Xem trạng thái hồ sơ"),
        ("📝", "Nộp hồ sơ", "Gửi hồ sơ trực tuyến"),
        ("📅", "Lịch hẹn", "Đặt & quản lý lịch"),
        ("👤", "Thông tin", "Cập nhật hồ sơ"),
    ]
    action_w = 280
    action_h = 100
    total_aw = 4 * action_w + 3 * gap
    ax = (w2 - total_aw) // 2

    for i, (icon, title, desc) in enumerate(actions):
        x = ax + i * (action_w + gap)
        draw_rounded_rect(draw2, (x, 400, x + action_w, 400 + action_h), radius=12, fill=(255, 255, 255))
        draw2.text((x + 20, 420), icon, fill=(0, 86, 166), font=get_font(20))
        draw2.text((x + 60, 420), title, fill=(30, 40, 60), font=get_font(13, bold=True))
        draw2.text((x + 20, 450), desc, fill=(140, 150, 170), font=get_font(11))

    img2.save(os.path.join(ASSETS_DIR, "web-citizen-dashboard.png"), "PNG")
    print("✅ Generated: web-citizen-dashboard.png")

    # ── 3. Officer Dashboard ──
    img3 = Image.new("RGB", (w2, h2), (248, 250, 252))
    draw3 = ImageDraw.Draw(img3)

    # Header
    draw_rounded_rect(draw3, (0, 0, w2, 60), radius=0, fill=(30, 40, 60))
    draw3.text((20, 22), "GovOne", fill=(255, 255, 255), font=get_font(16, bold=True))
    draw3.text((120, 25), "Bảng điều khiển — Cán bộ", fill=(180, 200, 220), font=get_font(12))

    # User info
    draw3.text((1020, 25), "Chị Hương ▼", fill=(200, 220, 240), font=get_font(12))
    draw3.ellipse([1190, 18, 1210, 38], fill=(0, 150, 80))

    # Page title
    draw3.text((30, 80), "Xin chào, Chị Hương", fill=(30, 40, 60), font=get_font(22, bold=True))
    draw3.text((30, 112), "Bảng điều khiển quản lý hồ sơ hành chính",
               fill=(100, 110, 130), font=get_font(13))

    # KPI cards
    kpis = [
        ("Hồ sơ chờ xử lý", "12", "+3 hôm nay", (210, 140, 20), (255, 248, 235)),
        ("Đang xử lý", "8", "5 hồ sơ của bạn", (0, 86, 166), (235, 245, 255)),
        ("Đã xử lý hôm nay", "6", "Đạt 75% chỉ tiêu", (34, 160, 80), (235, 250, 240)),
        ("Lịch hẹn hôm nay", "4", "2 lịch chưa xác nhận", (120, 60, 180), (245, 240, 255)),
    ]
    kpi_w = 280
    kpi_h = 150
    total_kw = 4 * kpi_w + 3 * gap
    kx = (w2 - total_kw) // 2

    for i, (label, val, change, clr, bg) in enumerate(kpis):
        x = kx + i * (kpi_w + gap)
        draw_rounded_rect(draw3, (x, 155, x + kpi_w, 155 + kpi_h), radius=12, fill=(255, 255, 255))
        draw3.ellipse([x + 20, 175, x + 56, 211], fill=clr)
        draw3.text((x + 38, 185), str(i + 1), fill=(255, 255, 255), font=get_font(14, bold=True), anchor="mm")
        draw3.text((x + 70, 181), label, fill=(100, 110, 130), font=get_font(11))
        draw3.text((x + 25, 220), val, fill=(30, 40, 60), font=get_font(28, bold=True))
        draw3.text((x + 25, 260), change, fill=(140, 150, 170), font=get_font(10))

    # Status distribution section
    draw3.text((30, 340), "Trạng thái hồ sơ", fill=(30, 40, 60), font=get_font(16, bold=True))

    statuses = [
        ("Chờ tiếp nhận", 5, (160, 170, 180)),
        ("Đang xử lý", 8, (0, 86, 166)),
        ("Chờ bổ sung", 3, (210, 140, 20)),
        ("Đã hoàn tất", 15, (34, 160, 80)),
        ("Đã hủy", 1, (200, 60, 50)),
    ]
    total = sum(s[1] for s in statuses)

    bar_x = 60
    bar_y = 370
    max_bar_w = 700
    bar_h = 28
    bar_gap = 8

    for i, (lbl, val, clr) in enumerate(statuses):
        y = bar_y + i * (bar_h + bar_gap)
        pct = val / total
        bw = int(max_bar_w * pct)
        draw3.rectangle([bar_x, y, bar_x + bw, y + bar_h], fill=clr)
        draw3.text((bar_x + 10, y + 3), lbl, fill=(255, 255, 255), font=get_font(11, bold=True))
        draw3.text((bar_x + bw + 10, y + 3), f"{val} ({pct * 100:.0f}%)",
                   fill=(60, 70, 90), font=get_font(11))

    # Recent activity
    draw3.text((30, 540), "Hoạt động gần đây", fill=(30, 40, 60), font=get_font(16, bold=True))

    activities = [
        ("HS-2024-12345 — Cấp lại CCCD", "Đã xử lý", "2 phút trước", (34, 160, 80)),
        ("HS-2024-12346 — Đăng ký khai sinh", "Đang xử lý", "15 phút trước", (0, 86, 166)),
        ("HS-2024-12347 — Xác nhận cư trú", "Chờ bổ sung", "1 giờ trước", (210, 140, 20)),
        ("HS-2024-12348 — Đăng ký kết hôn", "Chờ xử lý", "2 giờ trước", (160, 170, 180)),
        ("HS-2024-12349 — Giấy phép xây dựng", "Cảnh báo tồn đọng!", "4 ngày trước", (200, 60, 50)),
    ]

    for i, (doc_name, status, time, st_clr) in enumerate(activities):
        y = 570 + i * 35
        draw3.text((45, y), doc_name, fill=(30, 40, 60), font=get_font(12))
        draw3.text((500, y), status, fill=st_clr, font=get_font(11, bold=True))
        draw3.text((700, y), time, fill=(160, 170, 185), font=get_font(11))

    # Bottom status bar
    draw_rounded_rect(draw3, (0, h2 - 35, w2, h2), radius=0, fill=(30, 40, 60))
    draw3.text((20, h2 - 25), "GovOne v1.0 — Bảng B: Challenger — HackAIthon 2026",
               fill=(150, 170, 190), font=get_font(9))
    draw3.text((1000, h2 - 25), "Cập nhật gần nhất: 2 giây trước",
               fill=(150, 170, 190), font=get_font(9))

    img3.save(os.path.join(ASSETS_DIR, "web-officer-dashboard.png"), "PNG")
    print("✅ Generated: web-officer-dashboard.png")


def main():
    # ── Ensure output directory ──
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(ASSETS_DIR, exist_ok=True)

    # ── Generate web UI screenshots ──
    print("🖼️  Generating web UI screenshots...")
    generate_web_screenshots()
    print("🔄 Generating user flow diagrams...")
    from generate_user_flows import generate_citizen_flow, generate_officer_flow
    generate_citizen_flow()
    generate_officer_flow()

    # ── Create document ──
    print("📄 Creating document...")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Page setup — A4 size (21cm x 29.7cm)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # ── Build sections ──
    build_title_page(doc)
    add_page_break(doc)

    # Add Table of Contents (use Heading styles for automatic bookmarks)
    make_heading(doc, "Mục lục", 1)
    make_paragraph(doc, "(Nhấn Ctrl + Click để đến trang tương ứng trong file DOCX. Trong file PDF, "
                        "sử dụng thanh bookmarks (Ctrl+F) để điều hướng.)", size=11, first_line_indent=False,
                   color=RGBColor(0x88, 0x88, 0x88))
    toc_items = [
        "Chương 1: Đặt vấn đề",
        "Chương 2: Giải pháp GovOne",
        "Chương 3: Đổi mới & Khác biệt",
        "Chương 4: Thiết kế tổng quan",
        "Chương 5: Tính khả thi",
        "Chương 6: Tác động dự kiến",
        "Chương 7: Kết luận",
    ]
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.0)
        r = p.add_run(f"  {i}. {item}")
        set_font(r, size=13)

    add_page_break(doc)
    build_chapter1(doc)
    add_page_break(doc)
    build_chapter2(doc)
    add_page_break(doc)
    build_chapter3(doc)
    add_page_break(doc)
    build_chapter4(doc)
    add_page_break(doc)
    build_chapter5(doc)
    add_page_break(doc)
    build_chapter6(doc)
    build_chapter7(doc)

    # ── Add header & footer ──
    print("📑 Adding header/footer...")
    add_header_footer(doc)

    # ── Save DOCX ──
    print(f"💾 Saving DOCX to {DOCX_PATH}...")
    doc.save(DOCX_PATH)
    print(f"✅ DOCX saved: {os.path.getsize(DOCX_PATH) / 1024:.1f} KB")

    # ── Convert to PDF ──
    print("🔄 Converting to PDF via LibreOffice...")
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", OUTPUT_DIR, DOCX_PATH],
        capture_output=True, text=True, timeout=120
    )
    if result.returncode == 0:
        if os.path.exists(PDF_PATH):
            print(f"✅ PDF saved: {os.path.getsize(PDF_PATH) / 1024:.1f} KB")
        else:
            print(f"⚠️  PDF may have different name. Check: {OUTPUT_DIR}")
    else:
        print(f"❌ LibreOffice error: {result.stderr}")
        print("ℹ️  You can manually convert the DOCX to PDF using Word/Google Docs.")

    # ── Verify ──
    print("\n📋 Verification:")
    verify_script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "verify_content.py")
    subprocess.run(["python3", verify_script, DOCX_PATH])

    print(f"\n🎉 Done! Files in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
