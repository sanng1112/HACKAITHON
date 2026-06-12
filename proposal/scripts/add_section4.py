#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_bullet(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('\u2022 '); set_font(r, size=13)
    r = p.add_run(bold_part); set_font(r, size=13, bold=True)
    r = p.add_run(normal_part); set_font(r, size=13)

def make_table(doc, headers, data, col_aligns=None):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            if col_aligns and ci < len(col_aligns): p.alignment = col_aligns[ci]
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct); set_font(r, size=10)
            if ri % 2 == 1: c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))
    return table

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    doc_path = os.path.join(project_dir, 'proposal.docx')
    doc = Document(doc_path)
    add_page_break(doc)
    add_heading(doc, '4. T\u00cdNH KH\u1ea2 THI')
    add_heading(doc, '4.1 Ngu\u1ed3n d\u1eef li\u1ec7u', 2)
    add_para(doc, 'GovOne t\u1eadn d\u1ee5ng 4 ngu\u1ed3n d\u1eef li\u1ec7u ch\u00ednh \u0111\u1ec3 \u0111\u1ea3m b\u1ea3o \u0111\u1ed9 ch\u00ednh x\u00e1c v\u00e0 tin c\u1eady trong x\u1eed l\u00fd:')
    make_table(doc,
        ['Ngu\u1ed3n d\u1eef li\u1ec7u', 'M\u00f4 t\u1ea3', 'Ph\u1ea1m vi', 'Ph\u01b0\u01a1ng th\u1ee9c k\u1ebft n\u1ed1i'],
        [['API VNPT (SmartVoice, SmartReader, eKYC, SmartVision)','D\u1ecbch v\u1ee5 AI \u0111\u00e1m m\u00e2y do VNPT cung c\u1ea5p, h\u1ed7 tr\u1ee3 STT/TTS, OCR, eKYC, ph\u00e2n lo\u1ea1i gi\u1ea5y t\u1edd, nh\u1eadn di\u1ec7n c\u1ea3m x\u00fac','To\u00e0n qu\u1ed1c','REST API qua HTTPS, c\u00f3 SLA 99.5%'],
         ['H\u1ed3 s\u01a1 gi\u1ea5y t\u1edd l\u01b0u tr\u1eef','T\u00e0i li\u1ec7u v\u1eadt l\u00fd \u0111ang l\u01b0u t\u1ea1i UBND ph\u01b0\u1eddng — CCCD, s\u1ed5 h\u1ed9 kh\u1ea9u, gi\u1ea5y khai sinh, \u0111\u01a1n t\u1eeb h\u00e0nh ch\u00ednh','T\u1eebng UBND','Scan qua thi\u1ebft b\u1ecb Kiosk ho\u1eb7c m\u00e1y scan v\u0103n ph\u00f2ng'],
         ['C\u01a1 s\u1edf d\u1eef li\u1ec7u DVC Qu\u1ed1c gia','CSDL t\u1eadp trung c\u1ee7a C\u1ed5ng DVC Qu\u1ed1c gia — danh m\u1ee5c th\u1ee7 t\u1ee5c, bi\u1ec3u m\u1eabu, quy tr\u00ecnh x\u1eed l\u00fd','To\u00e0n qu\u1ed1c','API qua VNPT DVC Platform (\u0111ang \u0111\u00e0m ph\u00e1n)'],
         ['D\u1eef li\u1ec7u v\u1eadn h\u00e0nh n\u1ed9i b\u1ed9','D\u1eef li\u1ec7u giao d\u1ecbch, l\u1ecbch s\u1eed tra c\u1ee9u, ph\u1ea3n h\u1ed3i ng\u01b0\u1eddi d\u00f9ng — t\u00edch l\u0169y trong qu\u00e1 tr\u00ecnh v\u1eadn h\u00e0nh','N\u1ed9i b\u1ed9 h\u1ec7 th\u1ed1ng','PostgreSQL + Redis cache']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)
    add_heading(doc, '4.2 Nh\u00e2n l\u1ef1c', 2)
    add_para(doc, '\u0110\u1ed9i ng\u0169 ph\u00e1t tri\u1ec3n GovOne g\u1ed3m 5 vai tr\u00f2 ch\u00ednh, v\u1edbi t\u1ed5ng quy m\u00f4 5-7 ng\u01b0\u1eddi trong giai \u0111o\u1ea1n MVP:')
    make_table(doc,
        ['Vai tr\u00f2', 'S\u1ed1 l\u01b0\u1ee3ng', 'K\u1ef9 n\u0103ng ch\u00ednh', 'Kinh nghi\u1ec7m t\u1ed1i thi\u1ec3u'],
        [['PM (Product Manager)','1','Qu\u1ea3n l\u00fd s\u1ea3n ph\u1ea9m AI, agile, giao ti\u1ebfp kh\u00e1ch h\u00e0ng Gov','3 n\u0103m PM'],
         ['AI/Voice Engineer','1-2','STT/TTS, NLP, Python, API VNPT, x\u1eed l\u00fd h\u1ed9i tho\u1ea1i','2 n\u0103m AI'],
         ['OCR/Document AI Engineer','1','OCR, Document AI, Python, x\u1eed l\u00fd \u1ea3nh OpenCV','2 n\u0103m OCR'],
         ['Fullstack Developer','1-2','React/Next.js, Node.js/FastAPI, PostgreSQL, Docker','3 n\u0103m fullstack'],
         ['UI/UX Designer','1','Thi\u1ebft k\u1ebf giao di\u1ec7n Kiosk, WCAG, prototype Figma, nghi\u00ean c\u1ee9u ng\u01b0\u1eddi d\u00f9ng','2 n\u0103m UX']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 4.3 Ki\u1ebfn tr\u00fac k\u1ef9 thu\u1eadt
    add_heading(doc, '4.3 Ki\u1ebfn tr\u00fac k\u1ef9 thu\u1eadt', 2)
    add_para(doc, 'GovOne \u0111\u01b0\u1ee3c x\u00e2y d\u1ef1ng tr\u00ean n\u1ec1n t\u1ea3ng web hi\u1ec7n \u0111\u1ea1i, v\u1edbi ki\u1ebfn tr\u00fac microservices ph\u00eda backend v\u00e0 giao di\u1ec7n ng\u01b0\u1eddi d\u00f9ng linh ho\u1ea1t:')
    add_bullet(doc, 'Frontend: ', 'React/Next.js (TypeScript) \u2014 h\u1ed7 tr\u1ee3 SSR, SEO, v\u00e0 PWA. Th\u01b0 vi\u1ec7n giao di\u1ec7n Shadcn/UI, Tailwind CSS. T\u01b0\u01a1ng th\u00edch Kiosk m\u00e0n h\u00ecnh c\u1ea3m \u1ee9ng, Web App v\u00e0 Mobile App.')
    add_bullet(doc, 'Backend API: ', 'FastAPI (Python) cho AI/OCR endpoints \u2014 t\u1eadn d\u1ee5ng th\u01b0 vi\u1ec7n Python NLP, OpenCV. Node.js/Express cho API nghi\u1ec7p v\u1ee5 \u2014 real-time WebSocket, x\u1eed l\u00fd b\u1ea5t \u0111\u1ed3ng b\u1ed9.')
    add_bullet(doc, 'Database: ', 'PostgreSQL (15+) \u2014 l\u01b0u tr\u1eef giao d\u1ecbch, bi\u1ec3u m\u1eabu, l\u1ecbch s\u1eed. Redis \u2014 cache h\u00e0ng \u0111\u1ee3i STT/TTS, session management.')
    add_bullet(doc, 'File storage: ', 'MinIO (S3-compatible) \u2014 l\u01b0u tr\u1eef scan g\u1ed1c, \u1ea3nh gi\u1ea5y t\u1edd, log \u00e2m thanh. Object storage cluster 3 node.')
    add_bullet(doc, 'Container & Orchestration: ', 'Docker Compose (MVP) \u2192 Kubernetes (production). CI/CD qua GitLab Runner, t\u1ef1 \u0111\u1ed9ng build/test/deploy.')
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 4.4 MVP 7 ng\u00e0y
    add_heading(doc, '4.4 MVP 7 ng\u00e0y', 2)
    add_para(doc, 'Ch\u00fang t\u00f4i cam k\u1ebft ho\u00e0n th\u00e0nh MVP trong 7 ng\u00e0y v\u1edbi l\u1ed9 tr\u00ecnh ph\u00e2n chia theo t\u1eebng module:')
    make_table(doc,
        ['Ng\u00e0y', 'Module', 'C\u00f4ng vi\u1ec7c ch\u00ednh', 'K\u1ebft qu\u1ea3'],
        [['Ng\u00e0y 1-2', 'Voice + OCR Core',
          '\u00b7 T\u00edch h\u1ee3p SmartVoice STT\u00b7 SmartReader OCR\u00b7 eKYC Compare',
          'Voice capture + OCR scan + eKYC ho\u1ea1t \u0111\u1ed9ng'],
         ['Ng\u00e0y 3-4', 'Smartbot + eKYC',
          '\u00b7 X\u00e2y d\u1ef1ng Intent Engine\u00b7 T\u00edch h\u1ee3p Smartbot NLP\u00b7 eKYC Liveness',
          'Bot tr\u1ea3 l\u1eddi th\u1ee7 t\u1ee5c + x\u00e1c th\u1ef1c khu\u00f4n m\u1eb7t'],
         ['Ng\u00e0y 5-6', 'UI & Dashboard',
          '\u00b7 Thi\u1ebft k\u1ebf giao di\u1ec7n Kiosk\u00b7 Dashboard c\u00e1n b\u1ed9\u00b7 K\u1ebft n\u1ed1i API',
          'Giao di\u1ec7n voice-first + dashboard s\u1ed1 li\u1ec7u'],
         ['Ng\u00e0y 7', 'T\u00edch h\u1ee3p & Test',
          '\u00b7 Ch\u1ea1y end-to-end flow\u00b7 Fix bug\u00b7 Chu\u1ea9n b\u1ecb demo',
          'MVP ho\u00e0n ch\u1ec9nh, s\u1eb5n s\u00e0ng demo kh\u00e1ch h\u00e0ng']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 4.5 Chi ph\u00ed v\u1eadn h\u00e0nh
    add_heading(doc, '4.5 Chi ph\u00ed v\u1eadn h\u00e0nh', 2)
    add_para(doc, 'Chi ph\u00ed v\u1eadn h\u00e0nh h\u1ec7 th\u1ed1ng GovOne \u01b0\u1edbc t\u00ednh kho\u1ea3ng 3.7-5.7 tri\u1ec7u \u0111\u1ed3ng/th\u00e1ng, ph\u00f9 h\u1ee3p v\u1edbi ng\u00e2n s\u00e1ch CNTT c\u1ee7a UBND ph\u01b0\u1eddng/x\u00e3:')
    make_table(doc,
        ['H\u1ea1ng m\u1ee5c', 'Chi ph\u00ed (VN\u0110/th\u00e1ng)', 'Ghi ch\u00fa'],
        [['API VNPT AI (SmartVoice, SmartReader, eKYC, SmartVision)', '2.000.000 - 3.000.000', 'T\u00ednh theo l\u01b0\u1ee3t g\u1ecdi API, g\u00f3i c\u01a1 b\u1ea3n 10.000 l\u01b0\u1ee3t/th\u00e1ng'],
         ['Server Cloud (VPS 4CPU/8GB/100GB SSD)', '800.000 - 1.200.000', 'VPS ho\u1eb7c cloud instance, c\u00f3 th\u1ec3 d\u00f9ng server UBND s\u1eb5n c\u00f3'],
         ['MinIO Object Storage (500GB)', '300.000 - 500.000', 'L\u01b0u tr\u1eef scan g\u1ed1c, c\u00f3 th\u1ec3 d\u00f9ng NAS n\u1ed9i b\u1ed9'],
         ['T\u00ean mi\u1ec1n + SSL + Backup', '200.000 - 400.000', 'T\u00ean mi\u1ec1n .gov.vn \u0111\u01b0\u1ee3c h\u1ed7 tr\u1ee3, SSL mi\u1ec5n ph\u00ed Let\'s Encrypt'],
         ['\u0110i\u1ec7n + B\u1ea3o tr\u00ec thi\u1ebft b\u1ecb Kiosk', '400.000 - 600.000', 'Kiosk ch\u1ea1y mini PC, \u0111i\u1ec7n ~50W'],
         ['T\u1ed5ng c\u1ed9ng', '3.700.000 - 5.700.000', 'Ch\u01b0a bao g\u1ed3m nh\u00e2n s\u1ef1 v\u1eadn h\u00e0nh']],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 4.6 An to\u00e0n b\u1ea3o m\u1eadt & Ph\u00e1p l\u00fd
    add_heading(doc, '4.6 An to\u00e0n b\u1ea3o m\u1eadt & Ph\u00e1p l\u00fd', 2)
    add_para(doc, 'GovOne tu\u00e2n th\u1ee7 c\u00e1c quy \u0111\u1ecbnh ph\u00e1p lu\u1eadt hi\u1ec7n h\u00e0nh v\u1ec1 b\u1ea3o v\u1ec7 d\u1eef li\u1ec7u c\u00e1 nh\u00e2n v\u00e0 an to\u00e0n th\u00f4ng tin:')
    make_table(doc,
        ['V\u0103n b\u1ea3n ph\u00e1p l\u00fd', 'N\u1ed9i dung y\u00eau c\u1ea7u', 'Ph\u01b0\u01a1ng \u00e1n tu\u00e2n th\u1ee7 c\u1ee7a GovOne'],
        [['Ngh\u1ecb \u0111\u1ecbnh 13/2023/N\u0110-CP (B\u1ea3o v\u1ec7 d\u1eef li\u1ec7u c\u00e1 nh\u00e2n)',
          'Thu th\u1eadp, l\u01b0u tr\u1eef, x\u1eed l\u00fd d\u1eef li\u1ec7u c\u00e1 nh\u00e2n ph\u1ea3i c\u00f3 s\u1ef1 \u0111\u1ed3ng \u00fd; th\u00f4ng b\u00e1o m\u1ee5c \u0111\u00edch; \u0111\u1ea3m b\u1ea3o an to\u00e0n',
          '\u00b7 C\u00f3 m\u00e0n h\u00ecnh x\u00e1c nh\u1eadn \u0111\u1ed3ng \u00fd tr\u01b0\u1edbc khi scan\u00b7 M\u00e3 h\u00f3a d\u1eef li\u1ec7u c\u00e1 nh\u00e2n (AES-256)\u00b7 Ph\u00e2n quy\u1ec1n truy c\u1eadp ch\u1eb7t ch\u1ebd\u00b7 Audit log \u0111\u1ea7y \u0111\u1ee7'],
         ['Ngh\u1ecb \u0111\u1ecbnh 59/2022/N\u0110-CP (\u0110\u1ecbnh danh v\u00e0 x\u00e1c th\u1ef1c \u0111i\u1ec7n t\u1eed)',
          'Quy \u0111\u1ecbnh v\u1ec1 nh\u1eadn d\u1ea1ng, x\u00e1c th\u1ef1c ng\u01b0\u1eddi d\u00f9ng qua t\u00e0i kho\u1ea3n \u0111\u1ecbnh danh \u0111i\u1ec7n t\u1eed',
          '\u00b7 T\u00edch h\u1ee3p VNeID l\u00e0m ph\u01b0\u01a1ng th\u1ee9c x\u00e1c th\u1ef1c ch\u00ednh\u00b7 eKYC Liveness \u0111\u00e1p \u1ee9ng y\u00eau c\u1ea7u x\u00e1c th\u1ef1c sinh tr\u1eafc'],
         ['Lu\u1eadt An to\u00e0n th\u00f4ng tin m\u1ea1ng 2015',
          '\u0110\u1ea3m b\u1ea3o an to\u00e0n h\u1ec7 th\u1ed1ng th\u00f4ng tin; b\u1ea3o v\u1ec7 d\u1eef li\u1ec7u; ph\u00f2ng ch\u1ed1ng t\u1ea5n c\u00f4ng m\u1ea1ng',
          '\u00b7 Firewall + WAF\u00b7 Security audit \u0111\u1ecbnh k\u1ef3\u00b7 Backup d\u1eef li\u1ec7u 3-2-1\u00b7 ISO 27001 (l\u1ed9 tr\u00ecnh)']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 4.7 L\u1ed9 tr\u00ecnh 12 th\u00e1ng
    add_heading(doc, '4.7 L\u1ed9 tr\u00ecnh 12 th\u00e1ng', 2)
    add_para(doc, 'Sau khi ho\u00e0n th\u00e0nh MVP, GovOne s\u1ebd tri\u1ec3n khai m\u1edf r\u1ed9ng theo l\u1ed9 tr\u00ecnh 12 th\u00e1ng:')
    add_bullet(doc, 'Th\u00e1ng 1-2 \u2014 Pilot: ', 'Tri\u1ec3n khai th\u00ed \u0111i\u1ec3m t\u1ea1i 2-3 UBND ph\u01b0\u1eddng t\u1ea1i TP.HCM. Thu th\u1eadp ph\u1ea3n h\u1ed3i, \u0111o l\u01b0\u1eddng KPI (th\u1eddi gian x\u1eed l\u00fd, \u0111\u1ed9 ch\u00ednh x\u00e1c OCR, t\u1ef7 l\u1ec7 h\u00e0i l\u00f2ng). \u0110i\u1ec1u ch\u1ec9nh s\u1ea3n ph\u1ea9m theo th\u1ef1c t\u1ebf.')
    add_bullet(doc, 'Th\u00e1ng 3-4 \u2014 Scale: ', 'M\u1edf r\u1ed9ng l\u00ean 10 UBND t\u1ea1i TP.HCM v\u00e0 H\u00e0 N\u1ed9i. T\u1ed1i \u01b0u chi ph\u00ed API, c\u1ea3i thi\u1ec7n \u0111\u1ed9 ch\u00ednh x\u00e1c Smartbot (t\u1eeb 85% \u2192 92%). B\u1ed5 sung th\u00eam 20 th\u1ee7 t\u1ee5c h\u00e0nh ch\u00ednh ph\u1ed5 bi\u1ebfn.')
    add_bullet(doc, 'Th\u00e1ng 5-6 \u2014 Public Beta: ', 'M\u1edf r\u1ed9ng ra 50 UBND tr\u00ean c\u1ea3 n\u01b0\u1edbc. Ra m\u1eaft Web App cho ng\u01b0\u1eddi d\u00e2n t\u1ef1 tra c\u1ee9u t\u1ea1i nh\u00e0. H\u1ee3p t\u00e1c v\u1edbi S\u1edf TT&TT \u0111\u1ecba ph\u01b0\u01a1ng \u0111\u1ec3 \u0111\u01b0\u1ee3c h\u1ed7 tr\u1ee3 ch\u00ednh s\u00e1ch.')
    add_bullet(doc, 'Th\u00e1ng 7-12 \u2014 M\u1edf r\u1ed9ng: ', 'Nh\u00e2n r\u1ed9ng ra 200+ UBND. N\u00e2ng c\u1ea5p AI Core: x\u1eed l\u00fd th\u00eam 50 th\u1ee7 t\u1ee5c, h\u1ed7 tr\u1ee3 \u0111a ng\u00f4n ng\u1eef (Anh, Hoa). M\u1edf r\u1ed9ng l\u00ean c\u1ea5p qu\u1eadn/huy\u1ec7n v\u00e0 s\u1edf ban ng\u00e0nh.')

    doc.save(doc_path)
    print('\u2705 Section 4: "T\u00ednh kh\u1ea3 thi" added (7 subsections)')

if __name__ == '__main__':
    main()
