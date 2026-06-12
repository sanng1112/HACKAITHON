#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16, 3: 14}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_bullet(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('\u2022 '); set_font(r, size=13)
    r = p.add_run(bold_part); set_font(r, size=13, bold=True)
    r = p.add_run(normal_part); set_font(r, size=13)

def add_sub_heading(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text); set_font(r, size=14, bold=True, color=RGBColor(0x00, 0x66, 0xCC))

def make_table(doc, headers, data, col_aligns=None):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            if col_aligns and ci < len(col_aligns): p.alignment = col_aligns[ci]
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct); set_font(r, size=10)
            if ri % 2 == 1: c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))
    return table

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    doc_path = os.path.join(project_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '5. \u0110\u1ed4I M\u1edaI & KH\u00c1C BI\u1ec6T')

    # 5.1 So s\u00e1nh v\u1edbi gi\u1ea3i ph\u00e1p hi\u1ec7n t\u1ea1i
    add_heading(doc, '5.1 So s\u00e1nh v\u1edbi gi\u1ea3i ph\u00e1p hi\u1ec7n t\u1ea1i', 2)
    add_para(doc, 'GovOne \u0111\u01b0\u1ee3c so s\u00e1nh v\u1edbi 3 nh\u00f3m gi\u1ea3i ph\u00e1p \u0111ang hi\u1ec7n di\u1ec7n tr\u00ean th\u1ecb tr\u01b0\u1eddng \u2014 Chatbot DVC, OCR truy\u1ec1n th\u1ed1ng v\u00e0 Gia c\u00f4ng CNTT \u2014 theo 7 ti\u00eau ch\u00ed \u0111\u00e1nh gi\u00e1:')
    make_table(doc,
        ['Ti\u00eau ch\u00ed', 'Chatbot DVC', 'OCR truy\u1ec1n th\u1ed1ng', 'Gia c\u00f4ng CNTT', 'GovOne (\u0110\u1ec1 xu\u1ea5t)'],
        [['K\u00eanh t\u01b0\u01a1ng t\u00e1c', 'Text chat', 'Kh\u00f4ng c\u00f3', 'Theo y\u00eau c\u1ea7u', 'Voice-first + Vision'],
         ['Ng\u01b0\u1eddi d\u00f9ng m\u1ee5c ti\u00eau', 'Ng\u01b0\u1eddi tr\u1ebb, bi\u1ebft g\u00f5', 'C\u00e1n b\u1ed9 v\u0103n th\u01b0', 'Doanh nghi\u1ec7p CNTT', 'Ng\u01b0\u1eddi gi\u00e0, khuy\u1ebft t\u1eadt + c\u00e1n b\u1ed9'],
         ['X\u1eed l\u00fd gi\u1ea5y t\u1edd', 'Kh\u00f4ng', 'OCR \u0111\u01a1n gi\u1ea3n, 1 lo\u1ea1i', 'Theo d\u1ef1 \u00e1n', '\u0110a d\u1ea1ng gi\u1ea5y t\u1edd, t\u1ef1 \u0111\u1ed9ng ph\u00e2n lo\u1ea1i'],
         ['X\u00e1c th\u1ef1c danh t\u00ednh', 'Kh\u00f4ng', 'Kh\u00f4ng', 'Th\u1ee7 c\u00f4ng', 'eKYC Compare + Liveness'],
         ['\u0110o l\u01b0\u1eddng h\u00e0i l\u00f2ng', 'Kh\u00f4ng', 'Kh\u00f4ng', 'Kh\u00f4ng', 'Sentiment AI qua camera'],
         ['T\u00edch h\u1ee3p CSDL', 'R\u1eddi r\u1ea1c', 'C\u1ee5c b\u1ed9', 'Th\u1ee7 c\u00f4ng', '\u0110\u1ed3ng b\u1ed9 DVC Qu\u1ed1c gia + n\u1ed9i b\u1ed9'],
         ['Chi ph\u00ed tri\u1ec3n khai', '50-100tr', '100-300tr', '300-500tr/d\u1ef1 \u00e1n', '150-250tr (tr\u1ecdn g\u00f3i)']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 5.2 B\u1ed1n \u0111i\u1ec3m \u0111\u1ed5i m\u1edbi c\u1ed1t l\u00f5i
    add_heading(doc, '5.2 B\u1ed1n \u0111i\u1ec3m \u0111\u1ed5i m\u1edbi c\u1ed1t l\u00f5i', 2)
    add_para(doc, 'GovOne t\u1ea1o ra kh\u00e1c bi\u1ec7t so v\u1edbi c\u00e1c gi\u1ea3i ph\u00e1p hi\u1ec7n t\u1ea1i nh\u1edd 4 \u0111i\u1ec3m \u0111\u1ed5i m\u1edbi c\u1ed1t l\u00f5i:', indent=False)

    add_sub_heading(doc, '5.2.1 Voice-First + Vision \u2014 Giao ti\u1ebfp t\u1ef1 nhi\u00ean nh\u1ea5t')
    add_para(doc, 'Kh\u00f4ng gi\u1ed1ng c\u00e1c chatbot DVC hi\u1ec7n t\u1ea1i bu\u1ed9c ng\u01b0\u1eddi d\u00e2n ph\u1ea3i g\u00f5 ch\u1eef, GovOne \u0111\u1eb7t gi\u1ecdng n\u00f3i l\u00e0m k\u00eanh t\u01b0\u01a1ng t\u00e1c ch\u00ednh. K\u1ebft h\u1ee3p v\u1edbi camera AI (SmartVision), h\u1ec7 th\u1ed1ng c\u00f3 th\u1ec3: (a) ph\u00e1t hi\u1ec7n ng\u01b0\u1eddi d\u00f9ng khi \u0111\u1ebfn g\u1ea7n, (b) ph\u00e2n t\u00edch c\u1ea3m x\u00fac \u0111\u1ec3 \u0111i\u1ec1u ch\u1ec9nh c\u00e1ch h\u01b0\u1edbng d\u1eabn, (c) nh\u1eadn di\u1ec7n c\u1eed ch\u1ec9 tay \u0111\u1ec3 h\u1ed7 tr\u1ee3 ng\u01b0\u1eddi khuy\u1ebft t\u1eadt. \u0110\u00e2y l\u00e0 gi\u1ea3i ph\u00e1p \u201cZero UI\u201d \u0111\u1ea7u ti\u00ean cho h\u00e0nh ch\u00ednh c\u00f4ng t\u1ea1i Vi\u1ec7t Nam.')

    add_sub_heading(doc, '5.2.2 Zero UI \u2014 Thi\u1ebft k\u1ebf cho t\u1ea5t c\u1ea3 m\u1ecdi ng\u01b0\u1eddi')
    add_para(doc, 'Ng\u01b0\u1eddi d\u00e2n kh\u00f4ng c\u1ea7n ch\u1ea1m m\u00e0n h\u00ecnh, kh\u00f4ng c\u1ea7n g\u00f5 ph\u00edm, kh\u00f4ng c\u1ea7n \u0111\u1ecdc h\u01b0\u1edbng d\u1eabn. Giao di\u1ec7n Kiosk \u0111\u01b0\u1ee3c thi\u1ebft k\u1ebf theo nguy\u00ean t\u1eafc WCAG AA: font ch\u1eef t\u1ed1i thi\u1ec3u 18px, \u0111\u1ed9 t\u01b0\u01a1ng ph\u1ea3n 4.5:1, n\u00fat b\u1ea5m t\u1ed1i thi\u1ec3u 48x48px. M\u1ecdi thao t\u00e1c \u0111\u1ec1u c\u00f3 th\u1ec3 th\u1ef1c hi\u1ec7n b\u1eb1ng gi\u1ecdng n\u00f3i, v\u1edbi x\u00e1c nh\u1eadn b\u1eb1ng \u00e2m thanh. Ng\u01b0\u1eddi m\u00f9, ng\u01b0\u1eddi gi\u00e0, ng\u01b0\u1eddi kh\u00f4ng bi\u1ebft ch\u1eef \u0111\u1ec1u c\u00f3 th\u1ec3 s\u1eed d\u1ee5ng.')

    add_sub_heading(doc, '5.2.3 Orchestration 7 API \u2014 N\u1ec1n t\u1ea3ng h\u1ee3p nh\u1ea5t')
    add_para(doc, 'GovOne kh\u00f4ng ph\u1ea3i l\u00e0 s\u1ea3n ph\u1ea9m \u0111\u01a1n l\u1ebb \u2014 l\u00e0 n\u1ec1n t\u1ea3ng orchestration k\u1ebft h\u1ee3p 7 API AI c\u1ee7a VNPT (SmartVoice STT/TTS, Smartbot, SmartReader, eKYC Compare/Liveness, SmartVision Classification/Face/Sentiment) trong m\u1ed9t h\u1ec7 sinh th\u00e1i th\u1ed1ng nh\u1ea5t. Kh\u00f4ng c\u1ea7n chuy\u1ec3n \u0111\u1ed5i gi\u1eefa nhi\u1ec1u \u1ee9ng d\u1ee5ng, m\u1ecdi d\u1eef li\u1ec7u \u0111\u01b0\u1ee3c \u0111\u1ed3ng b\u1ed9 real-time. Ki\u1ebfn tr\u00fac microservices cho ph\u00e9p m\u1ed7i module ho\u1ea1t \u0111\u1ed9ng \u0111\u1ed9c l\u1eadp, d\u1ec5 d\u00e0ng n\u00e2ng c\u1ea5p ho\u1eb7c thay th\u1ebf.')

    add_sub_heading(doc, '5.2.4 Feedback Loop \u2014 C\u1ea3i ti\u1ebfn li\u00ean t\u1ee5c')
    add_para(doc, 'Kh\u00e1c v\u1edbi c\u00e1c gi\u1ea3i ph\u00e1p \u201cc\u00e0i xong \u0111\u1ec3 \u0111\u1ea5y\u201d, GovOne c\u00f3 v\u00f2ng l\u1eb7p ph\u1ea3n h\u1ed3i t\u1ef1 \u0111\u1ed9ng: Sentiment AI \u0111o m\u1ee9c \u0111\u1ed9 h\u00e0i l\u00f2ng sau m\u1ed7i giao d\u1ecbch \u2192 D\u1eef li\u1ec7u \u0111\u01b0\u1ee3c ph\u00e2n t\u00edch \u2192 Dashboard c\u00e1n b\u1ed9 hi\u1ec3n th\u1ecb xu h\u01b0\u1edbng \u2192 \u0110i\u1ec1u ch\u1ec9nh k\u1ecbch b\u1ea3n Smartbot v\u00e0 quy tr\u00ecnh x\u1eed l\u00fd. \u0110\u00e2y l\u00e0 c\u01a1 ch\u1ebf Kaizen (c\u1ea3i ti\u1ebfn li\u00ean t\u1ee5c) \u1ee9ng d\u1ee5ng trong h\u00e0nh ch\u00ednh c\u00f4ng, gi\u00fap ch\u1ea5t l\u01b0\u1ee3ng d\u1ecbch v\u1ee5 ng\u00e0y c\u00e0ng \u0111\u01b0\u1ee3c n\u00e2ng cao.')

    doc.save(doc_path)
    print('\u2705 Section 5: \"\u0110\u1ed5i m\u1edbi & Kh\u00e1c bi\u1ec7t\" added (comparison + 4 innovations)')

if __name__ == '__main__':
    main()
