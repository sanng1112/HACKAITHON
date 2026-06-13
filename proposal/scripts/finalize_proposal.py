#!/usr/bin/env python3
import os, subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_heading(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); set_font(r, size=18, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
        '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_header_footer(doc):
    for section in doc.sections:
        header = section.header; header.is_linked_to_previous = False
        hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run('GovOne — Đội thi [Tên đội] — Hackathon ĐMST 2026')
        set_font(r, size=10, color=RGBColor(0x66, 0x66, 0x66))
        footer = section.footer; footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run('Trang '); set_font(r, size=10, color=RGBColor(0x66, 0x66, 0x66))
        for ftype in ['PAGE', 'NUMPAGES']:
            r2 = fp.add_run(); fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); r2._r.append(fc1)
            r3 = fp.add_run(); it1 = OxmlElement('w:instrText'); it1.set(qn('xml:space'), 'preserve'); it1.text = f' {ftype} '; r3._r.append(it1)
            r4 = fp.add_run(); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); r4._r.append(fc2)
            if ftype == 'PAGE': r5 = fp.add_run(' / Tổng '); set_font(r5, size=10, color=RGBColor(0x66, 0x66, 0x66))

def add_conclusion(doc):
    add_heading(doc, '7. KẾT LUẬN')
    add_para(doc, 'GovOne là nền tảng quản lý hành chính công thông minh đầu tiên tại Việt Nam tích hợp 7 API AI cốt lõi của VNPT — SmartVoice, Smartbot, SmartReader, eKYC và SmartVision — trong một hệ thống duy nhất, phục vụ đồng thời cả người dân và cán bộ.')
    for bold_p, normal_p in [
        ('1. Giải quyết 4 pain-point cốt lõi:',' Voice + NLP giúp người dân nói thay vì gõ. OCR + Doc AI số hóa hồ sơ tồn đọng. eKYC tự động xác thực. Sentiment AI đo lường hài lòng.'),
        ('2. Khác biệt 100% so với giải pháp hiện tại:',' Voice-first + Vision, Zero UI, orchestration 7 API VNPT, pipeline OCR thông minh, vòng phản hồi tự động — không sản phẩm nào có đầy đủ 5 điểm này.'),
        ('3. Tính khả thi cao, tác động rõ ràng:',' Chi phí từ 3.7 triệu/tháng, MVP 7 ngày, tuân thủ pháp lý. TAM ~18.000 tỷ, ROI 3 năm ~250%, giảm 70% thời gian giao dịch.')]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(bold_p); set_font(r, size=14, bold=True); r = p.add_run(normal_p); set_font(r, size=14)
    add_para(doc, 'Chúng tôi kêu gọi sự hợp tác của VNPT và các cơ quan nhà nước để đưa GovOne đến mọi bộ phận một cửa trên cả nước, góp phần xây dựng nền hành chính công hiện đại, không rào cản, không ai bị bỏ lại phía sau.')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('\nGovOne — Hành chính công thông minh.\n"Không chạm, không gõ, không rào cản."')
    set_font(r, size=13, color=RGBColor(0x00, 0x66, 0xCC))

def convert_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path], check=True, timeout=60, capture_output=True)
        print(f'✅ PDF exported: {pdf_path}'); return True
    except: print('⚠️ PDF auto-export unavailable. Export manually.'); return False

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    docx_path = os.path.join(project_dir, 'proposal.docx')
    pdf_path = os.path.join(project_dir, 'proposal.pdf')
    doc = Document(docx_path)
    print('📝 Adding header/footer...'); add_header_footer(doc)
    print('📝 Adding conclusion...'); add_conclusion(doc)
    doc.save(docx_path)
    print(f'✅ Final proposal saved: {docx_path} ({os.path.getsize(docx_path)/1024:.1f} KB)')
    convert_to_pdf(docx_path, pdf_path)
    if os.path.exists(pdf_path): print(f'📊 PDF size: {os.path.getsize(pdf_path)/1024:.1f} KB')
    print('\n🎯 GovOne Proposal finalized!')

if __name__ == '__main__':
    main()
