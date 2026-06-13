import os, sys, unittest
from docx import Document

SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(SCRIPT_DIR, 'scripts'))

class TestGovOne(unittest.TestCase):
    def setUp(self):
        self.project_dir = SCRIPT_DIR
        self.assets_dir = os.path.join(self.project_dir, 'assets')

    def test_assets_directory_exists(self):
        self.assertTrue(os.path.isdir(self.assets_dir))

    def test_logo_generated(self):
        logo_path = os.path.join(self.assets_dir, 'logo-govone.png')
        self.assertTrue(os.path.isfile(logo_path))
        from PIL import Image
        img = Image.open(logo_path)
        self.assertEqual(img.size, (400, 400))
        self.assertEqual(img.mode, 'RGBA')

    def test_proposal_created(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        self.assertTrue(os.path.isfile(proposal_path))
        doc = Document(proposal_path)
        self.assertGreaterEqual(len(doc.paragraphs), 10)

    def test_cover_page_has_title(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        doc = Document(proposal_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any('GovOne' in t for t in texts))

    def test_cover_page_has_topic(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        doc = Document(proposal_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any('Đề tài 6' in t for t in texts))

    def test_section1_exists(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any('ĐẶT VẤN ĐỀ' in t for t in texts))

    def test_section1_has_painpoints_table(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        for table in doc.tables:
            if '#' in table.rows[0].cells[0].text:
                self.assertGreaterEqual(len(table.rows), 5)
                return
        self.fail("Không tìm thấy bảng pain-point")

    def test_section1_has_why_ai(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any('TẠI SAO AI' in t for t in texts))

    def test_section2_exists(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any('GIẢI PHÁP' in t for t in texts))

    def test_section2_has_voice_and_ocr(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text.upper() for p in doc.paragraphs]
        self.assertTrue(any('STT' in t for t in texts) and any('OCR' in t for t in texts))

    def test_section2_has_user_story(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any('Bước 1' in t for t in texts))

    def test_section3_exists(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip()]
        self.assertTrue(any('THIẾT KẾ' in t for t in texts))

    def test_section3_has_4_layers(self):
        doc = Document(os.path.join(self.project_dir, 'proposal.docx'))
        texts_upper = [p.text.upper() for p in doc.paragraphs]
        for layer in ['USER LAYER', 'AI CORE', 'PROCESSING', 'DATA LAYER']:
            self.assertTrue(any(layer in t for t in texts_upper))

if __name__ == '__main__':
    unittest.main()
