#!/usr/bin/env python3
"""Add section 5. DOI MOI & KHAC BIET to AutoCheck proposal."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16, 3: 14}.get(lv, 14)
    r = p.add_run(t); sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def mt(doc, hd, data):
    t = doc.add_table(rows=1+len(data), cols=len(hd))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(h)
        sf(r, s=11, b=True, c=RGBColor(0xFF,0xFF,0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT; r = p.add_run(ct); sf(r, s=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(sd, 'proposal.docx'))
    apb(doc); ah(doc, '5. TINH DOI MOI & KHAC BIET')

    ah(doc, '5.1 So sanh voi giai phap hien tai', 2)
    ap(doc, 'So sanh AutoCheck voi cac giai phap so hoa tai lieu hien co:')
    mt(doc, ['Tieu chi', 'Giai phap hien tai', 'AutoCheck'], [
        ['Toc do xu ly', '30-60 phut/ho so', '2-5 phut (nhanh 10-15x)'],
        ['Do chinh xac', 'Phu thuoc tay nghe', '>95% (AI OCR + doi chieu)'],
        ['Phan loai', 'Thu cong', 'AI tu dong (SmartVision)'],
        ['Boc tach', 'Nhap lieu thu cong', 'AI tu dong, co cau truc'],
        ['Doi chieu CSDL', 'Tra cuu thu cong', 'AI tu dong, real-time'],
        ['Phat hien sai lech', 'Mat thuong', 'AI tu dong + canh bao'],
        ['Luu tru', 'Kho giay', 'So hoa, backup, tim kiem nhanh'],
    ])
    ap(doc, 'AutoCheck vuot troi 7/7 tieu chi (~100% khac biet), vuot xa nguong 30% yeu cau.')

    ah(doc, '5.2 Bon diem doi moi cot loi', 2)

    ah(doc, '5.2.1 AI OCR hieu ngu canh giay to', 3)
    ap(doc, 'SmartReader Doc AI hieu cau truc tung loai giay to, trich xuat co cau truc.')

    ah(doc, '5.2.2 Tu dong phan loai & dinh tuyen', 3)
    ap(doc, 'SmartVision phan loai giay to ngay sau scan, dinh tuyen quy trinh rieng.')

    ah(doc, '5.2.3 Doi chieu thong minh & phat hien sai lech', 3)
    ap(doc, 'Doi chieu tu dong voi CSDL, canh bao tung truong sai lech.')

    ah(doc, '5.2.4 Pipeline xu ly hang loat', 3)
    ap(doc, 'Kien truc queue-based, xu ly den 10.000 ho so/ngay.')

    doc.save(os.path.join(sd, 'proposal.docx'))
    print('Task 6: Section "Doi moi & Khac biet" added')


if __name__ == '__main__':
    main()
