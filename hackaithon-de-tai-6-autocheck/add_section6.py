#!/usr/bin/env python3
"""Add section 6. TAC DONG DU KIEN to AutoCheck proposal."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(lv, 14)
    r = p.add_run(t); sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def mt(doc, hd, data):
    t = doc.add_table(rows=1+len(data), cols=len(hd))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(h)
        sf(r, s=11, b=True, c=RGBColor(0xFF,0xFF,0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT; r = p.add_run(ct); sf(r, s=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(sd, 'proposal.docx'))
    apb(doc); ah(doc, '6. TAC DONG DU KIEN')

    ah(doc, '6.1 Phan tich thi truong (TAM-SAM-SOM)', 2)
    mt(doc, ['Chi so', 'Gia tri', 'Cach tinh', 'Nguon'], [
        ['TAM', '~5.000 ty VND', 'So hoa tai lieu HC 63 tinh, ~80 ty/tinh/nam', 'Bo TT&TT 2025'],
        ['SAM', '~300 ty VND', 'OCR+AI cho ho so luu tru cap phuong (6% TAM)', 'Phan tich noi bo'],
        ['SOM', '~15 ty VND', '5% SAM trong 2 nam (~50-100 UBND)', 'Du bao than trong'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '6.2 Loi ich xa hoi', 2)
    mt(doc, ['Loi ich', 'Chi so', 'Giai thich'], [
        ['Tang toc xu ly ho so', 'Nhanh 10-15x', '30-60 phut -> 2-5 phut/ho so'],
        ['Tiet kiem nhan cong', 'Giam 60%', '1 he thong = 3-5 nhan vien'],
        ['Giam sai sot', 'Giam 90%', 'AI tu dong kiem tra + canh bao'],
        ['Tiet kiem khong gian', 'Giam 95%', 'Kho giay -> o cung 1TB ~500.000 ho so'],
        ['Tra cuu tuc thi', '< 5 giay', 'Full-text search tu CSDL so hoa'],
        ['Bao ton du lieu', 'Vinh vien', 'Backup tu dong, khong lo hu hong'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '6.3 Mo hinh doanh thu (B2G)', 2)
    mt(doc, ['Goi', 'Gia', 'Dich vu'], [
        ['Basic', '8.000.000/thang', '1 may scan, 1.000 ho so/thang, ho tro 8h'],
        ['Pro', '20.000.000/thang', 'Da may scan, 5.000 ho so/thang, 24/7'],
        ['Enterprise', 'Theo yeu cau', 'Khong gioi han, tich hop CSDL rieng, SLA'],
    ])
    ap(doc, 'Phi trien khai: 40-60 trieu. Hoa von 12 thang voi 15 KH Basic (~120 trieu/thang). ROI 3 nam: ~250%.')

    ah(doc, '6.4 Phan tich canh tranh', 2)
    ap(doc, 'Doi thu tren thi truong so hoa tai lieu:')
    for c in [
        'FPT.eDoc: OCR co ban, khong AI phan loai. Gia 15-30 trieu/thang.',
        'VNPT eDoc: OCR co ban, chua tich hop SmartReader Doc AI.',
        'Google Doc AI: OCR manh nhung chua toi uu tieng Viet, khong dat chuan VN.',
        'Nhan cong thu cong: 10-15 trieu/thang/nguoi, cham, sai sot.',
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
        r = p.add_run('* '); sf(r, s=13); r = p.add_run(c); sf(r, s=13)

    doc.save(os.path.join(sd, 'proposal.docx'))
    print('Task 7: Section "Tac dong du kien" added')


if __name__ == '__main__':
    main()
