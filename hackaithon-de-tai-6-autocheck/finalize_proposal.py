#!/usr/bin/env python3
"""Task 8: Finalize AutoCheck proposal - header, footer, conclusion, PDF."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os, subprocess


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def ah(doc, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(t); sf(r, s=18, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
    pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def add_header_footer(doc):
    for sec in doc.sections:
        h = sec.header; h.is_linked_to_previous = False
        hp = h.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run('AutoCheck - Doi thi [Ten doi] - Hackathon DMST 2026')
        sf(r, s=10, c=RGBColor(0x66, 0x66, 0x66))
        f = sec.footer; f.is_linked_to_previous = False
        fp = f.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run('Trang '); sf(r, s=10, c=RGBColor(0x66, 0x66, 0x66))
        for ftype in ['PAGE', 'NUMPAGES']:
            r2 = fp.add_run(); fc1 = OxmlElement('w:fldChar')
            fc1.set(qn('w:fldCharType'), 'begin'); r2._r.append(fc1)
            r3 = fp.add_run(); it1 = OxmlElement('w:instrText')
            it1.set(qn('xml:space'), 'preserve'); it1.text = f' {ftype} '; r3._r.append(it1)
            r4 = fp.add_run(); fc2 = OxmlElement('w:fldChar')
            fc2.set(qn('w:fldCharType'), 'end'); r4._r.append(fc2)
            if ftype == 'PAGE':
                r5 = fp.add_run(' / Tong '); sf(r5, s=10, c=RGBColor(0x66, 0x66, 0x66))


def add_conclusion(doc):
    ah(doc, '7. KET LUAN')
    ap(doc, 'AutoCheck la he thong OCR thong minh dau tien tai Viet Nam ket hop 3 cong nghe AI cot loi cua VNPT - SmartReader, SmartVision va eKYC - trong mot pipeline xu ly ho so luu tru tu dong.')
    for b, n in [
        ('1. Giai quyet 3 pain-point:', ' So hoa ho so ton dong, tu dong boc tach & doi chieu.'),
        ('2. Khac biet 100%:', ' AI OCR hieu ngu canh, phan loai, doi chieu thong minh.'),
        ('3. Tinh kha thi cao:', ' Chi phi tu 8 trieu/thang, MVP 7 ngay, TAM ~5.000 ty.'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(b); sf(r, s=14, b=True); r = p.add_run(n); sf(r, s=14)
    ap(doc, 'Chung toi keu goi su hop tac cua VNPT de dua AutoCheck den moi bo phan mot cua.')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('\nAutoCheck\n"So hoa hom nay - Gia tri ngay mai."')
    sf(r, s=13, c=RGBColor(0x00, 0x66, 0xCC))


def convert_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path), docx_path],
            check=True, timeout=60, capture_output=True)
        print(f'PDF exported: {pdf_path}'); return True
    except:
        print('PDF auto-export unavailable. Export manually from Word.'); return False


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    dx = os.path.join(sd, 'proposal.docx'); px = os.path.join(sd, 'proposal.pdf')
    doc = Document(dx)
    print('Adding header/footer...'); add_header_footer(doc)
    print('Adding conclusion...'); add_conclusion(doc)
    doc.save(dx)
    print(f'Final: {dx} ({os.path.getsize(dx)/1024:.1f} KB)')
    convert_to_pdf(dx, px)
    if os.path.exists(px): print(f'PDF: {os.path.getsize(px)/1024:.1f} KB')
    print('Task 8 done!')


if __name__ == '__main__':
    main()
