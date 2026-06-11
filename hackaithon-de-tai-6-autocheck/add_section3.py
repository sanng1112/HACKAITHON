#!/usr/bin/env python3
"""Add section 3. Architecture to AutoCheck proposal."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(lv, 14)
    r = p.add_run(t); sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def add_image(doc, img_path, width_cm=15, caption=None):
    if os.path.exists(img_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.add_picture(img_path, width=Cm(width_cm))
        if caption:
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(caption)
            sf(r2, s=11, b=True, c=RGBColor(0x66, 0x66, 0x66))


def ly(doc, b, n):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(b); sf(r, s=13, b=True)
    r = p.add_run(n); sf(r, s=13)


def bl(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('* '); sf(r, s=13)
    r = p.add_run(text); sf(r, s=13)


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(sd, 'proposal.docx'))
    ad = os.path.join(sd, 'assets')

    apb(doc)
    ah(doc, '3. THIET KE TONG QUAN')

    ah(doc, '3.1 Kien truc he thong', 2)
    ap(doc, 'AutoCheck duoc thiet ke theo mo hinh 3 tang, toi uu cho viec xu ly hang loat ho so giay to luu tru.')
    for b, n in [
        ('Tang 1 - Input Layer:', ' Tiep nhan tu may scan, upload PDF/JPEG, camera.'),
        ('Tang 2 - AI Processing:', ' SmartReader OCR, SmartVision, eKYC, AI Rules.'),
        ('Tang 3 - Output & Storage:', ' PostgreSQL, MinIO/S3, Redis.'),
    ]:
        ly(doc, b, n)

    arch_img = os.path.join(ad, 'architecture-diagram.png')
    add_image(doc, arch_img, width_cm=16, caption='Figure 1: So do kien truc tong the AutoCheck')

    ah(doc, '3.2 Giao dien nguoi dung', 2)
    ap(doc, 'AutoCheck co 2 giao dien chinh: Man hinh Scan (nhap lieu) va Dashboard Kiem tra & Xac thuc.')
    bl(doc, 'Man hinh Scan: Keo tha file, chon nguon scan, cai dat, xem preview.')
    bl(doc, 'Dashboard: Thong ke, danh sach ho so, chi tiet canh bao sai lech.')

    scan_img = os.path.join(ad, 'wireframe-scan-interface.png')
    add_image(doc, scan_img, width_cm=14, caption='Figure 2: Giao dien Scan')
    uf_img = os.path.join(ad, 'user-flow.png')
    add_image(doc, uf_img, width_cm=16, caption='Figure 3: Luong xu ly')
    dash_img = os.path.join(ad, 'wireframe-validation-dashboard.png')
    add_image(doc, dash_img, width_cm=16, caption='Figure 4: Dashboard xac thuc')

    ah(doc, '3.3 Quy trinh xu ly du lieu', 2)
    ap(doc, 'Pipeline: Scan -> Phan loai (SmartVision) -> OCR & Boc tach (SmartReader) -> Doi chieu (eKYC) -> Kiem tra (AI Rules) -> Xuat du lieu.')

    doc.save(os.path.join(sd, 'proposal.docx'))
    print('Task 4: Section "Architecture & Wireframe" added')


if __name__ == '__main__':
    main()
