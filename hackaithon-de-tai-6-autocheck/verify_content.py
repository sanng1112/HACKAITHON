#!/usr/bin/env python3
"""Verify AutoCheck proposal content."""
from docx import Document
import os

doc = Document(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposal.docx'))
print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
print()

print('=== Content ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'P{i}: {text[:120]}')

print()
print('=== Sections Check ===')
required = ['DAT VAN DE', 'GIAI PHAP', 'THIET KE', 'KHA THI', 'DOI MOI', 'TAC DONG', 'KET LUAN']
texts = [p.text.strip().upper() for p in doc.paragraphs if p.text.strip()]
for s in required:
    found = any(s in t for t in texts)
    print(f'  [{"OK" if found else "MISSING"}] Section "{s}"')

print()
if all(any(s in t for t in texts) for s in required):
    print('All 7 sections present! Proposal is complete.')
else:
    missing = [s for s in required if not any(s in t for t in texts)]
    print(f'Missing sections: {missing}')
