#!/usr/bin/env python3
"""Add section 2. GIAI PHAP AUTOCHECK to proposal."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16, 3: 14}.get(lv, 14)
    r = p.add_run(t); sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" '
            'w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def st(doc, bd, nml):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(bd); sf(r, s=13, b=True)
    r = p.add_run(nml); sf(r, s=13)


def comp(doc, bd, nml):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('* '); sf(r, s=13)
    r = p.add_run(bd); sf(r, s=13, b=True)
    r = p.add_run(nml); sf(r, s=13)


def mt(doc, hd, data):
    t = doc.add_table(rows=1+len(data), cols=len(hd))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(h)
        sf(r, s=11, b=True, c=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(
            parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT; r = p.add_run(ct); sf(r, s=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(
                    parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(sd, 'proposal.docx')
    doc = Document(doc_path)


    apb(doc)
    ah(doc, '2. GIAI PHAP AUTOCHECK')

    ah(doc, '2.1 Tong quan giai phap', 2)
    ap(doc, 'AutoCheck la mot he thong OCR thong minh cho phep so hoa hang loat ho so '
        'giay to luu tru tai cac co quan nha nuoc.')
    ap(doc, 'Khac voi giai phap OCR truyen thong, AutoCheck hieu ngu canh cua tung '
        'loai giay to (CCCD, so ho khau, giay khai sinh...) va tu dong dien thong tin.')

    ah(doc, '2.2 Tinh nang cot loi', 2)
    mt(doc, ['Tinh nang', 'Mo ta', 'Cong nghe VNPT'], [
        ['Scan & OCR', 'OCR nhan dang -> xuat text co cau truc', 'SmartReader (Doc AI)'],
        ['Phan loai tu dong', 'AI nhan dien loai giay to', 'SmartVision'],
        ['Boc tach thong tin', 'Trich xuat ho ten, CCCD, dia chi...', 'SmartReader'],
        ['Doi chieu & Xac thuc', 'So sanh voi CSDL hien co', 'eKYC (Compare)'],
        ['Kiem tra hop le', 'Phat hien sai lech, canh bao', 'SmartReader + Rules'],
        ['Xuat bao cao', 'Excel/JSON/CSV', '-'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '2.3 Quy trinh xu ly 6 buoc', 2)
    ap(doc, 'Quy trinh xu ly ho so gom 6 buoc:', ind=True)
    for b, n in [
        ('Buoc 1 - Nap ho so:', ' Dua ho so giay vao may scan hoac upload file.'),
        ('Buoc 2 - Phan loai:', ' SmartVision tu dong nhan dien loai giay to.'),
        ('Buoc 3 - OCR:', ' SmartReader OCR nhan dang, trich xuat cac truong.'),
        ('Buoc 4 - Doi chieu:', ' So sanh voi CSDL, danh dau sai lech.'),
        ('Buoc 5 - Kiem tra:', ' Can bo xem dashboard, xac nhan ket qua.'),
        ('Buoc 6 - Xuat du lieu:', ' Xuat ra CSDL, luu ban scan tren MinIO.'),
    ]:
        st(doc, b, n)

    ah(doc, '2.4 Vai tro cac thanh phan AI', 2)
    ap(doc, 'AutoCheck su dung 5 thanh phan AI tu VNPT:', ind=True)
    for b, n in [
        ('SmartReader (OCR): ', 'Nhan dang chu tieng Viet >95% .'),
        ('SmartReader (Doc AI): ', 'Boc tach thong tin co cau truc tu giay to.'),
        ('SmartVision: ', 'Phan loai giay to theo chung loai.'),
        ('eKYC (Compare): ', 'So sanh thong tin voi CSDL.'),
        ('eKYC (Liveness): ', 'Kiem tra anh chan dung that/gia.'),
    ]:
        comp(doc, b, n)

    doc.save(doc_path)
    print('Task 3: Section "Giai phap" added')


if __name__ == '__main__':
    main()
