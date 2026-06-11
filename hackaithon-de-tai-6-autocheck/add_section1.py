#!/usr/bin/env python3
"""Add section 1. DAT VAN DE to AutoCheck proposal."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n
    run.font.size = Pt(s)
    run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c:
        run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(lv, 14)
    r = p.add_run(t)
    sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" '
            'w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind:
        p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t)
    sf(r)
    return p


def bl(doc, bold_prefix, normal_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run('* ')
    sf(r, s=14, b=True)
    r = p.add_run(bold_prefix)
    sf(r, s=14, b=True)
    r = p.add_run(normal_text)
    sf(r, s=14)


def mt(doc, hd, data):
    t = doc.add_table(rows=1 + len(data), cols=len(hd))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        sf(r, s=11, b=True, c=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(
            parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct)
            sf(r, s=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(
                    parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(sd, 'proposal.docx')
    doc = Document(doc_path)

    apb(doc)
    ah(doc, '1. DAT VAN DE')

    ah(doc, '1.1 Boi canh', 2)
    ap(doc,
       'Chuyen doi so hanh chinh cong la nhiem vu trong tam cua Chinh phu giai '
       'doan 2026-2030. Mac du Cong Dich vu cong Quoc gia da dat hon 4.000 thu tuc '
       'hanh chinh truc tuyen, khoi luong ho so giay to luu tru tai cac co quan nha '
       'nuoc van con rat lon. Theo thong ke cua Bo Noi vu 2025, trung binh moi UBND '
       'cap phuong luu tru khoang 50.000-200.000 ho so giay to tu cac nam truoc, '
       'phan lon chua duoc so hoa.')

    ap(doc,
       'Viectra cuu, doi chieu thong tin tu nhung ho so giay to cu nay hoan toan '
       'dua vao thu cong - can bo phai luc tim ho so, doc va nhap lieu tung thong '
       'tin. Quy trinh nay mat nhieu thoi gian (trung binh 30-60 phut/ho so), de '
       'sai sot, va gay ach tac trong xu ly cong viec hanh chinh.')

    ah(doc, '1.2 Ba Pain-point chinh', 2)
    mt(doc,
       ['#', 'Pain-point', 'Minh chung', 'Doi tuong'],
       [
           ['PP1', 'Ton dong ho so giay chua so hoa',
            '~70% ho so luu tru tai UBND cap phuong chua so hoa (Bo TT&TT 2025)',
            'Can bo van thu, luu tru'],
           ['PP2', 'Tra cuu & doi chieu thu cong mat nhieu thoi gian',
            'Moi lan tra cuu ho so cu mat 30-60 phut (Khao sat UBND TP.HCM)',
            'Can bo mot cua, nguoi dan cho doi'],
           ['PP3', 'Rui ro sai sot, that lac, hu hong ho so giay',
            '~15% ho so giay sau 5 nam bi phai mo, rach, mat chu (Luu tru QG 2024)',
            'Co quan nha nuoc, nguoi dan'],
       ])

    p = doc.add_paragraph()
    sf(p.add_run(''), s=6)

    ah(doc, '1.3 Tai sao AI la giai phap?', 2)
    ap(doc, 'Ba pain-point tren deu co the giai quyet bang AI:', ind=True)
    bl(doc, 'PP1 -> ',
       'OCR & Document AI: SmartReader tu dong nhan dang va boc tach thong tin '
       'tu ho so giay, chuyen doi thanh du lieu so co cau truc.')
    bl(doc, 'PP2 -> ',
       'Intelligent Search: AI doi chieu thong tin trich xuat voi CSDL hien co, '
       'tu dong tim kiem va goi y ket qua.')
    bl(doc, 'PP3 -> ',
       'Auto-validate & Backup: AI kiem tra tinh hop le, phat hien sai lech, '
       'tu dong sao luu - loai bo rui ro mat mat du lieu.')

    ap(doc,
       'Cong nghe AI cua VNPT duoc chon vi SmartReader co kha nang OCR tieng Viet '
       'voi do chinh xac cao (>95%), nhan dang duoc nhieu loai giay to, va dap ung '
       'tieu chuan bao mat cua co quan nha nuoc.')

    ah(doc, '1.4 Tu van de den giai phap', 2)
    ap(doc,
       'Xuat phat tu thuc te do, chung toi de xuat AutoCheck - he thong OCR thong '
       'minh cho phep so hoa hang loat ho so giay to luu tru, tu dong boc tach thong '
       'tin, doi chieu voi CSDL va danh gia tinh hop le. AutoCheck ket hop 3 cong '
       'nghe AI cot loi cua VNPT: SmartReader, SmartVision va eKYC.')

    doc.save(doc_path)
    print('Task 2: Section "Dat van de" added')


if __name__ == '__main__':
    main()
