#!/usr/bin/env python3
"""
create_proposal.py - Generate proposal.docx cover page for AutoCheck (HACKAITHON 2026)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

WORK_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(WORK_DIR, "assets")
OUTPUT_PATH = os.path.join(WORK_DIR, "proposal.docx")
LOGO_PATH = os.path.join(ASSETS_DIR, "logo-autocheck.png")


def set_cell_border(cell, **kwargs):
    """Set cell border properties (optional utility)."""
    pass


def add_cover_page(doc):
    """Build the cover page content."""
    # ── Logo (if exists) ──
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.space_after = Pt(6)
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Cm(3.5), height=Cm(3.5))

    # ── Blank spacer ──
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(0)
    spacer.space_after = Pt(6)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_before = Pt(0)
    title_para.space_after = Pt(6)
    run = title_para.add_run("DỰ THI HACKATHON ĐỔI MỚI SÁNG TẠO 2026")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.name = "Times New Roman"

    # ── Topic ──
    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_para.space_before = Pt(4)
    topic_para.space_after = Pt(4)
    run = topic_para.add_run(
        "Đề tài 6: Ứng dụng trí tuệ nhân tạo (AI) trong việc số hóa và quản lý\n"
        "hồ sơ lưu trữ tại các cơ quan hành chính nhà nước"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = "Times New Roman"

    # ── Spacer ──
    spacer2 = doc.add_paragraph()
    spacer2.space_before = Pt(6)
    spacer2.space_after = Pt(6)

    # ── Product name ──
    prod_para = doc.add_paragraph()
    prod_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prod_para.space_before = Pt(0)
    prod_para.space_after = Pt(4)
    run = prod_para.add_run("AutoCheck")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.name = "Times New Roman"

    # ── Subtitle ──
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.space_before = Pt(0)
    sub_para.space_after = Pt(12)
    run = sub_para.add_run("Hệ thống OCR & Xử lý Hồ sơ Lưu trữ Thông minh")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = "Times New Roman"

    # ── Horizontal rule (line) ──
    hr_para = doc.add_paragraph()
    hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr_para.space_before = Pt(4)
    hr_para.space_after = Pt(12)
    run = hr_para.add_run("─" * 60)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.name = "Times New Roman"

    # ── Team info ──
    team_info = [
        ("Bảng thi:", "Bảng B (Challenger)"),
        ("Đội thi:", "[Tên đội]"),
        ("Thành viên 1:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 2:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 3:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 4:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 5:", "[Họ tên] — [Vai trò]"),
        ("Ngày nộp:", "16/06/2026"),
    ]

    for label, value in team_info:
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.space_before = Pt(1)
        info_para.space_after = Pt(1)

        run_label = info_para.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(14)
        run_label.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run_label.font.name = "Times New Roman"

        run_value = info_para.add_run(value)
        run_value.font.size = Pt(14)
        run_value.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run_value.font.name = "Times New Roman"


def create_proposal():
    doc = Document()

    # ── Page setup: A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # ── Set default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # ── Add cover page ──
    add_cover_page(doc)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Proposal saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    create_proposal()
