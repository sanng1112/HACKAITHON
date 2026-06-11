#!/usr/bin/env python3
"""Add section 4. TINH KHA THI to AutoCheck proposal."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def sf(run, n='Times New Roman', s=14, b=False, c=None):
    run.font.name = n; run.font.size = Pt(s); run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), n)
    if c: run.font.color.rgb = c


def apb(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)


def ah(doc, t, lv=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16, 3: 14}.get(lv, 14)
    r = p.add_run(t); sf(r, s=sz, b=True, c=RGBColor(0x00, 0x66, 0xCC))
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml('<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)


def ap(doc, t, ind=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if ind: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(t); sf(r); return p


def mt(doc, hd, data):
    t = doc.add_table(rows=1+len(data), cols=len(hd))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(h)
        sf(r, s=11, b=True, c=RGBColor(0xFF,0xFF,0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT; r = p.add_run(ct); sf(r, s=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))


def main():
    sd = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(sd, 'proposal.docx'))
    apb(doc); ah(doc, '4. TINH KHA THI')

    ah(doc, '4.1 Nguon du lieu', 2)
    ap(doc, 'AutoCheck su dung du lieu tu nhieu nguon:')
    mt(doc, ['Yeu to', 'Mo ta'], [
        ['Ho so giay luu tru', 'Co san tai UBND, moi phuong ~50.000-200.000 ho so.'],
        ['Du lieu huan luyen', 'VNPT SmartReader da huan luyen san tieng Viet.'],
        ['CSDL doi chieu', 'Cong DVC Quoc gia, CSDL Quoc gia ve dan cu.'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '4.2 Nhan luc', 2)
    mt(doc, ['Vai tro', 'SL', 'Ky nang chinh'], [
        ['Project Manager', '1', 'Agile/Scrum'],
        ['AI/OCR Developer', '2', 'Python, OCR, Xu ly anh'],
        ['Fullstack Developer', '1', 'React, FastAPI, PostgreSQL'],
        ['UI/UX Designer', '1', 'Figma'],
        ['BA', '1', 'Nghiep vu luu tru, mot cua'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '4.3 Kien truc ky thuat', 2)
    ap(doc, 'Frontend: React + TypeScript. Backend: Python FastAPI. '
        'AI: VNPT API. Database: PostgreSQL, MinIO/S3, Redis. DevOps: Docker, CI/CD.')

    ah(doc, '4.4 Ke hoach MVP 7 ngay (Vong 2)', 2)
    mt(doc, ['Ngay', 'Cong viec', 'Ket qua'], [
        ['1-2', 'Setup + SmartReader OCR', 'OCR anh -> text'],
        ['3-4', 'SmartVision + eKYC', 'Phan loai + doi chieu'],
        ['5-6', 'UI Scan + Dashboard', 'Giao dien hoan chinh'],
        ['7', 'E2E test + Fix + Dong goi', 'MVP deploy duoc'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '4.5 Chi phi van hanh', 2)
    mt(doc, ['Hang muc', 'VNĐ/thang', 'Ghi chu'], [
        ['Server (2 VPS 8GB)', '~2.000.000', 'AWS/VNPT Cloud'],
        ['API VNPT', '~1.000.000-3.000.000', 'Tuy so luong ho so'],
        ['MinIO/S3 Storage', '~500.000', 'Luu scan + du lieu'],
        ['Domain + SSL', '~200.000', '.gov.vn'],
        ['DevOps tools', 'Mien phi', 'GitHub/Docker Free'],
        ['Tong', '~3.700.000-5.700.000', '~$150-230/thang'],
    ])
    ap(doc, 'Chi phi setup: 20-30 trieu (may scan ong). Tiet kiem ~60% so voi nhan cong.')

    ah(doc, '4.6 An toan bao mat & Phap ly', 2)
    mt(doc, ['Yeu cau', 'Giai phap'], [
        ['Bao ve du lieu CN', 'Nghi dinh 13/2023 - AES-256, TLS 1.3'],
        ['So hoa tai lieu', 'Thong tu 01/2019/TT-BNV'],
        ['An toan TT', 'Luat ATTT 2015 - Audit log'],
        ['Luu tru so', 'Luat Luu tru 2011'],
    ])
    p = doc.add_paragraph(); sf(p.add_run(''), s=6)

    ah(doc, '4.7 Lo trinh phat trien', 2)
    for b, n in [
        ('Thang 1-2:', ' MVP -> Pilot 1 phuong, 5.000 ho so'),
        ('Thang 3-4:', ' Feedback -> Scale 5-10 quan/huyen'),
        ('Thang 5-6:', ' Tich hop CSDL QG -> Mo rong loai giay to'),
        ('Thang 7-12:', ' Trieu khai dien rong -> Hop tac VNPT'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(b); sf(r, s=13, b=True)
        r = p.add_run(n); sf(r, s=13)

    doc.save(os.path.join(sd, 'proposal.docx'))
    print('Task 5: Section "Tinh kha thi" added')


if __name__ == '__main__':
    main()
