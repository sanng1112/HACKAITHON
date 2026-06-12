#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_image(doc, img_path, width_cm=15, caption=None):
    if os.path.exists(img_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(); r.add_picture(img_path, width=Cm(width_cm))
        if caption: p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; r2 = p2.add_run(caption); set_font(r2, size=11, bold=True, color=RGBColor(0x66, 0x66, 0x66))

def add_layer(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(bold_part); set_font(r, size=13, bold=True); r = p.add_run(normal_part); set_font(r, size=13)

def add_bullet(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('• '); set_font(r, size=13); r = p.add_run(text); set_font(r, size=13)

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    doc_path = os.path.join(project_dir, 'proposal.docx')
    assets_dir = os.path.join(project_dir, 'assets')
    doc = Document(doc_path)
    add_page_break(doc)
    add_heading(doc, '3. THIẾT KẾ TỔNG QUAN')

    add_heading(doc, '3.1 Kiến trúc hệ thống', 2)
    add_para(doc, 'GovOne được thiết kế theo mô hình microservices 4 tầng, đảm bảo khả năng mở rộng và bảo trì độc lập giữa các thành phần. Mỗi tầng đảm nhận một vai trò riêng biệt và giao tiếp qua RESTful API/WebSocket.')
    for b,n in [('Tầng 1 — User Layer:',' Đa kênh: Kiosk touchscreen (voice-first), Web App, Mobile App. Giao diện tối giản, font lớn, tương phản cao — WCAG AA.'),
        ('Tầng 2 — AI Core (VNPT APIs):',' SmartVoice STT/TTS, Smartbot NLP, SmartReader OCR/Doc AI, eKYC OCR/Compare/Liveness, SmartVision Classification/Face/Sentiment.'),
        ('Tầng 3 — Processing Layer:',' Voice Gateway, Intent Engine, Document Processor, Validation Rules Engine, Sentiment Analyzer.'),
        ('Tầng 4 — Data Layer:',' PostgreSQL (giao dịch), Redis (cache), MinIO/S3 (scan gốc), Knowledge Base (thủ tục HC).')]:
        add_layer(doc,b,n)
    add_image(doc, os.path.join(assets_dir, 'architecture-diagram.png'), 16, 'Figure 1: Sơ đồ kiến trúc tổng thể GovOne (4 tầng)')

    add_heading(doc, '3.2 Giao diện người dùng', 2)
    add_para(doc, 'GovOne có 3 giao diện chính, thiết kế theo nguyên tắc tối giản, ưu tiên khả năng tiếp cận cho người già và người khuyết tật:')
    for t in ['Giao diện Kiosk: Voice-first, font ≥18px, nút ≥48x48px, tương phản 4.5:1 (WCAG AA).','Giao diện Scan OCR: Kéo thả file, chọn nguồn scan, xem preview. Dành cho cán bộ.','Dashboard cán bộ: Card KPI, biểu đồ xu hướng, bảng hồ sơ real-time, cảnh báo sai lệch.']:
        add_bullet(doc, t)
    add_image(doc, os.path.join(assets_dir, 'wireframe-kiosk.png'), 14, 'Figure 2: Wireframe giao diện Kiosk Voice-first')
    add_image(doc, os.path.join(assets_dir, 'wireframe-scan.png'), 14, 'Figure 3: Wireframe giao diện Scan OCR cho cán bộ')
    add_image(doc, os.path.join(assets_dir, 'wireframe-dashboard.png'), 16, 'Figure 4: Wireframe Dashboard cán bộ')

    add_heading(doc, '3.3 Luồng xử lý nghiệp vụ', 2)
    add_para(doc, 'GovOne vận hành 2 luồng xử lý song song:')
    add_layer(doc, 'Luồng Citizen (Voice-first): ','Người dân đến Kiosk → Camera phát hiện → Chào giọng nói → STT → Smartbot → Scan CCCD → eKYC → Xác nhận → TTS kết quả → Sentiment AI.')
    add_layer(doc, 'Luồng Officer (OCR): ','Scan hồ sơ → SmartVision phân loại → SmartReader OCR → eKYC đối chiếu → Rules Engine → Dashboard duyệt → Xuất CSDL/MinIO.')

    doc.save(doc_path)
    print('✅ Section 3: "Thiết kế tổng quan" added')

if __name__ == '__main__':
    main()
