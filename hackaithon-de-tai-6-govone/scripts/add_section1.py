#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_bullet(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('• '); set_font(r, size=14)
    r = p.add_run(bold_part); set_font(r, size=14, bold=True)
    r = p.add_run(normal_part); set_font(r, size=14)

def make_table(doc, headers, data):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct); set_font(r, size=10)
            if ri % 2 == 1:
                c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))
    return table

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(os.path.dirname(script_dir), 'proposal.docx'))

    add_page_break(doc)
    add_heading(doc, '1. ĐẶT VẤN ĐỀ')

    add_heading(doc, '1.1 Bối cảnh', 2)
    add_para(doc, 'Chuyển đổi số hành chính công là nhiệm vụ trọng tâm của Chính phủ giai đoạn 2026-2030. Cổng Dịch vụ công Quốc gia đã đạt hơn 4.000 thủ tục hành chính trực tuyến, nhưng tỷ lệ người dân sử dụng còn thấp (~30%) do rào cản công nghệ và giao diện phức tạp.')
    add_para(doc, 'Về phía cơ quan nhà nước, khối lượng hồ sơ giấy tờ lưu trữ tại các UBND cấp phường vẫn còn rất lớn — trung bình 50.000-200.000 hồ sơ/đơn vị, phần lớn chưa được số hóa (Bộ TT&TT 2025). Việc tra cứu, đối chiếu thông tin từ những hồ sơ giấy cũ hoàn toàn dựa vào thủ công, mất 30-60 phút/hồ sơ, dễ sai sót và gây ách tắc.')

    add_heading(doc, '1.2 Bốn Pain-point chính', 2)
    add_para(doc, 'Qua khảo sát thực tế tại các UBND phường và bộ phận một cửa, chúng tôi xác định 4 pain-point cốt lõi:', indent=False)
    make_table(doc,
        ['#', 'Pain-point', 'Minh chứng', 'Đối tượng'],
        [['PP1', 'Người dân gặp rào cản công nghệ — giao diện phức tạp, ngôn ngữ hành chính khó hiểu', '65% người >60 tuổi không tự tra cứu được thủ tục online (Bộ TT&TT 2025)', 'Người già, người khuyết tật'],
         ['PP2', 'Tồn đọng hồ sơ giấy chưa số hóa — tra cứu, đối chiếu thủ công', '~70% hồ sơ lưu trữ chưa số hóa; 30-60 phút/tra cứu (Bộ TT&TT 2025)', 'Cán bộ văn thư, lưu trữ'],
         ['PP3', 'Cán bộ một cửa quá tải — hướng dẫn lặp lại, kiểm tra hồ sơ thủ công', '1 cán bộ tiếp ~50-70 lượt/ngày, 60% là hướng dẫn thủ tục (UBND TP.HCM)', 'Cán bộ một cửa, người dân'],
         ['PP4', 'Rủi ro sai sót, thất lạc, hư hỏng hồ sơ giấy', '~15% hồ sơ giấy sau 5 năm bị phai mờ, rách, mất chữ (Lưu trữ QG 2024)', 'Cơ quan nhà nước, người dân']])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    add_heading(doc, '1.3 Tại sao AI là giải pháp?', 2)
    add_para(doc, 'Bốn pain-point trên đều có thể giải quyết bằng AI:', indent=False)
    add_bullet(doc, 'PP1 → ', 'Xử lý ngôn ngữ tự nhiên & Voice: SmartVoice STT giúp người dân nói thay vì gõ. Smartbot hướng dẫn thủ tục từng bước.')
    add_bullet(doc, 'PP2 → ', 'OCR & Document AI: SmartReader tự động nhận dạng và bóc tách thông tin từ hồ sơ giấy, chuyển đổi thành dữ liệu số.')
    add_bullet(doc, 'PP3 → ', 'Tự động hóa quy trình: AI xử lý câu hỏi lặp lại, tự động tra cứu và đối chiếu thông tin — giảm tải 40% cho cán bộ.')
    add_bullet(doc, 'PP4 → ', 'Auto-validate & Sentiment AI: AI kiểm tra tính hợp lệ, phát hiện sai lệch. Camera AI phân tích cảm xúc — đo hài lòng real-time.')
    add_para(doc, 'Công nghệ AI của VNPT được chọn vì: SmartReader OCR tiếng Việt >95%, SmartVoice STT/TTS đa vùng miền, eKYC đáp ứng tiêu chuẩn bảo mật nhà nước.')

    add_heading(doc, '1.4 Từ vấn đề đến giải pháp', 2)
    add_para(doc, 'Xuất phát từ thực tế đó, chúng tôi đề xuất GovOne — hệ thống quản lý hành chính công thông minh, tích hợp cả hai luồng: (1) Giao tiếp giọng nói cho người dân và (2) OCR & xử lý hồ sơ cho cán bộ. GovOne kết hợp 7 API AI của VNPT — SmartVoice, Smartbot, SmartReader, eKYC và SmartVision — trong một nền tảng thống nhất.')

    doc.save(os.path.join(os.path.dirname(script_dir), 'proposal.docx'))
    print('✅ Section 1: "Đặt vấn đề" added (4 pain-point)')

if __name__ == '__main__':
    main()