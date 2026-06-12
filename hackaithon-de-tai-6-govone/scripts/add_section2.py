#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16, 3: 14}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True, bold_prefix=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    if bold_prefix: r = p.add_run(bold_prefix); set_font(r, size=14, bold=True)
    r = p.add_run(text); set_font(r, size=14); return p
def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('• '); set_font(r, size=13)
    if bold_prefix: r = p.add_run(bold_prefix); set_font(r, size=13, bold=True)
    r = p.add_run(text); set_font(r, size=13)

def add_step(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(bold_part); set_font(r, size=13, bold=True)
    r = p.add_run(normal_part); set_font(r, size=13)

def make_table(doc, headers, data, col_aligns=None):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            if col_aligns and ci < len(col_aligns): p.alignment = col_aligns[ci]
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct); set_font(r, size=10)
            if ri % 2 == 1: c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))
    return table
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc = Document(os.path.join(os.path.dirname(script_dir), 'proposal.docx'))
    add_page_break(doc)
    add_heading(doc, '2. GIẢI PHÁP GOVONE')

    add_heading(doc, '2.1 Tổng quan giải pháp', 2)
    add_para(doc, 'GovOne là nền tảng quản lý hành chính công thông minh đầu tiên tại Việt Nam tích hợp 3 luồng AI trong một hệ thống: (1) Voice-first cho người dân — giao tiếp hoàn toàn bằng giọng nói; (2) OCR & Document AI cho cán bộ — số hóa và xử lý hồ sơ giấy tờ tự động; (3) Sentiment AI — đo lường mức độ hài lòng qua camera.')
    add_para(doc, 'Khác với các giải pháp đơn lẻ (chatbot text thuần túy hoặc OCR truyền thống), GovOne kết hợp cả hai trong một nền tảng duy nhất, cho phép người dân tương tác bằng giọng nói trong khi cán bộ xử lý hồ sơ qua dashboard OCR thông minh.')

    add_heading(doc, '2.2 Tính năng cốt lõi', 2)
    add_para(doc, 'GovOne tích hợp 6 tính năng chính, tận dụng 7 API VNPT:', indent=False)
    make_table(doc,
        ['Tính năng', 'Mô tả', 'API VNPT'],
        [['Voice Tra cứu & Khai báo', 'Người dân nói → STT → Smartbot xử lý → TTS trả lời', 'SmartVoice, Smartbot'],
         ['Scan & OCR thông minh', 'Nhận dạng, bóc tách thông tin từ CCCD, sổ hộ khẩu...', 'SmartReader, eKYC OCR'],
         ['Xác thực danh tính', 'So sánh khuôn mặt, phát hiện người thật/giả', 'eKYC Compare, Liveness'],
         ['Phân loại & Định tuyến', 'AI nhận diện loại giấy tờ, phân luồng xử lý', 'SmartVision Classification'],
         ['Đối chiếu & Kiểm tra', 'So sánh với CSDL, phát hiện sai lệch, cảnh báo', 'SmartReader + Rules Engine'],
         ['Đo lường hài lòng', 'Camera phân tích cảm xúc → báo cáo real-time', 'SmartVision Face/Sentiment']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)
    add_heading(doc, '2.3 Kịch bản người dùng: Câu chuyện Bác A', 2)
    add_para(doc, 'Để minh họa cách GovOne vận hành, xin giới thiệu kịch bản điển hình: Bác Nguyễn Văn A (65 tuổi) đến UBND phường làm thủ tục xác nhận tình trạng hôn nhân:', indent=False)
    for b, n in [('Bước 1 — Phát hiện:', ' Camera tại Kiosk GovOne phát hiện bác A → Hệ thống phát giọng chào.'),
        ('Bước 2 — Tra cứu:', ' Bác A nói yêu cầu → STT → Smartbot nhận diện ý định.'),
        ('Bước 3 — Hướng dẫn:', ' Smartbot xác định thủ tục → TTS hướng dẫn đưa CCCD vào khay scan.'),
        ('Bước 4 — Scan & OCR:', ' eKYC OCR + SmartReader tự động nhận dạng, bóc tách thông tin.'),
        ('Bước 5 — Xác thực:', ' eKYC Compare + Liveness: so sánh khuôn mặt + kiểm tra người thật.'),
        ('Bước 6 — Xác nhận:', ' TTS đọc lại thông tin, bác A xác nhận → tự động điền form.'),
        ('Bước 7 — Kết quả:', ' Hệ thống kiểm tra hợp lệ → TTS thông báo kết quả + mã hồ sơ.'),
        ('Bước 8 — Đo lường:', ' SmartVision phân tích cảm xúc → ghi nhận hài lòng → dashboard.')]:
        add_step(doc, b, n)
    add_para(doc, 'Toàn bộ quy trình chỉ mất 5-7 phút, giảm 70% thời gian so với 20-30 phút theo cách truyền thống. Bác A không cần chạm màn hình hay gõ phím.')
    add_heading(doc, '2.4 Luồng xử lý hồ sơ cho cán bộ', 2)
    add_para(doc, 'Song song với luồng giao tiếp dân, GovOne cung cấp quy trình xử lý hồ sơ 6 bước cho cán bộ:', indent=False)
    for b, n in [('Bước 1 — Nạp hồ sơ:', ' Scan hàng loạt giấy tờ hoặc upload file.'),
        ('Bước 2 — Phân loại:', ' SmartVision tự động nhận diện loại giấy tờ.'),
        ('Bước 3 — OCR & Bóc tách:', ' SmartReader OCR nhận dạng, trích xuất thông tin.'),
        ('Bước 4 — Đối chiếu CSDL:', ' So sánh với dữ liệu hiện có, đánh dấu sai lệch.'),
        ('Bước 5 — Kiểm tra & Duyệt:', ' Cán bộ xem dashboard, xác nhận kết quả.'),
        ('Bước 6 — Xuất dữ liệu:', ' Xuất ra CSDL, lưu bản scan trên MinIO/S3.')]:
        add_step(doc, b, n)

    add_heading(doc, '2.5 Vai trò các thành phần AI', 2)
    add_para(doc, 'GovOne sử dụng 7 thành phần AI từ VNPT:', indent=False)
    for b, n in [('SmartVoice STT: ', 'Chuyển giọng nói tiếng Việt thành văn bản. Hỗ trợ giọng địa phương.'),
        ('SmartVoice TTS: ', 'Chuyển văn bản thành giọng nói tự nhiên, thân thiện.'),
        ('Smartbot (NLP): ', 'Nhận diện ý định, tra cứu thủ tục, xử lý hội thoại đa lượt.'),
        ('SmartReader OCR: ', 'Nhận dạng ký tự quang học, bóc tách thông tin có cấu trúc.'),
        ('eKYC (Compare + Liveness): ', 'So sánh khuôn mặt, phát hiện người thật và giấy tờ thật/giả.'),
        ('SmartVision Classification: ', 'Phân loại giấy tờ theo chủng loại.'),
        ('SmartVision Face/Sentiment: ', 'Nhận diện cảm xúc khuôn mặt, đo lường hài lòng.')]:
        add_bullet(doc, n, bold_prefix=b)

    doc.save(os.path.join(os.path.dirname(script_dir), 'proposal.docx'))
    print('✅ Section 2: "Giải pháp GovOne" added (Voice + OCR + Sentiment)')

if __name__ == '__main__':
    main()
