#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_page_break(doc):
    p = doc.add_paragraph(); r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: 18, 2: 16}.get(level, 14)
    r = p.add_run(text); set_font(r, size=sz, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pPr.append(parse_xml(
            '<w:pBdr %s><w:bottom w:val="single" w:sz="8" w:space="1" w:color="0066CC"/></w:pBdr>' % nsdecls('w')))

def add_para(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); set_font(r, size=14); return p

def add_bullet(doc, bold_part, normal_part):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.line_spacing = 1.5
    r = p.add_run('\u2022 '); set_font(r, size=13)
    r = p.add_run(bold_part); set_font(r, size=13, bold=True)
    r = p.add_run(normal_part); set_font(r, size=13)

def make_table(doc, headers, data, col_aligns=None):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="0066CC" w:val="clear"/>' % nsdecls('w')))
    for ri, rd in enumerate(data):
        for ci, ct in enumerate(rd):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            if col_aligns and ci < len(col_aligns): p.alignment = col_aligns[ci]
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(ct); set_font(r, size=10)
            if ri % 2 == 1: c._tc.get_or_add_tcPr().append(parse_xml('<w:shd %s w:fill="E8F0FE" w:val="clear"/>' % nsdecls('w')))
    return table

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    doc_path = os.path.join(project_dir, 'proposal.docx')
    doc = Document(doc_path)

    add_page_break(doc)
    add_heading(doc, '6. T\u00c1C \u0110\u1ed8NG D\u1ef0 KI\u1ebeN')

    # 6.1 TAM-SAM-SOM
    add_heading(doc, '6.1 Ph\u00e2n t\u00edch th\u1ecb tr\u01b0\u1eddng TAM-SAM-SOM', 2)
    add_para(doc, 'Th\u1ecb tr\u01b0\u1eddng gi\u1ea3i ph\u00e1p h\u00e0nh ch\u00ednh c\u00f4ng th\u00f4ng minh t\u1ea1i Vi\u1ec7t Nam \u0111\u01b0\u1ee3c ph\u00e2n t\u00edch theo m\u00f4 h\u00ecnh TAM-SAM-SOM:')
    make_table(doc,
        ['Ch\u1ec9 s\u1ed1', 'Gi\u00e1 tr\u1ecb (VN\u0110)', 'C\u01a1 s\u1edf t\u00ednh to\u00e1n'],
        [['TAM (Total Addressable Market) \u2014 T\u1ed5ng th\u1ecb tr\u01b0\u1eddng', '18.000 t\u1ef7',
          'To\u00e0n b\u1ed9 chi ti\u00eau CNTT c\u1ee7a 63 t\u1ec9nh/th\u00e0nh + 705 qu\u1eadn/huy\u1ec7n + 10.599 x\u00e3/ph\u01b0\u1eddng. \u01af\u1edbc t\u00ednh 300 tri\u1ec7u \u0111\u1ed3ng/\u0111\u01a1n v\u1ecb/n\u0103m cho chuy\u1ec3n \u0111\u1ed5i s\u1ed1.'],
         ['SAM (Serviceable Addressable Market) \u2014 Th\u1ecb tr\u01b0\u1eddng ph\u1ee5c v\u1ee5', '800 t\u1ef7',
          'Ph\u00e2n kh\u00fac UBND ph\u01b0\u1eddng/x\u00e3 c\u00f3 nhu c\u1ea7u c\u1ea5p b\u00e1ch v\u1ec1 voice + OCR (5.000 \u0111\u01a1n v\u1ecb \u00d7 160 tri\u1ec7u \u0111\u1ed3ng). T\u1eadp trung TP.HCM, H\u00e0 N\u1ed9i, \u0110\u00e0 N\u1eb5ng, C\u1ea7n Th\u01a1.'],
         ['SOM (Serviceable Obtainable Market) \u2014 Th\u1ecb tr\u01b0\u1eddng \u0111\u1ea1t \u0111\u01b0\u1ee3c', '40 t\u1ef7',
          'M\u1ee5c ti\u00eau n\u0103m 1: 200 UBND \u00d7 200 tri\u1ec7u \u0111\u1ed3ng (ph\u1ea7n c\u1ee9ng + ph\u1ea7n m\u1ec1m + API). N\u0103m 2: 500 UBND. N\u0103m 3: 1.000 UBND.']],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 6.2 L\u1ee3i \u00edch x\u00e3 h\u1ed9i
    add_heading(doc, '6.2 L\u1ee3i \u00edch x\u00e3 h\u1ed9i', 2)
    add_para(doc, 'GovOne mang l\u1ea1i 7 t\u00e1c \u0111\u1ed9ng t\u00edch c\u1ef1c \u0111o l\u01b0\u1eddng \u0111\u01b0\u1ee3c cho x\u00e3 h\u1ed9i:')
    make_table(doc,
        ['Ch\u1ec9 s\u1ed1', 'Hi\u1ec7n t\u1ea1i', 'M\u1ee5c ti\u00eau GovOne', 'Ph\u01b0\u01a1ng ph\u00e1p \u0111o'],
        [['Th\u1eddi gian x\u1eed l\u00fd th\u1ee7 t\u1ee5c', '20-30 ph\u00fat', '5-7 ph\u00fat (gi\u1ea3m 70%)', 'Timing real-time t\u1eeb h\u1ec7 th\u1ed1ng'],
         ['\u0110\u1ed9 ph\u1ee7 ng\u01b0\u1eddi d\u00f9ng', '~30% d\u00e2n s\u1ed1', '>95% (bao g\u1ed3m ng\u01b0\u1eddi gi\u00e0, khuy\u1ebft t\u1eadt)', 'Kh\u1ea3o s\u00e1t sau giao d\u1ecbch'],
         ['T\u1ef7 l\u1ec7 h\u00e0i l\u00f2ng', '~65%', '>90%', 'Sentiment AI + Kh\u1ea3o s\u00e1t'],
         ['Th\u1eddi gian tra c\u1ee9u h\u1ed3 s\u01a1 c\u0169', '30-60 ph\u00fat', '<5 ph\u00fat (gi\u1ea3m 90%)', 'Log h\u1ec7 th\u1ed1ng OCR'],
         ['T\u1ef7 l\u1ec7 sai s\u00f3t h\u1ed3 s\u01a1', '~15%', '<2%', 'Audit \u0111\u1ecbnh k\u1ef3'],
         ['Chi ph\u00ed v\u1eadn h\u00e0nh/UBND/n\u0103m', '~200 tri\u1ec7u', '~50 tri\u1ec7u (gi\u1ea3m 75%)', 'B\u00e1o c\u00e1o t\u00e0i ch\u00ednh'],
         ['L\u01b0\u1ee3ng gi\u1ea5y t\u1edd l\u01b0u tr\u1eef v\u1eadt l\u00fd', '50.000-200.000 b\u1ed9', 'Gi\u1ea3m 80% (s\u1ed1 h\u00f3a to\u00e0n b\u1ed9)', 'Th\u1ed1ng k\u00ea scan h\u00e0ng th\u00e1ng']],
        col_aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 6.3 M\u00f4 h\u00ecnh doanh thu B2G
    add_heading(doc, '6.3 M\u00f4 h\u00ecnh doanh thu B2G', 2)
    add_para(doc, 'GovOne \u00e1p d\u1ee5ng m\u00f4 h\u00ecnh kinh doanh B2G (Business-to-Government) v\u1edbi 3 g\u00f3i d\u1ecbch v\u1ee5 linh ho\u1ea1t:')
    make_table(doc,
        ['G\u00f3i d\u1ecbch v\u1ee5', 'Gi\u00e1 (VN\u0110/th\u00e1ng)', 'Bao g\u1ed3m', '\u0110\u1ed1i t\u01b0\u1ee3ng'],
        [['Basic', '8.000.000',
          '\u00b7 Kiosk voice-first\u00b7 OCR 500 l\u01b0\u1ee3t/th\u00e1ng\u00b7 Smartbot 10 th\u1ee7 t\u1ee5c\u00b7 Dashboard c\u01a1 b\u1ea3n',
          'UBND ph\u01b0\u1eddng/x\u00e3 quy m\u00f4 nh\u1ecf'],
         ['Pro', '20.000.000',
          '\u00b7 Kiosk + Web App\u00b7 OCR 2.000 l\u01b0\u1ee3t/th\u00e1ng\u00b7 Smartbot 50 th\u1ee7 t\u1ee5c\u00b7 eKYC\u00b7 Dashboard n\u00e2ng cao',
          'UBND qu\u1eadn/huy\u1ec7n quy m\u00f4 trung b\u00ecnh'],
         ['Enterprise', 'Li\u00ean h\u1ec7',
          '\u00b7 T\u1ea5t c\u1ea3 t\u00ednh n\u0103ng\u00b7 OCR kh\u00f4ng gi\u1edbi h\u1ea1n\u00b7 Smartbot to\u00e0n b\u1ed9 th\u1ee7 t\u1ee5c\u00b7 T\u00edch h\u1ee3p CSDL s\u1edf/ban ng\u00e0nh\u00b7 SLA 99.9%',
          'S\u1edf TT&TT, UBND t\u1ec9nh']],
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    p = doc.add_paragraph(); set_font(p.add_run(''), size=6)

    # 6.4 Ph\u00e2n t\u00edch c\u1ea1nh tranh
    add_heading(doc, '6.4 Ph\u00e2n t\u00edch c\u1ea1nh tranh', 2)
    add_para(doc, 'Th\u1ecb tr\u01b0\u1eddng gi\u1ea3i ph\u00e1p h\u00e0nh ch\u00ednh c\u00f4ng th\u00f4ng minh t\u1ea1i Vi\u1ec7t Nam c\u00f3 5 nh\u00f3m \u0111\u1ed1i th\u1ee7 c\u1ea1nh tranh ch\u00ednh. GovOne c\u00f3 l\u1ee3i th\u1ebf kh\u00e1c bi\u1ec7t nh\u1edd t\u00edch h\u1ee3p Voice + OCR + Sentiment tr\u00ean c\u00f9ng m\u1ed9t n\u1ec1n t\u1ea3ng:')
    add_bullet(doc, 'FPT.AI \u2014 Smart Speech & OCR: ', '\u0110\u1ed1i th\u1ee7 m\u1ea1nh v\u1ec1 AI ti\u1ebfng Vi\u1ec7t. Tuy nhi\u00ean, FPT cung c\u1ea5p API \u0111\u01a1n l\u1ebb, kh\u00f4ng c\u00f3 gi\u1ea3i ph\u00e1p t\u00edch h\u1ee3p Kiosk + Dashboard ho\u00e0n ch\u1ec9nh. Kh\u00f4ng c\u00f3 Sentiment AI cho h\u00e0nh ch\u00ednh c\u00f4ng.')
    add_bullet(doc, 'VNPT eDoc \u2014 H\u1ed3 s\u01a1 \u0111i\u1ec7n t\u1eed: ', 'S\u1ea3n ph\u1ea9m c\u1ee7a VNPT t\u1eadp trung v\u00e0o l\u01b0u tr\u1eef h\u1ed3 s\u01a1, kh\u00f4ng c\u00f3 Voice-first. Kh\u00e1ch h\u00e0ng m\u1ee5c ti\u00eau l\u00e0 c\u1ea5p s\u1edf/ban ng\u00e0nh, kh\u00f4ng ph\u1ea3i UBND ph\u01b0\u1eddng.')
    add_bullet(doc, 'Zalo OA \u2014 Chatbot DVC: ', 'Gi\u1ea3i ph\u00e1p chatbot tr\u00ean Zalo ph\u1ed5 bi\u1ebfn nh\u01b0ng ch\u1ec9 h\u1ed7 tr\u1ee3 text, kh\u00f4ng c\u00f3 OCR, kh\u00f4ng c\u00f3 eKYC, kh\u00f4ng c\u00f3 Kiosk v\u1eadt l\u00fd. Ph\u1ee5 thu\u1ed9c v\u00e0o n\u1ec1n t\u1ea3ng Zalo.')
    add_bullet(doc, 'Google Document AI: ', 'OCR ti\u1ebfng Vi\u1ec7t t\u1ed1t nh\u01b0ng ch\u01b0a t\u1ed1i \u01b0u cho gi\u1ea5y t\u1edd h\u00e0nh ch\u00ednh Vi\u1ec7t Nam (CCCD, s\u1ed5 h\u1ed9 kh\u1ea9u). Kh\u00f4ng c\u00f3 voice, kh\u00f4ng \u0111\u00e1p \u1ee9ng Ngh\u1ecb \u0111\u1ecbnh 13/2023 v\u1ec1 l\u01b0u tr\u1eef d\u1eef li\u1ec7u trong n\u01b0\u1edbc.')
    add_bullet(doc, 'Gia c\u00f4ng CNTT (outsourcing): ', 'C\u00e1c c\u00f4ng ty gia c\u00f4ng x\u00e2y d\u1ef1ng gi\u1ea3i ph\u00e1p theo y\u00eau c\u1ea7u \u2014 chi ph\u00ed cao (300-500tr/d\u1ef1 \u00e1n), th\u1eddi gian d\u00e0i (3-6 th\u00e1ng), kh\u00f4ng c\u00f3 s\u1ea3n ph\u1ea9m chu\u1ea9n h\u00f3a, kh\u00f3 b\u1ea3o tr\u00ec.')

    doc.save(doc_path)
    print('\u2705 Section 6: \"T\u00e1c \u0111\u1ed9ng d\u1ef1 ki\u1ebfn\" added (TAM-SAM-SOM, l\u1ee3i \u00edch, doanh thu, c\u1ea1nh tranh)')

if __name__ == '__main__':
    main()
