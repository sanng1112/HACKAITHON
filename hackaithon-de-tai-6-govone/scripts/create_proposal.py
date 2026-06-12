#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORK_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(WORK_DIR, "assets")
OUTPUT_PATH = os.path.join(WORK_DIR, "proposal.docx")
LOGO_PATH = os.path.join(ASSETS_DIR, "logo-govone.png")

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color: run.font.color.rgb = color

def add_header(doc):
    for section in doc.sections:
        header = section.header; header.is_linked_to_previous = False
        hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO_PATH):
            run = hp.add_run(); run.add_picture(LOGO_PATH, width=Cm(1.5), height=Cm(1.5))
        run = hp.add_run('  GovOne — Hành chính công thông minh')
        set_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))

def add_cover_page(doc):
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(6)
        run = p.add_run(); run.add_picture(LOGO_PATH, width=Cm(4.0), height=Cm(4.0))
    doc.add_paragraph().space_after = Pt(6)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(6)
    run = p.add_run("DỰ THI HACKATHON ĐỔI MỚI SÁNG TẠO 2026")
    set_font(run, size=20, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    run = p.add_run("Đề tài 6: Ứng dụng AI nhằm nâng cao năng suất\nxử lý hồ sơ, thủ tục hành chính cho cơ quan nhà nước")
    set_font(run, size=14, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '12'); b.set(qn('w:space'), '1'); b.set(qn('w:color'), '0066CC')
    pBdr.append(b); pPr.append(pBdr)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    run = p.add_run("GovOne"); set_font(run, size=26, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(12)
    run = p.add_run("Hệ thống quản lý hành chính công thông minh\nVoice-first • OCR • eKYC • Sentiment AI")
    set_font(run, size=16, bold=True, color=RGBColor(0x00, 0x66, 0xCC))
    for label, value in [
        ("Bảng thi:", "Bảng B (Challenger)"), ("Đội thi:", "[Tên đội]"),
        ("Thành viên 1:", "[Họ tên] — [Vai trò]"), ("Thành viên 2:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 3:", "[Họ tên] — [Vai trò]"), ("Thành viên 4:", "[Họ tên] — [Vai trò]"),
        ("Thành viên 5:", "[Họ tên] — [Vai trò]"), ("Ngày nộp:", "16/06/2026"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(1)
        r1 = p.add_run(label + " "); set_font(r1, size=14, bold=True)
        r2 = p.add_run(value); set_font(r2, size=14)

def create_proposal():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    add_header(doc); add_cover_page(doc)
    doc.save(OUTPUT_PATH)
    print(f"✅ Proposal saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    create_proposal()
