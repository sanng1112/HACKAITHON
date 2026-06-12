import os, sys, unittest
from docx import Document

SCRIPT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(SCRIPT_DIR, 'scripts'))

class TestGovOne(unittest.TestCase):
    def setUp(self):
        self.project_dir = SCRIPT_DIR
        self.assets_dir = os.path.join(self.project_dir, 'assets')

    def test_assets_directory_exists(self):
        self.assertTrue(os.path.isdir(self.assets_dir))

    def test_logo_generated(self):
        logo_path = os.path.join(self.assets_dir, 'logo-govone.png')
        self.assertTrue(os.path.isfile(logo_path))
        from PIL import Image
        img = Image.open(logo_path)
        self.assertEqual(img.size, (400, 400))
        self.assertEqual(img.mode, 'RGBA')

    def test_proposal_created(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        self.assertTrue(os.path.isfile(proposal_path))
        doc = Document(proposal_path)
        self.assertGreaterEqual(len(doc.paragraphs), 10)

    def test_cover_page_has_title(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        doc = Document(proposal_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any('GovOne' in t for t in texts))

    def test_cover_page_has_topic(self):
        from docx import Document
        proposal_path = os.path.join(self.project_dir, 'proposal.docx')
        doc = Document(proposal_path)
        texts = [p.text for p in doc.paragraphs]
        self.assertTrue(any('Đề tài 6' in t for t in texts))

if __name__ == '__main__':
    unittest.main()
